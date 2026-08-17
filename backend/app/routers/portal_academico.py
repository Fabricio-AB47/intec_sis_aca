from datetime import date, datetime, timezone
from decimal import Decimal
from html import escape
from hashlib import sha256
from io import BytesIO
import json
import logging
from pathlib import Path
import re
from tempfile import TemporaryDirectory
import textwrap
import unicodedata
from typing import Annotated, Any, Literal
from zipfile import ZIP_DEFLATED, ZipFile

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook
from PIL import Image as PILImage
from pydantic import BaseModel, Field, ValidationError
import pyodbc
from starlette.concurrency import run_in_threadpool
from asn1crypto import pkcs12 as asn1_pkcs12
from cryptography import x509 as crypto_x509
from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.primitives.serialization import pkcs12 as crypto_pkcs12
from pyhanko.pdf_utils.incremental_writer import IncrementalPdfFileWriter
from pyhanko.pdf_utils.reader import PdfFileReader
from pyhanko.pdf_utils.text import TextBoxStyle
from pyhanko.sign import signers
from pyhanko.sign.fields import SigFieldSpec, SigSeedSubFilter
from pyhanko.stamp import QRStampStyle
from reportlab.graphics import renderPDF
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
from reportlab.lib.pagesizes import A4, landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import Flowable, Image as PdfImage, Indenter, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from svglib.svglib import svg2rlg

from app.core.security import SessionUser, require_roles
from app.services.db import get_connection, get_finance_connection
from app.services.graph_documents import (
    delete_item as delete_graph_document_item,
    ensure_folder as ensure_graph_document_folder,
    safe_folder_part as graph_safe_folder_part,
    upload_bytes as upload_graph_document_bytes,
)
from app.services.grade_calculation import (
    calculate_homologation_grade_with_recovery,
    calculate_regular_grade_with_recovery,
)

router = APIRouter(prefix="/api/portal", tags=["portal-academico"])
logger = logging.getLogger(__name__)

_STUDENT_ACCESS = require_roles("ESTUDIANTE")
_TEACHER_ACCESS = require_roles("DOCENTE")
_PORTAL_ADMIN_ACCESS = require_roles("ADMINISTRADOR", "ACADEMICO", "RECTOR")
_GRADES_ADMIN_ACCESS = require_roles("ADMINISTRADOR", "ACADEMICO", "BIENESTAR")
_BACKEND_ROOT = Path(__file__).resolve().parents[2]
_PROJECT_ROOT = _BACKEND_ROOT.parent
_REPORT_TEMPLATE_PATH = _PROJECT_ROOT / "frontend" / "doc" / "Plantilla word (1) - copia (1).docx"
_TEACHER_COMPLIANCE_WORD_TEMPLATE_PATH = Path.home() / "Documents" / "FABRICIO BORJA" / "FABRICIO BORJA CUMPLIMIENTO GRSI.docx"
_TEACHER_COMPLIANCE_BACKGROUND_PATH = _BACKEND_ROOT.parent / "backend" / ".codex_template_image1.png"
_LOGO_PATH = _PROJECT_ROOT / "frontend" / "public" / "Intec-Logowithslogangray.svg"
_ACADEMIC_PLANNING_PEA_BACKGROUND_PATH = _BACKEND_ROOT / "app" / "assets" / "academic_planning_pea_background.png"
_ACADEMIC_PLANNING_SYLLABUS_BACKGROUND_PATH = _BACKEND_ROOT / "app" / "assets" / "academic_planning_silabo_background.png"
_PORTAL_CONFIG_ROOT = _BACKEND_ROOT / "data"
_TEACHER_COMPLIANCE_FORMAT_PATH = _PORTAL_CONFIG_ROOT / "teacher_compliance_format.json"
_TEACHER_CONTRACT_STORAGE_ROOT = _BACKEND_ROOT / "private_uploads" / "teacher_contracts"
_TEACHER_CONTRACT_MAX_FILE_SIZE = 15 * 1024 * 1024
_TEACHER_CONTRACT_CERTIFICATE_MAX_FILE_SIZE = 2 * 1024 * 1024
_TEACHER_CONTRACT_MAX_PAGES = 60
_SIGNED_TEACHER_DOCUMENT_MAX_FILE_SIZE = 100 * 1024 * 1024
_TEACHER_DOCUMENT_ONEDRIVE_ROOT = "DOCENTES"

_CONTRACT_SPANISH_MONTHS = {
    "ENE": 1,
    "ENERO": 1,
    "FEB": 2,
    "FEBRERO": 2,
    "MAR": 3,
    "MARZO": 3,
    "ABR": 4,
    "ABRIL": 4,
    "MAY": 5,
    "MAYO": 5,
    "JUN": 6,
    "JUNIO": 6,
    "JUL": 7,
    "JULIO": 7,
    "AGO": 8,
    "AGOSTO": 8,
    "SEP": 9,
    "SEPT": 9,
    "SEPTIEMBRE": 9,
    "SETIEMBRE": 9,
    "OCT": 10,
    "OCTUBRE": 10,
    "NOV": 11,
    "NOVIEMBRE": 11,
    "DIC": 12,
    "DICIEMBRE": 12,
}


class TeacherGradePayload(BaseModel):
    codigo_estud: int
    cod_anio_basica: int
    codigo_periodo: int
    codigo_materia: int
    paralelo: str = Field(min_length=1, max_length=10)
    num_matricula: int | None = None
    num_grupo: int | None = None
    teoria_homo: float | None = Field(default=None, ge=0, le=10)
    practica_homo: float | None = Field(default=None, ge=0, le=10)
    p1_tareas: float | None = Field(default=None, ge=0, le=10)
    p1_proyectos: float | None = Field(default=None, ge=0, le=10)
    p1_examen: float | None = Field(default=None, ge=0, le=10)
    prom_p1: float | None = Field(default=None, ge=0, le=10)
    p2_tareas: float | None = Field(default=None, ge=0, le=10)
    p2_proyectos: float | None = Field(default=None, ge=0, le=10)
    p2_examen: float | None = Field(default=None, ge=0, le=10)
    prom_p2: float | None = Field(default=None, ge=0, le=10)
    p3_tareas: float | None = Field(default=None, ge=0, le=10)
    p3_proyectos: float | None = Field(default=None, ge=0, le=10)
    p3_examen: float | None = Field(default=None, ge=0, le=10)
    prom_p3: float | None = Field(default=None, ge=0, le=10)
    promedio: float | None = Field(default=None, ge=0, le=10)
    asistencia: float | None = Field(default=None, ge=0, le=100)
    recuperacion: float | None = Field(default=None, ge=0, le=10)
    promedio_final: float | None = Field(default=None, ge=0, le=10)
    caprueba: str | None = Field(default=None, max_length=10)


class AdminGradeCourseSelectionPayload(BaseModel):
    codigo_periodo: int
    cod_anio_basica: int
    codigo_materia: str = Field(min_length=1, max_length=100)
    paralelo: str = Field(min_length=1, max_length=10)
    cod_jornada: int | None = None


class AdminGradeCourseBatchPayload(BaseModel):
    courses: list[AdminGradeCourseSelectionPayload] = Field(min_length=1, max_length=3)


class AcademicPlanningTopicPayload(BaseModel):
    tema: str = Field(min_length=1, max_length=500)
    semana: int = Field(ge=1, le=52)
    horas_docencia: int = Field(default=0, ge=0, le=100)
    horas_practica: int = Field(default=0, ge=0, le=100)
    horas_autonomo: int = Field(default=0, ge=0, le=100)
    actividad_docencia: str = Field(default="", max_length=1000)
    actividad_practica: str = Field(default="", max_length=1000)
    actividad_autonoma: str = Field(default="", max_length=1000)
    evaluacion: str = Field(default="", max_length=500)


class AcademicPlanningUnitPayload(BaseModel):
    nombre: str = Field(min_length=1, max_length=300)
    resultado_aprendizaje: str = Field(default="", max_length=1500)
    temas: list[AcademicPlanningTopicPayload] = Field(default_factory=list, max_length=30)


class AcademicPlanningPayload(BaseModel):
    document_type: Literal["pea", "silabo"]
    codigo_periodos: list[int] = Field(min_length=1, max_length=4)
    codigo_materia: str = Field(min_length=1, max_length=100)
    paralelo: str = Field(min_length=1, max_length=10)
    cod_anio_basica: int | None = None
    cod_jornada: int | None = None
    nivel: str = Field(default="", max_length=100)
    unidad_curricular: str = Field(default="", max_length=150)
    campo_formacion: str = Field(default="", max_length=150)
    modalidad: str = Field(default="Presencial / En línea", max_length=150)
    prerrequisitos: str = Field(default="", max_length=500)
    correquisitos: str = Field(default="", max_length=500)
    horario_clases: str = Field(default="", max_length=500)
    horario_tutorias: str = Field(default="", max_length=500)
    descripcion: str = Field(default="", max_length=5000)
    objetivo_general: str = Field(default="", max_length=3000)
    resultados_aprendizaje: str = Field(default="", max_length=5000)
    mision_intec: str = Field(default="", max_length=3000)
    mision_escuela: str = Field(default="", max_length=3000)
    mision_carrera: str = Field(default="", max_length=3000)
    unidades: list[AcademicPlanningUnitPayload] = Field(min_length=1, max_length=12)
    estrategias_metodologicas: str = Field(default="", max_length=5000)
    formacion_ciudadana: str = Field(default="", max_length=3000)
    sostenibilidad: str = Field(default="", max_length=3000)
    recursos_didacticos: str = Field(default="", max_length=5000)
    evaluacion_tareas: int = Field(default=30, ge=0, le=100)
    evaluacion_individual: int = Field(default=15, ge=0, le=100)
    evaluacion_colaborativo: int = Field(default=15, ge=0, le=100)
    evaluacion_acumulativa: int = Field(default=40, ge=0, le=100)
    bibliografia_basica: str = Field(default="", max_length=5000)
    bibliografia_complementaria: str = Field(default="", max_length=5000)
    proyecto_tema: str = Field(default="", max_length=1000)
    proyecto_tiempo: str = Field(default="Un semestre", max_length=300)
    proyecto_objetivo: str = Field(default="", max_length=2000)
    proyecto_contexto: str = Field(default="", max_length=5000)
    version: str = Field(default="001", max_length=20)
    fecha_elaboracion: date = Field(default_factory=date.today)


class TeacherComplianceReportFormat(BaseModel):
    title: str = Field(default="REPORTE ACADÉMICO", max_length=180)
    pea_heading: str = Field(default="Cumplimiento del PEA y sílabo", max_length=180)
    pea_instruction: str = Field(
        default="Evidenciar el sílabo y PEA cargado en el sistema de aulas virtuales, debidamente firmado electrónicamente.",
        max_length=1000,
    )
    syllabus_update_heading: str = Field(default="Reporte de actualización del sílabo", max_length=180)
    syllabus_update_default: str = Field(default="Sin cambios realizados.", max_length=1000)
    virtual_classroom_heading: str = Field(default="Reporte del aula virtual", max_length=180)
    virtual_classroom_intro: str = Field(
        default="En el reporte consolidado se evidencia en el sistema de aulas virtuales que se cargaron los siguientes recursos en material académico:",
        max_length=1200,
    )
    resources: list[str] = Field(
        default_factory=lambda: [
            "Bibliografía del material académico",
            "Presentación PPT cargada como PDF por cada clase",
            "Link de grabaciones de cada clase o tutoría impartida",
            "Simulador de examen y banco de preguntas, para los casos que aplique",
            "Evaluación(es) teórica(s)",
            "Componente(s) práctico(s)",
        ]
    )
    teams_heading: str = Field(default="Evidencia de clases grabadas en TEAMS", max_length=180)
    attendance_heading: str = Field(default="Asistencias", max_length=180)
    grades_heading: str = Field(default="Reporte de notas", max_length=180)
    grades_instruction: str = Field(
        default="Se incluye resumen de nota máxima, nota mínima y casos reprobados según el reporte de notas registrado en el sistema académico.",
        max_length=1000,
    )
    annexes_heading: str = Field(default="Anexos", max_length=180)
    annexes_intro: str = Field(default="El presente informe debe ir acompañado de la siguiente documentación de respaldo:", max_length=1000)
    annexes: list[str] = Field(
        default_factory=lambda: [
            "Contrato firmado electrónicamente",
            "Reporte de notas firmado electrónicamente",
            "Factura electrónica emitida de acuerdo al número de contrato y valor",
        ]
    )
    closing: str = Field(default="Saludos cordiales,", max_length=300)
    signature_label: str = Field(default="Firma electrónica", max_length=120)
    signature_role: str = Field(default="DOCENTE", max_length=120)


class TeacherTeamsRecordingEvidence(BaseModel):
    team_id: str = Field(default="", max_length=100)
    team_name: str = Field(default="", max_length=300)
    recording_id: str = Field(default="", max_length=300)
    name: str = Field(min_length=1, max_length=300)
    date: str = Field(default="No disponible", max_length=60)
    start_hour: str = Field(default="No disponible", max_length=40)
    end_hour: str = Field(default="No disponible", max_length=40)
    call_duration: str = Field(default="No disponible", max_length=80)
    recording_duration: str = Field(default="No disponible", max_length=80)
    modified_by: str = Field(default="No disponible", max_length=200)
    web_url: str = Field(default="", max_length=2000)
    source: Literal["Microsoft Graph"] = "Microsoft Graph"


def _parse_teacher_teams_recordings(raw_payload: str) -> list[dict[str, Any]]:
    if not raw_payload.strip():
        return []
    try:
        decoded = json.loads(raw_payload)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail="La evidencia de Teams no contiene JSON válido") from exc
    if not isinstance(decoded, list):
        raise HTTPException(status_code=400, detail="La evidencia de Teams debe ser una lista")
    if len(decoded) > 50:
        raise HTTPException(status_code=400, detail="Solo se pueden anexar hasta 50 grabaciones de Teams")

    parsed: list[dict[str, Any]] = []
    try:
        for item in decoded:
            recording = TeacherTeamsRecordingEvidence.model_validate(item).model_dump()
            web_url = str(recording.get("web_url") or "").strip()
            recording["web_url"] = web_url if web_url.lower().startswith("https://") else ""
            parsed.append(recording)
    except (ValidationError, TypeError) as exc:
        raise HTTPException(status_code=400, detail="La evidencia de Teams contiene campos inválidos") from exc
    return parsed


_GRADE_COLUMN_MAP = {
    "teoria_homo": "teoriaHomo",
    "practica_homo": "practicahomo",
    "p1_tareas": "P1Tareas",
    "p1_proyectos": "P1Proyectos",
    "p1_examen": "P1Examen",
    "prom_p1": "promP1",
    "p2_tareas": "P2Tareas",
    "p2_proyectos": "P2Proyectos",
    "p2_examen": "P2Examen",
    "prom_p2": "promP2",
    "p3_tareas": "P3Tareas",
    "p3_proyectos": "P3Proyectos",
    "p3_examen": "P3Examen",
    "prom_p3": "promP3",
    "promedio": "Promedio",
    "asistencia": "Asistencia",
    "recuperacion": "Recuperacion",
    "promedio_final": "PromedioFinal",
    "caprueba": "caprueba",
}


def _default_teacher_compliance_format() -> dict[str, Any]:
    return TeacherComplianceReportFormat().model_dump()


def _sanitize_text_list(values: list[str]) -> list[str]:
    return [_clean(item) for item in values if _clean(item)]


def _read_teacher_compliance_format() -> dict[str, Any]:
    defaults = _default_teacher_compliance_format()
    if not _TEACHER_COMPLIANCE_FORMAT_PATH.exists():
        return defaults
    try:
        payload = json.loads(_TEACHER_COMPLIANCE_FORMAT_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return defaults
    if not isinstance(payload, dict):
        return defaults
    merged = {**defaults, **payload}
    merged["resources"] = _sanitize_text_list(merged.get("resources") or defaults["resources"]) or defaults["resources"]
    merged["annexes"] = _sanitize_text_list(merged.get("annexes") or defaults["annexes"]) or defaults["annexes"]
    return TeacherComplianceReportFormat(**merged).model_dump()


def _write_teacher_compliance_format(payload: TeacherComplianceReportFormat) -> dict[str, Any]:
    data = payload.model_dump()
    data["resources"] = _sanitize_text_list(data.get("resources") or [])
    data["annexes"] = _sanitize_text_list(data.get("annexes") or [])
    _PORTAL_CONFIG_ROOT.mkdir(parents=True, exist_ok=True)
    _TEACHER_COMPLIANCE_FORMAT_PATH.write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return data


def _image_evidence_flowables(
    images: list[dict[str, Any]],
    styles: dict[str, ParagraphStyle],
    max_width: float = 16.8 * cm,
    max_height: float = 10.5 * cm,
    show_labels: bool = False,
) -> list[Any]:
    flowables: list[Any] = []
    for item in images:
        content = item.get("content")
        if not content:
            continue
        try:
            image = PILImage.open(BytesIO(content))
            image.verify()
            image = PILImage.open(BytesIO(content))
        except Exception:
            continue
        width, height = image.size
        if width <= 0 or height <= 0:
            continue
        ratio = min(max_width / width, max_height / height, 1)
        label = _clean(item.get("label")) or "Captura de pantalla"
        buffer = BytesIO(content)
        flowables.append(Spacer(1, 0.08 * cm))
        if show_labels:
            flowables.append(Paragraph(f"<b>{_pdf_text(label)}</b>", styles["ComplianceSmall"]))
        flowable = PdfImage(buffer, width=width * ratio, height=height * ratio)
        flowable._evidence_buffer = buffer  # type: ignore[attr-defined]
        flowables.append(flowable)
        flowables.append(Spacer(1, 0.12 * cm))
    return flowables


def _clean(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\xa0", " ").strip()


def _date_text(value: Any) -> str:
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    return _clean(value)


def _number(value: Any) -> float | None:
    if value is None or value == "":
        return None
    if isinstance(value, Decimal):
        return float(value)
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


_PASSING_GRADE = 7.0


def _grade_result(value: Any) -> str:
    final_grade = _number(value)
    if final_grade is None:
        return "PENDIENTE"
    return "APROBADO" if final_grade >= _PASSING_GRADE else "REPROBADO"


def _weighted_regular_partial(tareas: Any, proyectos: Any, examen: Any) -> float | None:
    tareas_value = _number(tareas)
    proyectos_value = _number(proyectos)
    examen_value = _number(examen)
    if tareas_value is None or proyectos_value is None or examen_value is None:
        return None
    return round((tareas_value * 0.30) + (proyectos_value * 0.30) + (examen_value * 0.40), 2)


def _weighted_homologation_final(teoria: Any, practica: Any) -> float | None:
    teoria_value = _number(teoria)
    practica_value = _number(practica)
    if teoria_value is None or practica_value is None:
        return None
    return round((teoria_value * 0.40) + (practica_value * 0.60), 2)


def _int(value: Any) -> int | None:
    number = _number(value)
    return int(number) if number is not None else None


def _is_homologation_type(*values: Any) -> bool:
    text = " ".join(_clean(value).upper() for value in values)
    return "HOMO" in text or text in {"H", "HOMOLOGACION", "HOMOLOGADO"}


class _SvgLogo(Flowable):
    def __init__(self, path: Path, width: float) -> None:
        super().__init__()
        self.drawing = svg2rlg(str(path)) if path.exists() else None
        if self.drawing:
            self.scale = width / float(self.drawing.width or width)
            self.width = width
            self.height = float(self.drawing.height or 0) * self.scale
        else:
            self.scale = 1
            self.width = width
            self.height = 1.0 * cm

    def draw(self) -> None:
        if not self.drawing:
            self.canv.setFont("Helvetica-Bold", 22)
            self.canv.setFillColor(colors.HexColor("#808285"))
            self.canv.drawString(0, 0.25 * cm, "intec")
            return
        self.canv.saveState()
        self.canv.scale(self.scale, self.scale)
        renderPDF.draw(self.drawing, self.canv, 0, 0)
        self.canv.restoreState()


def _template_logo(width: float) -> Flowable:
    if _REPORT_TEMPLATE_PATH.exists():
        try:
            with ZipFile(_REPORT_TEMPLATE_PATH) as archive:
                with archive.open("word/media/image1.png") as source:
                    image = PILImage.open(source).convert("RGBA")
                    full_width, full_height = image.size
                    logo = image.crop(
                        (
                            int(full_width * 0.04),
                            int(full_height * 0.035),
                            int(full_width * 0.45),
                            int(full_height * 0.16),
                        )
                    )
                    buffer = BytesIO()
                    logo.save(buffer, format="PNG")
                    buffer.seek(0)
                    height = width * (logo.height / max(logo.width, 1))
                    flowable = PdfImage(buffer, width=width, height=height)
                    flowable._template_buffer = buffer  # type: ignore[attr-defined]
                    return flowable
        except Exception:
            pass
    return _SvgLogo(_LOGO_PATH, width)


def _template_page_image() -> bytes:
    with ZipFile(_REPORT_TEMPLATE_PATH) as archive:
        media_names = sorted(name for name in archive.namelist() if name.startswith("word/media/"))
        if not media_names:
            raise HTTPException(status_code=500, detail="La plantilla Word no contiene imagen base")
        return archive.read(media_names[0])


def _safe_filename(value: Any) -> str:
    text = _clean(value).lower()
    text = re.sub(r"[^a-z0-9._-]+", "-", text)
    return text.strip("-") or "reporte"


def _grade_status(final: Any, fallback: Any = "") -> str:
    value = _number(final)
    if value is not None:
        return "Aprobada" if value >= _PASSING_GRADE else "Reprobada"
    fallback_text = _clean(fallback)
    return fallback_text or "Pendiente"


def _grade_text(value: Any, decimals: int = 2) -> str:
    number = _number(value)
    if number is None:
        return "-"
    return f"{number:.{decimals}f}"


def _pdf_text(value: Any) -> str:
    text = _clean(value)
    return escape(text) if text else "-"


def _pdf_paragraph(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(_pdf_text(value), style)


def _report_styles() -> dict[str, ParagraphStyle]:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontSize=14,
            leading=17,
            textColor=colors.HexColor("#0c1f42"),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportSubtitle",
            parent=styles["BodyText"],
            alignment=TA_CENTER,
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#4d5a78"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Meta",
            parent=styles["BodyText"],
            fontSize=8.2,
            leading=10,
            textColor=colors.HexColor("#0c1f42"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="Cell",
            parent=styles["BodyText"],
            fontSize=6.4,
            leading=7.6,
            textColor=colors.HexColor("#0c1f42"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CellBold",
            parent=styles["Cell"],
            fontName="Helvetica-Bold",
        )
    )
    return styles


def _student_meta_table(profile: dict[str, Any], career: str, styles: dict[str, ParagraphStyle]) -> Table:
    data = [
        [
            Paragraph(f"<b>Estudiante:</b> {_pdf_text(profile.get('nombre_estudiante'))}", styles["Meta"]),
            Paragraph(f"<b>Cedula:</b> {_pdf_text(profile.get('cedula'))}", styles["Meta"]),
            Paragraph(f"<b>Codigo:</b> {_pdf_text(profile.get('codigo_estud'))}", styles["Meta"]),
        ],
        [
            Paragraph(f"<b>Correo INTEC:</b> {_pdf_text(profile.get('correo_intec'))}", styles["Meta"]),
            Paragraph(f"<b>Carrera:</b> {_pdf_text(career)}", styles["Meta"]),
            Paragraph(f"<b>Fecha:</b> {date.today().strftime('%d/%m/%Y')}", styles["Meta"]),
        ],
    ]
    table = Table(data, colWidths=[6.5 * cm, 5.8 * cm, 5.8 * cm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f5fafc")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#b7c8cf")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d5e0e5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _subject_code_text(item: dict[str, Any]) -> str:
    code = _clean(item.get("cod_materia"))
    internal = _clean(item.get("codigo_materia"))
    if code and internal and code != internal:
        return f"{escape(code)}<br/><font size=\"6.5\">{escape(internal)}</font>"
    return escape(code or internal or "-")


def _build_student_report_pdf(
    title: str,
    subtitle: str,
    profile: dict[str, Any],
    career: str,
    headers: list[str],
    rows: list[list[Any]],
    col_widths: list[float],
) -> bytes:
    if not _REPORT_TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="No se encontro la plantilla Word para generar el PDF")

    styles = _report_styles()
    template_reader = ImageReader(BytesIO(_template_page_image()))
    story: list[Any] = [
        Paragraph(title, styles["ReportTitle"]),
        Paragraph(subtitle, styles["ReportSubtitle"]),
        _student_meta_table(profile, career, styles),
        Spacer(1, 0.3 * cm),
    ]

    table_data: list[list[Any]] = [[Paragraph(f"<b>{escape(header)}</b>", styles["CellBold"]) for header in headers]]
    table_data.extend(rows)
    if len(table_data) == 1:
        table_data.append([Paragraph("No hay informacion para mostrar.", styles["Cell"]) for _ in headers])

    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e9f3f6")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0c1f42")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cad5dc")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fbfc")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)

    def draw_template(canvas: Any, _doc: Any) -> None:
        page_width, page_height = A4
        canvas.saveState()
        canvas.drawImage(template_reader, 0, 0, width=page_width, height=page_height, mask="auto")
        if hasattr(canvas, "setFillAlpha"):
            canvas.setFillAlpha(0.96)
        canvas.setFillColor(colors.white)
        canvas.roundRect(0.65 * cm, 1.1 * cm, page_width - 1.3 * cm, page_height - 4.9 * cm, 8, stroke=0, fill=1)
        canvas.restoreState()

    output = BytesIO()
    SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=0.75 * cm,
        leftMargin=0.75 * cm,
        topMargin=3.8 * cm,
        bottomMargin=1.3 * cm,
        title=title,
    ).build(story, onFirstPage=draw_template, onLaterPages=draw_template)
    output.seek(0)
    return output.getvalue()


def _academic_pdf_rows(items: list[dict[str, Any]]) -> list[list[Any]]:
    styles = _report_styles()
    cell = styles["Cell"]
    rows: list[list[Any]] = []
    for item in items:
        is_homo = _is_homologation_type(
            item.get("tipo_matricula"),
            item.get("ultimo_periodo"),
            item.get("esquema_calificacion"),
        )
        final = _number(item.get("promedio_final"))
        status = _grade_status(final, item.get("estado_academico"))
        rows.append(
            [
                _pdf_paragraph(item.get("semestre"), cell),
                Paragraph(_subject_code_text(item), cell),
                _pdf_paragraph(item.get("nombre_materia"), cell),
                _pdf_paragraph(_grade_text(item.get("creditos")), cell),
                _pdf_paragraph(item.get("esquema_calificacion") or ("HOMOLOGACION" if is_homo else "REGULAR"), cell),
                _pdf_paragraph(
                    f"T: {_grade_text(item.get('teoria_homo'))} / P: {_grade_text(item.get('practica_homo'))}"
                    if is_homo
                    else "-",
                    cell,
                ),
                _pdf_paragraph("-" if is_homo else _grade_text(item.get("prom_p1")), cell),
                _pdf_paragraph("-" if is_homo else _grade_text(item.get("prom_p2")), cell),
                _pdf_paragraph("-" if is_homo else _grade_text(item.get("prom_p3")), cell),
                _pdf_paragraph(_grade_text(final) if final is not None else "0", cell),
                _pdf_paragraph(status, cell),
            ]
        )
    return rows


def _calificaciones_pdf_rows(items: list[dict[str, Any]], homologation_only: bool = False) -> list[list[Any]]:
    styles = _report_styles()
    cell = styles["Cell"]
    rows: list[list[Any]] = []
    for item in items:
        is_homo = _is_homologation_type(
            item.get("tipo_matricula"),
            item.get("detalle_periodo"),
            item.get("esquema_calificacion"),
        )
        final = _number(item.get("promedio_final"))
        base_row = [
            _pdf_paragraph(item.get("semestre"), cell),
            _pdf_paragraph(item.get("detalle_periodo") or item.get("codigo_periodo"), cell),
            Paragraph(_subject_code_text(item), cell),
            _pdf_paragraph(item.get("nombre_materia"), cell),
            _pdf_paragraph(item.get("esquema_calificacion") or ("HOMOLOGACION" if is_homo else "REGULAR"), cell),
        ]
        if not homologation_only:
            base_row.extend(
                [
                    _pdf_paragraph("-" if is_homo else _grade_text(item.get("prom_p1")), cell),
                    _pdf_paragraph("-" if is_homo else _grade_text(item.get("prom_p2")), cell),
                    _pdf_paragraph("-" if is_homo else _grade_text(item.get("prom_p3")), cell),
                ]
            )
        base_row.extend(
            [
                _pdf_paragraph(_grade_text(final), cell),
                _pdf_paragraph(_grade_status(final, item.get("estado_academico")), cell),
            ]
        )
        rows.append(base_row)
    return rows


def _practice_search_text(value: Any) -> str:
    text = _clean(value).upper()
    replacements = {
        "Á": "A",
        "É": "E",
        "Í": "I",
        "Ó": "O",
        "Ú": "U",
        "Ü": "U",
        "Ñ": "N",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text


def _practice_requirement_code(item: dict[str, Any]) -> str:
    text = _practice_search_text(
        " ".join(
            [
                _clean(item.get("nombre_materia")),
                _clean(item.get("cod_materia")),
                _clean(item.get("codigo_materia")),
            ]
        )
    )
    if "VINCULACION" in text or "SERVICIO COMUNITARIO" in text:
        return "VIN"
    if "PRACTICA" in text:
        return "PPF"
    return ""


def _academic_grid_with_practice_requirements(items: list[dict[str, Any]], career: str = "") -> list[dict[str, Any]]:
    result = [dict(item) for item in items]
    existing_codes = {_practice_requirement_code(item) for item in result}
    defaults = [
        {
            "code": "PPF",
            "cod_materia": "PPF-240",
            "codigo_materia": "PPF-240",
            "nombre_materia": "PRÁCTICAS PREPROFESIONALES - 240 HORAS",
            "horas": 240,
        },
        {
            "code": "VIN",
            "cod_materia": "VIN-060",
            "codigo_materia": "VIN-060",
            "nombre_materia": "VINCULACIÓN - 60 HORAS",
            "horas": 60,
        },
    ]
    for default in defaults:
        if default["code"] in existing_codes:
            continue
        result.append(
            {
                "semestre": 3,
                "orden": 99980 if default["code"] == "VIN" else 99970,
                "cod_materia": default["cod_materia"],
                "codigo_materia": default["codigo_materia"],
                "nombre_materia": default["nombre_materia"],
                "nombre_carrera": career,
                "creditos": None,
                "horas": default["horas"],
                "esquema_calificacion": "REQUISITO",
                "estado_academico": "Pendiente",
                "promedio_final": None,
                "nota_aprobar": 7,
                "ultimo_periodo": "Requisito institucional",
                "codigo_periodo": "PRACTICAS",
                "detalle_periodo": "Requisito institucional",
            }
        )
    return result


def _student_profile_from_row(row: Any) -> dict[str, Any]:
    return {
        "codigo_estud": _clean(getattr(row, "codigo_estud", "")),
        "cedula": _clean(getattr(row, "cedula", "")),
        "nombre_estudiante": _clean(getattr(row, "nombre_estudiante", "")),
        "correo_personal": _clean(getattr(row, "correo_personal", "")),
        "correo_intec": _clean(getattr(row, "correo_intec", "")),
        "estado_codigo": _clean(getattr(row, "estado_codigo", "")),
    }


def _fetch_student_profile(cursor: pyodbc.Cursor, codigo_estud: int) -> dict[str, Any]:
    cursor.execute(
        """
        SELECT TOP (1)
            TRY_CONVERT(varchar(50), de.codigo_estud) AS codigo_estud,
            TRY_CONVERT(nvarchar(100), de.Cedula_Est) AS cedula,
            TRY_CONVERT(nvarchar(4000), de.Apellidos_nombre) AS nombre_estudiante,
            TRY_CONVERT(nvarchar(255), de.correo) AS correo_personal,
            COALESCE(
                NULLIF(TRY_CONVERT(nvarchar(255), ce.CorreoIntec), N''),
                TRY_CONVERT(nvarchar(255), de.correointec)
            ) AS correo_intec,
            TRY_CONVERT(nvarchar(50), de.Estado) AS estado_codigo
        FROM dbo.DATOS_ESTUD de
        LEFT JOIN dbo.CorreosEstudIntec ce
          ON TRY_CONVERT(int, ce.codestud) = TRY_CONVERT(int, de.codigo_estud)
        WHERE TRY_CONVERT(int, de.codigo_estud) = ?
        """,
        codigo_estud,
    )
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="No se encontro el estudiante vinculado a la sesion")
    return _student_profile_from_row(row)


def _record_item(row: Any) -> dict[str, Any]:
    es_homologacion = _is_homologation_type(
        getattr(row, "tipo_matricula", ""),
        getattr(row, "detalle_periodo", ""),
    )
    nota_final = _number(getattr(row, "nota_final", None))
    if nota_final is None:
        nota_final = _number(getattr(row, "promedio_final_raw", None))
    nota_aprobar = _PASSING_GRADE
    estado_nota = _grade_result(nota_final)
    if estado_nota == "PENDIENTE":
        estado = "En curso"
        aprobada = False
    elif estado_nota == "APROBADO":
        estado = "Aprobada"
        aprobada = True
    else:
        estado = "Reprobada"
        aprobada = False

    return {
        "codigo_estud": _clean(row.codigo_estud),
        "cod_anio_basica": _clean(row.cod_anio_basica),
        "nombre_carrera": _clean(row.nombre_carrera),
        "codigo_periodo": _clean(row.codigo_periodo),
        "detalle_periodo": _clean(row.detalle_periodo),
        "anio_periodo": _int(getattr(row, "anio_periodo", None)),
        "codigo_materia": _clean(row.codigo_materia),
        "cod_materia": _clean(row.cod_materia),
        "nombre_materia": _clean(row.nombre_materia),
        "semestre": _int(row.semestre),
        "creditos": _number(row.creditos),
        "horas": _number(getattr(row, "horas", None)),
        "orden": _int(getattr(row, "orden", None)),
        "num_malla": _int(getattr(row, "num_malla", None)),
        "paralelo": _clean(row.paralelo),
        "num_grupo": _int(row.num_grupo),
        "num_matricula": _clean(row.num_matricula),
        "fecha_matricula": _date_text(row.fecha_matricula),
        "tipo_matricula": _clean(row.tipo_matricula),
        "es_homologacion": es_homologacion,
        "esquema_calificacion": "HOMOLOGACION" if es_homologacion else "REGULAR",
        "teoria_homo": _number(getattr(row, "teoria_homo", None)),
        "practica_homo": _number(getattr(row, "practica_homo", None)),
        "p1_tareas": _number(row.p1_tareas),
        "p1_proyectos": _number(row.p1_proyectos),
        "p1_examen": _number(row.p1_examen),
        "prom_p1": _number(row.prom_p1),
        "p2_tareas": _number(row.p2_tareas),
        "p2_proyectos": _number(row.p2_proyectos),
        "p2_examen": _number(row.p2_examen),
        "prom_p2": _number(row.prom_p2),
        "p3_tareas": _number(row.p3_tareas),
        "p3_proyectos": _number(row.p3_proyectos),
        "p3_examen": _number(row.p3_examen),
        "prom_p3": _number(row.prom_p3),
        "promedio": _number(row.promedio),
        "asistencia": _number(row.asistencia),
        "recuperacion": _number(row.recuperacion),
        "promedio_final": nota_final,
        "nota_aprobar": nota_aprobar,
        "aprobada": aprobada,
        "estado_nota": estado_nota,
        "estado_academico": estado,
        "observaciones": _clean(getattr(row, "observaciones", "")),
        "seguimiento": _clean(getattr(row, "seguimiento", "")),
    }


def _fetch_student_record(cursor: pyodbc.Cursor, codigo_estud: int) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    cursor.execute(
        """
        SELECT TOP (3000)
            TRY_CONVERT(varchar(50), de.codigo_estud) AS codigo_estud,
            TRY_CONVERT(nvarchar(100), de.Cedula_Est) AS cedula,
            TRY_CONVERT(nvarchar(4000), de.Apellidos_nombre) AS nombre_estudiante,
            TRY_CONVERT(nvarchar(255), de.correo) AS correo_personal,
            COALESCE(
                NULLIF(TRY_CONVERT(nvarchar(255), ce.CorreoIntec), N''),
                TRY_CONVERT(nvarchar(255), de.correointec)
            ) AS correo_intec,
            TRY_CONVERT(nvarchar(50), de.Estado) AS estado_codigo,
            TRY_CONVERT(varchar(50), cxe.cod_anio_Basica) AS cod_anio_basica,
            TRY_CONVERT(nvarchar(4000), c.Nombre_Basica) AS nombre_carrera,
            TRY_CONVERT(varchar(50), cxe.codigo_periodo) AS codigo_periodo,
            TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo) AS detalle_periodo,
            TRY_CONVERT(int, pe.anio) AS anio_periodo,
            TRY_CONVERT(varchar(50), cxe.codigo_materia) AS codigo_materia,
            TRY_CONVERT(varchar(100), p.cod_materia) AS cod_materia,
            TRY_CONVERT(nvarchar(4000), p.Nomb_Materia) AS nombre_materia,
            TRY_CONVERT(int, p.Semestre) AS semestre,
            TRY_CONVERT(int, p.NumMalla) AS num_malla,
            TRY_CONVERT(float, p.Horas) AS horas,
            TRY_CONVERT(int, p.Orden) AS orden,
            TRY_CONVERT(float, COALESCE(NULLIF(cxe.Num_Creditos, 0), p.Creditos)) AS creditos,
            TRY_CONVERT(nvarchar(50), cxe.paralelo) AS paralelo,
            TRY_CONVERT(int, cxe.NumGrupo) AS num_grupo,
            TRY_CONVERT(varchar(50), cxe.Num_Matricula) AS num_matricula,
            cxe.Fecha_Matricula AS fecha_matricula,
            TRY_CONVERT(nvarchar(20), cxe.TipoMatricula) AS tipo_matricula,
            TRY_CONVERT(float, cxe.teoriaHomo) AS teoria_homo,
            TRY_CONVERT(float, cxe.practicahomo) AS practica_homo,
            TRY_CONVERT(float, cxe.P1Tareas) AS p1_tareas,
            TRY_CONVERT(float, cxe.P1Proyectos) AS p1_proyectos,
            TRY_CONVERT(float, cxe.P1Examen) AS p1_examen,
            TRY_CONVERT(float, cxe.promP1) AS prom_p1,
            TRY_CONVERT(float, cxe.P2Tareas) AS p2_tareas,
            TRY_CONVERT(float, cxe.P2Proyectos) AS p2_proyectos,
            TRY_CONVERT(float, cxe.P2Examen) AS p2_examen,
            TRY_CONVERT(float, cxe.promP2) AS prom_p2,
            TRY_CONVERT(float, cxe.P3Tareas) AS p3_tareas,
            TRY_CONVERT(float, cxe.P3Proyectos) AS p3_proyectos,
            TRY_CONVERT(float, cxe.P3Examen) AS p3_examen,
            TRY_CONVERT(float, cxe.promP3) AS prom_p3,
            TRY_CONVERT(float, cxe.Promedio) AS promedio,
            TRY_CONVERT(float, cxe.Asistencia) AS asistencia,
            TRY_CONVERT(float, cxe.Recuperacion) AS recuperacion,
            TRY_CONVERT(float, cxe.PromedioFinal) AS promedio_final_raw,
            COALESCE(
                TRY_CONVERT(float, cxe.PromedioFinal),
                CASE
                    WHEN (
                            UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(50), cxe.TipoMatricula), N'')))) = N'H'
                         OR UPPER(COALESCE(TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo), N'')) LIKE N'%HOMO%'
                         )
                     AND TRY_CONVERT(float, cxe.teoriaHomo) IS NOT NULL
                     AND TRY_CONVERT(float, cxe.practicahomo) IS NOT NULL
                    THEN (TRY_CONVERT(float, cxe.teoriaHomo) * 0.4) + (TRY_CONVERT(float, cxe.practicahomo) * 0.6)
                END,
                CASE
                    WHEN TRY_CONVERT(float, cxe.promP1) IS NOT NULL
                     AND TRY_CONVERT(float, cxe.promP2) IS NOT NULL
                     AND TRY_CONVERT(float, cxe.promP3) IS NOT NULL
                    THEN (TRY_CONVERT(float, cxe.promP1) + TRY_CONVERT(float, cxe.promP2) + TRY_CONVERT(float, cxe.promP3)) / 3
                END,
                TRY_CONVERT(float, cxe.Promedio),
                TRY_CONVERT(float, cxe.PromedioAux)
            ) AS nota_final,
            COALESCE(TRY_CONVERT(float, pe.NotaAprobar), 7) AS nota_aprobar,
            CASE
                WHEN UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(50), cxe.caprueba), N'')))) LIKE N'A%' THEN 1
                WHEN COALESCE(
                        TRY_CONVERT(float, cxe.PromedioFinal),
                        CASE
                            WHEN (
                                    UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(50), cxe.TipoMatricula), N'')))) = N'H'
                                 OR UPPER(COALESCE(TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo), N'')) LIKE N'%HOMO%'
                                 )
                             AND TRY_CONVERT(float, cxe.teoriaHomo) IS NOT NULL
                             AND TRY_CONVERT(float, cxe.practicahomo) IS NOT NULL
                            THEN (TRY_CONVERT(float, cxe.teoriaHomo) * 0.4) + (TRY_CONVERT(float, cxe.practicahomo) * 0.6)
                        END,
                        CASE
                            WHEN TRY_CONVERT(float, cxe.promP1) IS NOT NULL
                             AND TRY_CONVERT(float, cxe.promP2) IS NOT NULL
                             AND TRY_CONVERT(float, cxe.promP3) IS NOT NULL
                            THEN (TRY_CONVERT(float, cxe.promP1) + TRY_CONVERT(float, cxe.promP2) + TRY_CONVERT(float, cxe.promP3)) / 3
                        END,
                        TRY_CONVERT(float, cxe.Promedio),
                        TRY_CONVERT(float, cxe.PromedioAux)
                     ) >= COALESCE(TRY_CONVERT(float, pe.NotaAprobar), 7)
                THEN 1
                ELSE 0
            END AS aprobada,
            TRY_CONVERT(nvarchar(max), cxe.observaciones) AS observaciones,
            TRY_CONVERT(nvarchar(255), cxe.seguimiento) AS seguimiento
        FROM dbo.CARRERAXESTUD cxe
        INNER JOIN dbo.DATOS_ESTUD de
          ON TRY_CONVERT(int, de.codigo_estud) = TRY_CONVERT(int, cxe.codigo_estud)
        LEFT JOIN dbo.CorreosEstudIntec ce
          ON TRY_CONVERT(int, ce.codestud) = TRY_CONVERT(int, de.codigo_estud)
        LEFT JOIN dbo.CARRERAS c
          ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
        LEFT JOIN dbo.PERIODO pe
          ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cxe.codigo_periodo)
        LEFT JOIN dbo.PENSUM p
          ON TRY_CONVERT(int, p.codigo_materia) = TRY_CONVERT(int, cxe.codigo_materia)
         AND TRY_CONVERT(int, p.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
        WHERE TRY_CONVERT(int, cxe.codigo_estud) = ?
        ORDER BY
            COALESCE(TRY_CONVERT(int, pe.Orden), TRY_CONVERT(int, pe.cod_periodo)) DESC,
            TRY_CONVERT(int, pe.cod_periodo) DESC,
            TRY_CONVERT(nvarchar(4000), c.Nombre_Basica),
            TRY_CONVERT(int, p.Semestre),
            TRY_CONVERT(int, p.Orden),
            TRY_CONVERT(nvarchar(4000), p.Nomb_Materia)
        """,
        codigo_estud,
    )
    rows = cursor.fetchall()
    profile = _student_profile_from_row(rows[0]) if rows else _fetch_student_profile(cursor, codigo_estud)
    return profile, [_record_item(row) for row in rows]


def _record_summary(items: list[dict[str, Any]]) -> dict[str, Any]:
    total = len(items)
    aprobadas = sum(1 for item in items if item["aprobada"])
    reprobadas = sum(1 for item in items if item["estado_academico"] == "Reprobada")
    en_curso = sum(1 for item in items if item["estado_academico"] == "En curso")
    notas = [item["promedio_final"] for item in items if item["promedio_final"] is not None]
    creditos_aprobados = sum(float(item["creditos"] or 0) for item in items if item["aprobada"])
    return {
        "total_materias": total,
        "aprobadas": aprobadas,
        "reprobadas": reprobadas,
        "en_curso": en_curso,
        "creditos_aprobados": round(creditos_aprobados, 2),
        "promedio_general": round(sum(notas) / len(notas), 2) if notas else None,
        "cumplimiento_academico": round((aprobadas / total) * 100, 2) if total else 0,
    }


def _record_sort_value(item: dict[str, Any]) -> tuple[int, int, float, int]:
    final = _number(item.get("promedio_final"))
    try:
        period = int(str(item.get("codigo_periodo") or "0").strip())
    except ValueError:
        period = 0
    return (
        1 if item.get("aprobada") else 0,
        1 if final is not None else 0,
        final if final is not None else -1,
        period,
    )


def _curriculum_item(row: Any) -> dict[str, Any]:
    return {
        "cod_anio_basica": _clean(row.cod_anio_basica),
        "nombre_carrera": _clean(row.nombre_carrera),
        "codigo_materia": _clean(row.codigo_materia),
        "cod_materia": _clean(row.cod_materia),
        "nombre_materia": _clean(row.nombre_materia),
        "semestre": _int(row.semestre),
        "creditos": _number(row.creditos),
        "horas": _number(row.horas),
        "orden": _int(row.orden),
        "num_malla": _int(row.num_malla),
        "unidad_organiza": _clean(getattr(row, "unidad_organiza", "")),
        "estado_materia": _clean(getattr(row, "estado_materia", "")),
    }


def _curriculum_from_record_items(record_items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_subject: dict[str, dict[str, Any]] = {}
    for item in record_items:
        subject_code = _clean(item.get("codigo_materia"))
        if not subject_code or subject_code in by_subject:
            continue
        by_subject[subject_code] = {
            "cod_anio_basica": _clean(item.get("cod_anio_basica")),
            "nombre_carrera": _clean(item.get("nombre_carrera")),
            "codigo_materia": subject_code,
            "cod_materia": _clean(item.get("cod_materia")),
            "nombre_materia": _clean(item.get("nombre_materia")),
            "semestre": _int(item.get("semestre")),
            "creditos": _number(item.get("creditos")),
            "horas": _number(item.get("horas")),
            "orden": _int(item.get("orden")),
            "num_malla": _int(item.get("num_malla")),
            "unidad_organiza": "",
            "estado_materia": "Desde record academico",
        }
    return sorted(
        by_subject.values(),
        key=lambda item: (
            _int(item.get("semestre")) or 999,
            _clean(item.get("nombre_materia")),
        ),
    )


def _student_career_from_record(items: list[dict[str, Any]]) -> int | None:
    for item in items:
        value = _int(item.get("cod_anio_basica"))
        if value is not None:
            return value
    return None


def _fetch_student_current_career(
    cursor: pyodbc.Cursor,
    codigo_estud: int,
    items: list[dict[str, Any]],
) -> int | None:
    cursor.execute(
        """
        SELECT TOP (1) TRY_CONVERT(int, cm.cod_anio_Basica) AS cod_anio_basica
        FROM dbo.CABECERA_MATRICULA cm
        LEFT JOIN dbo.PERIODO pe
          ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cm.codigo_periodo)
        WHERE TRY_CONVERT(int, cm.codigo_estud) = ?
          AND TRY_CONVERT(int, cm.cod_anio_Basica) IS NOT NULL
        ORDER BY
            COALESCE(TRY_CONVERT(int, pe.Orden), TRY_CONVERT(int, cm.codigo_periodo)) DESC,
            TRY_CONVERT(int, cm.codigo_periodo) DESC,
            TRY_CONVERT(int, cm.cod_anio_Basica)
        """,
        codigo_estud,
    )
    row = cursor.fetchone()
    career_from_header = _int(row.cod_anio_basica) if row else None
    if career_from_header is not None:
        return career_from_header

    career_from_record = _student_career_from_record(items)
    if career_from_record is not None:
        return career_from_record
    cursor.execute(
        """
        SELECT TOP (1) TRY_CONVERT(int, cxe.cod_anio_Basica) AS cod_anio_basica
        FROM dbo.CARRERAXESTUD cxe
        WHERE TRY_CONVERT(int, cxe.codigo_estud) = ?
        ORDER BY
            TRY_CONVERT(int, cxe.codigo_periodo) DESC,
            TRY_CONVERT(int, cxe.cod_anio_Basica)
        """,
        codigo_estud,
    )
    row = cursor.fetchone()
    return _int(row.cod_anio_basica) if row else None


def _fetch_student_payments(cursor: pyodbc.Cursor, codigo_estud: int) -> list[dict[str, Any]]:
    cursor.execute(
        """
        SELECT TOP (12)
            cm.codigo_periodo,
            pe.Detalle_Periodo,
            cm.cod_anio_Basica,
            c.Nombre_Basica,
            cm.Num_Matricula,
            cm.numcodigo,
            cm.fecha_pago,
            cm.valor,
            cm.InscripValor,
            cm.MatriValor,
            cm.Cuota1,
            cm.Beca,
            cm.Descuento,
            COALESCE(cm.urlconvenio, pre.urlconvenio) AS urlconvenio,
            rp.Num AS pago_num,
            rp.Detalle AS pago_detalle,
            rp.fechapago AS pago_fecha,
            rp.ValorRegistrado AS pago_valor,
            rp.NoDeposito AS pago_referencia,
            rp.Banco AS pago_banco
        FROM dbo.CABECERA_MATRICULA cm
        LEFT JOIN dbo.PERIODO pe ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cm.codigo_periodo)
        LEFT JOIN dbo.CARRERAS c ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, cm.cod_anio_Basica)
        OUTER APPLY (
            SELECT TOP (1) pay.*
            FROM dbo.REGISTROPAGOS pay
            WHERE TRY_CONVERT(int, pay.Codestu) = TRY_CONVERT(int, cm.codigo_estud)
              AND TRY_CONVERT(int, pay.codperiodo) = TRY_CONVERT(int, cm.codigo_periodo)
              AND TRY_CONVERT(int, pay.cod_anio_Basica) = TRY_CONVERT(int, cm.cod_anio_Basica)
            ORDER BY TRY_CONVERT(int, pay.Num) DESC, pay.fechapago DESC
        ) rp
        OUTER APPLY (
            SELECT TOP (1) p.urlconvenio
            FROM dbo.PREINSCRIPCION p
            WHERE TRY_CONVERT(int, p.Codestu) = TRY_CONVERT(int, cm.codigo_estud)
               OR LTRIM(RTRIM(TRY_CONVERT(nvarchar(20), p.Cedula))) IN (
                    SELECT TOP (1) LTRIM(RTRIM(TRY_CONVERT(nvarchar(20), d.Cedula_Est)))
                    FROM dbo.DATOS_ESTUD d
                    WHERE TRY_CONVERT(int, d.codigo_estud) = TRY_CONVERT(int, cm.codigo_estud)
                )
            ORDER BY TRY_CONVERT(int, p.num) DESC
        ) pre
        WHERE TRY_CONVERT(int, cm.codigo_estud) = ?
        ORDER BY
            TRY_CONVERT(int, cm.codigo_periodo) DESC,
            TRY_CONVERT(int, cm.Num_Matricula) DESC,
            cm.fecha_pago DESC
        """,
        codigo_estud,
    )
    payments: list[dict[str, Any]] = []
    for row in cursor.fetchall():
        total = _number(getattr(row, "valor", None)) or 0
        beca = _number(getattr(row, "Beca", None)) or 0
        descuento = _number(getattr(row, "Descuento", None)) or 0
        saldo = max(round(total - beca - descuento, 2), 0)
        cuota = _number(getattr(row, "Cuota1", None)) or 0
        cuotas = int(round(saldo / cuota)) if saldo > 0 and cuota > 0 else 1
        payments.append(
            {
                "codigo_periodo": _clean(getattr(row, "codigo_periodo", "")),
                "periodo": _clean(getattr(row, "Detalle_Periodo", "")),
                "cod_anio_basica": _clean(getattr(row, "cod_anio_Basica", "")),
                "carrera": _clean(getattr(row, "Nombre_Basica", "")),
                "num_matricula": _clean(getattr(row, "Num_Matricula", "")),
                "codigo_documentacion": _clean(getattr(row, "numcodigo", "")),
                "fecha_pago": _date_text(getattr(row, "fecha_pago", "")),
                "total": total,
                "inscripcion": _number(getattr(row, "InscripValor", None)) or 0,
                "matricula": _number(getattr(row, "MatriValor", None)) or 0,
                "beca": beca,
                "descuento": descuento,
                "saldo": saldo,
                "cuota": cuota,
                "cuotas": cuotas,
                "convenio_url": _clean(getattr(row, "urlconvenio", "")),
                "pago_num": _int(getattr(row, "pago_num", None)),
                "pago_detalle": _clean(getattr(row, "pago_detalle", "")),
                "pago_fecha": _date_text(getattr(row, "pago_fecha", "")),
                "pago_valor": _number(getattr(row, "pago_valor", None)) or 0,
                "pago_referencia": _clean(getattr(row, "pago_referencia", "")),
                "pago_banco": _clean(getattr(row, "pago_banco", "")),
            }
        )
    return payments


def _fetch_student_current_pensum(
    cursor: pyodbc.Cursor,
    codigo_estud: int,
    career_code: int,
    items: list[dict[str, Any]],
) -> int | None:
    cursor.execute(
        """
        SELECT TOP (1) TRY_CONVERT(int, mp.Malla) AS num_malla
        FROM dbo.MALLA_PENSUM mp
        WHERE TRY_CONVERT(int, mp.Cod_Carrera) = ?
          AND UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(20), mp.Estado), N'')))) = N'A'
          AND TRY_CONVERT(int, mp.Malla) IS NOT NULL
        ORDER BY
            TRY_CONVERT(int, mp.Malla) DESC,
            TRY_CONVERT(int, mp.Num) DESC
        """,
        career_code,
    )
    row = cursor.fetchone()
    active_malla = _int(row.num_malla) if row else None
    if active_malla is not None:
        return active_malla

    counts: dict[int, int] = {}
    for item in items:
        item_career = _int(item.get("cod_anio_basica"))
        if item_career is not None and item_career != career_code:
            continue
        num_malla = _int(item.get("num_malla"))
        if num_malla is None:
            continue
        counts[num_malla] = counts.get(num_malla, 0) + 1
    if counts:
        return max(counts, key=lambda key: (counts[key], key))

    cursor.execute(
        """
        SELECT TOP (1) TRY_CONVERT(int, p.NumMalla) AS num_malla
        FROM dbo.CARRERAXESTUD cxe
        INNER JOIN dbo.PENSUM p
          ON TRY_CONVERT(int, p.codigo_materia) = TRY_CONVERT(int, cxe.codigo_materia)
         AND TRY_CONVERT(int, p.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
        WHERE TRY_CONVERT(int, cxe.codigo_estud) = ?
          AND TRY_CONVERT(int, cxe.cod_anio_Basica) = ?
          AND TRY_CONVERT(int, p.NumMalla) IS NOT NULL
        GROUP BY TRY_CONVERT(int, p.NumMalla)
        ORDER BY
            COUNT(*) DESC,
            MAX(TRY_CONVERT(int, cxe.codigo_periodo)) DESC,
            TRY_CONVERT(int, p.NumMalla) DESC
        """,
        codigo_estud,
        career_code,
    )
    row = cursor.fetchone()
    current_malla = _int(row.num_malla) if row else None
    if current_malla is not None:
        return current_malla

    cursor.execute(
        """
        SELECT TOP (1) TRY_CONVERT(int, p.NumMalla) AS num_malla
        FROM dbo.PENSUM p
        WHERE TRY_CONVERT(int, p.Cod_AnioBasica) = ?
          AND TRY_CONVERT(int, p.NumMalla) IS NOT NULL
        GROUP BY TRY_CONVERT(int, p.NumMalla)
        ORDER BY
            TRY_CONVERT(int, p.NumMalla) DESC,
            COUNT(*) DESC
        """,
        career_code,
    )
    row = cursor.fetchone()
    return _int(row.num_malla) if row else None


def _academic_grid_items(
    curriculum: list[dict[str, Any]],
    record_items: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    records_by_subject: dict[str, list[dict[str, Any]]] = {}
    for item in record_items:
        subject_code = _clean(item.get("codigo_materia"))
        if subject_code:
            records_by_subject.setdefault(subject_code, []).append(item)

    grid: list[dict[str, Any]] = []
    for subject in curriculum:
        subject_code = _clean(subject.get("codigo_materia"))
        attempts = records_by_subject.get(subject_code, [])
        best_attempt = max(attempts, key=_record_sort_value) if attempts else None
        best_final = _number(best_attempt.get("promedio_final")) if best_attempt else None
        has_attempt = bool(best_attempt)

        if best_final is not None:
            academic_status = "Aprobada" if best_final >= 7 else "Reprobada"
        elif has_attempt:
            academic_status = "En curso"
        else:
            academic_status = "Pendiente"

        if best_final is not None:
            approved = best_final >= 7
        elif best_attempt and best_attempt.get("aprobada"):
            approved = True
        else:
            approved = False

        if approved:
            academic_status = "Aprobada"

        grid.append(
            {
                **subject,
                "estado_academico": academic_status,
                "aprobada": approved,
                "faltante": not approved,
                "intentos": len(attempts),
                "ultimo_periodo": best_attempt.get("detalle_periodo") if best_attempt else "",
                "codigo_periodo": best_attempt.get("codigo_periodo") if best_attempt else "",
                "paralelo": best_attempt.get("paralelo") if best_attempt else "",
                "tipo_matricula": best_attempt.get("tipo_matricula") if best_attempt else "",
                "esquema_calificacion": best_attempt.get("esquema_calificacion") if best_attempt else "",
                "teoria_homo": best_attempt.get("teoria_homo") if best_attempt else None,
                "practica_homo": best_attempt.get("practica_homo") if best_attempt else None,
                "p1_tareas": best_attempt.get("p1_tareas") if best_attempt else None,
                "p1_proyectos": best_attempt.get("p1_proyectos") if best_attempt else None,
                "p1_examen": best_attempt.get("p1_examen") if best_attempt else None,
                "prom_p1": best_attempt.get("prom_p1") if best_attempt else None,
                "p2_tareas": best_attempt.get("p2_tareas") if best_attempt else None,
                "p2_proyectos": best_attempt.get("p2_proyectos") if best_attempt else None,
                "p2_examen": best_attempt.get("p2_examen") if best_attempt else None,
                "prom_p2": best_attempt.get("prom_p2") if best_attempt else None,
                "p3_tareas": best_attempt.get("p3_tareas") if best_attempt else None,
                "p3_proyectos": best_attempt.get("p3_proyectos") if best_attempt else None,
                "p3_examen": best_attempt.get("p3_examen") if best_attempt else None,
                "prom_p3": best_attempt.get("prom_p3") if best_attempt else None,
                "promedio_final": best_final if best_final is not None else None,
                "nota_aprobar": best_attempt.get("nota_aprobar") if best_attempt else 7.0,
            }
        )
    return grid


def _curriculum_summary(grid: list[dict[str, Any]]) -> dict[str, Any]:
    total = len(grid)
    aprobadas = 0
    en_curso = 0
    reprobadas = 0
    creditos_totales = sum(float(item.get("creditos") or 0) for item in grid)
    creditos_aprobados = 0.0
    for item in grid:
        final = _number(item.get("promedio_final"))
        if final is None:
            if item.get("estado_academico") == "En curso":
                en_curso += 1
            continue
        if final >= 7:
            aprobadas += 1
            creditos_aprobados += float(item.get("creditos") or 0)
        else:
            reprobadas += 1
    return {
        "total_materias": total,
        "aprobadas": aprobadas,
        "faltantes": max(total - aprobadas, 0),
        "en_curso": en_curso,
        "reprobadas": reprobadas,
        "creditos_totales": round(creditos_totales, 2),
        "creditos_aprobados": round(creditos_aprobados, 2),
        "porcentaje_avance": round((aprobadas / total) * 100, 2) if total else 0,
    }


def _fetch_student_curriculum(
    cursor: pyodbc.Cursor,
    codigo_estud: int,
    record_items: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    career_code = _fetch_student_current_career(cursor, codigo_estud, record_items)
    if career_code is None:
        curriculum = _curriculum_from_record_items(record_items)
        academic_grid = _academic_grid_items(curriculum, record_items)
        return curriculum, academic_grid, _curriculum_summary(academic_grid)

    num_malla = _fetch_student_current_pensum(cursor, codigo_estud, career_code, record_items)
    curriculum_sql = """
        SELECT
            TRY_CONVERT(varchar(50), p.Cod_AnioBasica) AS cod_anio_basica,
            TRY_CONVERT(nvarchar(4000), c.Nombre_Basica) AS nombre_carrera,
            TRY_CONVERT(varchar(50), p.codigo_materia) AS codigo_materia,
            TRY_CONVERT(varchar(100), p.cod_materia) AS cod_materia,
            TRY_CONVERT(nvarchar(4000), p.Nomb_Materia) AS nombre_materia,
            TRY_CONVERT(int, p.Semestre) AS semestre,
            TRY_CONVERT(float, p.Creditos) AS creditos,
            TRY_CONVERT(float, p.Horas) AS horas,
            TRY_CONVERT(int, p.Orden) AS orden,
            COALESCE(TRY_CONVERT(int, mp.Malla), TRY_CONVERT(int, p.NumMalla)) AS num_malla,
            TRY_CONVERT(nvarchar(255), p.Unidad_Organiza) AS unidad_organiza,
            TRY_CONVERT(nvarchar(100), p.estado_mat) AS estado_materia
        FROM dbo.PENSUM p
        LEFT JOIN dbo.CARRERAS c
          ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, p.Cod_AnioBasica)
        LEFT JOIN dbo.MALLA_PENSUM mp
          ON TRY_CONVERT(int, mp.Cod_Carrera) = TRY_CONVERT(int, p.Cod_AnioBasica)
         AND TRY_CONVERT(int, mp.Malla) = TRY_CONVERT(int, p.NumMalla)
        WHERE {where_clause}
        ORDER BY
            TRY_CONVERT(int, p.Semestre),
            TRY_CONVERT(int, p.Orden),
            TRY_CONVERT(int, p.codigo_materia)
        """

    def fetch_curriculum(where_clause: str, *params: Any) -> list[dict[str, Any]]:
        cursor.execute(curriculum_sql.format(where_clause=where_clause), *params)
        return [_curriculum_item(row) for row in cursor.fetchall()]

    if num_malla is not None:
        curriculum = fetch_curriculum(
            "TRY_CONVERT(int, p.Cod_AnioBasica) = ? "
            "AND COALESCE(TRY_CONVERT(int, mp.Malla), TRY_CONVERT(int, p.NumMalla)) = ?",
            career_code,
            num_malla,
        )
    else:
        curriculum = []

    if not curriculum:
        curriculum = fetch_curriculum("TRY_CONVERT(int, p.Cod_AnioBasica) = ?", career_code)

    if not curriculum:
        career_name = next((_clean(item.get("nombre_carrera")) for item in record_items if _clean(item.get("nombre_carrera"))), "")
        if career_name:
            if num_malla is not None:
                curriculum = fetch_curriculum(
                    "UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(4000), c.Nombre_Basica)))) = UPPER(LTRIM(RTRIM(?))) "
                    "AND COALESCE(TRY_CONVERT(int, mp.Malla), TRY_CONVERT(int, p.NumMalla)) = ?",
                    career_name,
                    num_malla,
                )
            if not curriculum:
                curriculum = fetch_curriculum(
                    "UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(4000), c.Nombre_Basica)))) = UPPER(LTRIM(RTRIM(?)))",
                    career_name,
                )
    if not curriculum:
        curriculum = _curriculum_from_record_items(record_items)
    academic_grid = _academic_grid_items(curriculum, record_items)
    return curriculum, academic_grid, _curriculum_summary(academic_grid)


@router.get("/student/me")
def student_profile(
    current_user: Annotated[SessionUser, Depends(_STUDENT_ACCESS)],
) -> dict[str, Any]:
    if current_user.codigo_estud is None:
        raise HTTPException(status_code=403, detail="La sesion no tiene estudiante vinculado")
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            profile, items = _fetch_student_record(cursor, current_user.codigo_estud)
        return {"student": profile, "summary": _record_summary(items)}
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error consultando perfil del estudiante: {exc}") from exc


@router.get("/student/record")
def student_record(
    current_user: Annotated[SessionUser, Depends(_STUDENT_ACCESS)],
    approved_only: Annotated[bool, Query(description="Mostrar solo materias aprobadas")] = False,
) -> dict[str, Any]:
    if current_user.codigo_estud is None:
        raise HTTPException(status_code=403, detail="La sesion no tiene estudiante vinculado")
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            profile, items = _fetch_student_record(cursor, current_user.codigo_estud)
            curriculum, academic_grid, curriculum_resume = _fetch_student_curriculum(
                cursor,
                current_user.codigo_estud,
                items,
            )
            payments = _fetch_student_payments(cursor, current_user.codigo_estud)
        summary = _record_summary(items)
        visible_items = [item for item in items if item["aprobada"]] if approved_only else items
        return {
            "student": profile,
            "summary": summary,
            "curriculum_summary": curriculum_resume,
            "curriculum": curriculum,
            "academic_grid": academic_grid,
            "payments": payments,
            "items": visible_items,
            "total": len(visible_items),
        }
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error consultando record academico: {exc}") from exc


@router.get("/student/record/export")
def student_record_export(
    current_user: Annotated[SessionUser, Depends(_STUDENT_ACCESS)],
    approved_only: Annotated[bool, Query(description="Parametro heredado; la exportacion incluye aprobadas y reprobadas")] = False,
    codigo_periodo: Annotated[str | None, Query(description="Periodo seleccionado para exportar calificaciones")] = None,
) -> StreamingResponse:
    _ = approved_only
    if current_user.codigo_estud is None:
        raise HTTPException(status_code=403, detail="La sesion no tiene estudiante vinculado")
    selected_period = _clean(codigo_periodo)
    if not selected_period:
        raise HTTPException(status_code=400, detail="Seleccione un periodo para exportar calificaciones")
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            profile, items = _fetch_student_record(cursor, current_user.codigo_estud)
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error exportando record academico: {exc}") from exc

    period_items = []
    for item in items:
        if item["codigo_periodo"] == selected_period or item["detalle_periodo"] == selected_period:
            period_items.append(item)

    visible_items = list(period_items)
    visible_items.sort(
        key=lambda item: (
            _int(item.get("semestre")) or 999,
            _int(item.get("orden")) or 9999,
            _int(item.get("codigo_materia")) or 999999,
            _clean(item.get("nombre_materia")),
        )
    )
    selected_period_label = next(
        (item["detalle_periodo"] for item in period_items if item["detalle_periodo"]),
        selected_period,
    )

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Calificaciones"
    sheet.append(["Estudiante", profile["nombre_estudiante"]])
    sheet.append(["Cedula", profile["cedula"]])
    sheet.append(["Codigo", profile["codigo_estud"]])
    sheet.append(["Correo INTEC", profile["correo_intec"]])
    sheet.append(["Periodo", selected_period_label])
    sheet.append(["Carrera", next((item["nombre_carrera"] for item in period_items if item["nombre_carrera"]), "")])
    sheet.append([])
    homologation_only = bool(visible_items) and all(
        _is_homologation_type(
            item.get("tipo_matricula"),
            item.get("detalle_periodo"),
            item.get("esquema_calificacion"),
        )
        for item in visible_items
    )
    if homologation_only:
        sheet.append(["#", "Periodo", "Nivel", "Materia", "Codigo materia", "Esquema", "Nota final", "Estado"])
    else:
        sheet.append([
            "#",
            "Periodo",
            "Nivel",
            "Materia",
            "Codigo materia",
            "Esquema",
            "Prom. 1",
            "Prom. 2",
            "Prom. 3",
            "Nota final",
            "Estado",
        ])
    for index, item in enumerate(visible_items, start=1):
        final = _number(item.get("promedio_final"))
        if final is None:
            estado = "En curso"
        elif final >= 7:
            estado = "Aprobada"
        else:
            estado = "Reprobada"
        is_homo = _is_homologation_type(item.get("tipo_matricula"), item.get("detalle_periodo"))
        row = [
            index,
            item["detalle_periodo"],
            item["semestre"],
            item["nombre_materia"],
            item["cod_materia"] or item["codigo_materia"],
            item["esquema_calificacion"],
        ]
        if not homologation_only:
            row.extend([None if is_homo else item["prom_p1"], None if is_homo else item["prom_p2"], None if is_homo else item["prom_p3"]])
        row.extend([final, estado])
        sheet.append(row)

    for worksheet in workbook.worksheets:
        for column in worksheet.columns:
            max_length = max(len(str(cell.value or "")) for cell in column)
            worksheet.column_dimensions[column[0].column_letter].width = min(max(max_length + 2, 12), 42)

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    period_suffix = selected_period.replace("/", "-").replace("\\", "-").replace(" ", "-")
    filename = f"calificaciones-{profile['codigo_estud']}-{period_suffix}.xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/student/record/export-pdf")
def student_record_pdf_export(
    current_user: Annotated[SessionUser, Depends(_STUDENT_ACCESS)],
    tipo: Annotated[str, Query(description="academica o calificaciones")] = "calificaciones",
    codigo_periodo: Annotated[str | None, Query(description="Periodo seleccionado para calificaciones")] = None,
) -> StreamingResponse:
    if current_user.codigo_estud is None:
        raise HTTPException(status_code=403, detail="La sesion no tiene estudiante vinculado")

    report_type = _clean(tipo).lower()
    if report_type not in {"academica", "calificaciones"}:
        raise HTTPException(status_code=400, detail="Tipo de reporte no valido")

    selected_period = _clean(codigo_periodo)
    if report_type == "calificaciones" and not selected_period:
        raise HTTPException(status_code=400, detail="Seleccione un periodo para exportar calificaciones")

    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            profile, items = _fetch_student_record(cursor, current_user.codigo_estud)
            curriculum, academic_grid, curriculum_resume = _fetch_student_curriculum(
                cursor,
                current_user.codigo_estud,
                items,
            )
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error generando PDF academico: {exc}") from exc

    career = (
        next((_clean(item.get("nombre_carrera")) for item in academic_grid if _clean(item.get("nombre_carrera"))), "")
        or next((_clean(item.get("nombre_carrera")) for item in curriculum if _clean(item.get("nombre_carrera"))), "")
        or next((_clean(item.get("nombre_carrera")) for item in items if _clean(item.get("nombre_carrera"))), "")
    )

    if report_type == "academica":
        academic_grid = _academic_grid_with_practice_requirements(academic_grid, career)
        academic_grid.sort(
            key=lambda item: (
                _int(item.get("semestre")) or 999,
                _int(item.get("orden")) or 9999,
                _int(item.get("codigo_materia")) or 999999,
                _clean(item.get("nombre_materia")),
            )
        )
        pdf_bytes = _build_student_report_pdf(
            "Malla academica",
            f"Malla y calificaciones consolidadas | Avance {curriculum_resume.get('porcentaje_avance', 0)}%",
            profile,
            career,
            ["Nivel", "Codigo materia", "Materia", "Creditos", "Esquema", "HOMO", "Prom. 1", "Prom. 2", "Prom. 3", "Final", "Estado"],
            _academic_pdf_rows(academic_grid),
            [0.8 * cm, 2.2 * cm, 5.4 * cm, 1.2 * cm, 1.6 * cm, 1.7 * cm, 1.1 * cm, 1.1 * cm, 1.1 * cm, 1.1 * cm, 1.5 * cm],
        )
        filename = f"malla-academica-{_safe_filename(profile.get('codigo_estud'))}.pdf"
    else:
        period_items = [
            item
            for item in items
            if item["codigo_periodo"] == selected_period or item["detalle_periodo"] == selected_period
        ]
        period_items.sort(
            key=lambda item: (
                _int(item.get("semestre")) or 999,
                _int(item.get("orden")) or 9999,
                _int(item.get("codigo_materia")) or 999999,
                _clean(item.get("nombre_materia")),
            )
        )
        selected_period_label = next(
            (item["detalle_periodo"] for item in period_items if item["detalle_periodo"]),
            selected_period,
        )
        homologation_only = bool(period_items) and all(
            _is_homologation_type(
                item.get("tipo_matricula"),
                item.get("detalle_periodo"),
                item.get("esquema_calificacion"),
            )
            for item in period_items
        )
        headers = (
            ["Nivel", "Periodo", "Codigo", "Materia", "Esquema", "Final", "Estado"]
            if homologation_only
            else ["Nivel", "Periodo", "Codigo", "Materia", "Esquema", "Prom. 1", "Prom. 2", "Prom. 3", "Final", "Estado"]
        )
        col_widths = (
            [0.8 * cm, 3.2 * cm, 2.0 * cm, 6.4 * cm, 1.9 * cm, 1.3 * cm, 1.7 * cm]
            if homologation_only
            else [0.8 * cm, 3.0 * cm, 2.0 * cm, 5.4 * cm, 1.6 * cm, 1.1 * cm, 1.1 * cm, 1.1 * cm, 1.2 * cm, 1.4 * cm]
        )
        pdf_bytes = _build_student_report_pdf(
            "Calificaciones",
            f"Periodo: {selected_period_label}",
            profile,
            career,
            headers,
            _calificaciones_pdf_rows(period_items, homologation_only),
            col_widths,
        )
        filename = f"calificaciones-{_safe_filename(profile.get('codigo_estud'))}-{_safe_filename(selected_period_label)}.pdf"

    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/student/record/export-secretaria-pdf")
def student_record_secretary_pdf_export(
    current_user: Annotated[SessionUser, Depends(_STUDENT_ACCESS)],
    codigo_periodo: Annotated[str | None, Query(description="Periodo seleccionado para reporte de notas formato Secretaria")] = None,
    tipo: Annotated[str, Query(description="calificaciones o malla")] = "calificaciones",
) -> StreamingResponse:
    if current_user.codigo_estud is None:
        raise HTTPException(status_code=403, detail="La sesion no tiene estudiante vinculado")

    selected_period = _clean(codigo_periodo)
    report_type = _clean(tipo).lower()
    if report_type not in {"calificaciones", "malla"}:
        raise HTTPException(status_code=400, detail="Tipo de reporte no valido")
    if report_type == "calificaciones" and not selected_period:
        raise HTTPException(status_code=400, detail="Seleccione un periodo para exportar el reporte de notas")

    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            profile, items = _fetch_student_record(cursor, current_user.codigo_estud)
            curriculum, academic_grid, _curriculum_resume = _fetch_student_curriculum(
                cursor,
                current_user.codigo_estud,
                items,
            )
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error generando reporte de notas formato Secretaria: {exc}") from exc

    if report_type == "malla":
        career = (
            next((_clean(item.get("nombre_carrera")) for item in academic_grid if _clean(item.get("nombre_carrera"))), "")
            or next((_clean(item.get("nombre_carrera")) for item in curriculum if _clean(item.get("nombre_carrera"))), "")
            or next((_clean(item.get("nombre_carrera")) for item in items if _clean(item.get("nombre_carrera"))), "")
        )
        report_items = list(academic_grid)
        report_items = _academic_grid_with_practice_requirements(report_items, career)
        selected_period_label = "Malla academica general"
        for item in report_items:
            item["detalle_periodo"] = item.get("ultimo_periodo") or selected_period_label
            item["codigo_periodo"] = item.get("codigo_periodo") or "MALLA"
            item["nombre_carrera"] = item.get("nombre_carrera") or career
    else:
        report_items = [
            item
            for item in items
            if item["codigo_periodo"] == selected_period or item["detalle_periodo"] == selected_period
        ]
        selected_period_label = next(
            (item["detalle_periodo"] for item in report_items if item["detalle_periodo"]),
            selected_period,
        )

    report_items.sort(
        key=lambda item: (
            _int(item.get("semestre")) or 999,
            _int(item.get("orden")) or 9999,
            _int(item.get("codigo_materia")) or 999999,
            _clean(item.get("nombre_materia")),
        )
    )
    for item in report_items:
        item["codigo_estud"] = item.get("codigo_estud") or profile.get("codigo_estud")
        item["cedula"] = item.get("cedula") or profile.get("cedula")
        item["nombre_estudiante"] = item.get("nombre_estudiante") or profile.get("nombre_estudiante")

    pdf_bytes = _student_secretaria_notes_pdf(profile, report_items, report_type, selected_period_label)
    filename_prefix = "malla-secretaria" if report_type == "malla" else "reporte-notas-secretaria"
    filename = f"{filename_prefix}-{_safe_filename(profile.get('codigo_estud'))}-{_safe_filename(selected_period_label)}.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _teacher_code(current_user: SessionUser) -> int:
    if current_user.codigo_doc is None:
        raise HTTPException(status_code=403, detail="La sesion no tiene docente vinculado")
    return current_user.codigo_doc


def _course_item(row: Any) -> dict[str, Any]:
    codigo_periodo = _clean(row.codigo_periodo)
    detalle_periodo = _clean(row.detalle_periodo)
    tipo_periodo = _clean(getattr(row, "tipo_periodo", ""))
    es_homologacion = _is_homologation_type(tipo_periodo, detalle_periodo)
    internal_code = _clean(row.codigo_materia)
    common_code = _clean(getattr(row, "cod_materia", "")) or internal_code
    return {
        "codigo_doc": _clean(row.codigo_doc),
        "cod_anio_basica": _clean(row.cod_anio_basica),
        "cod_anio_basicas": [_clean(row.cod_anio_basica)] if _clean(row.cod_anio_basica) else [],
        "nombre_carrera": _clean(row.nombre_carrera),
        "codigo_materia": common_code,
        "codigo_materias": [internal_code] if internal_code else [],
        "cod_materia": common_code,
        "nombre_materia": _clean(row.nombre_materia),
        "codigo_periodo": codigo_periodo,
        "codigo_periodos": [codigo_periodo] if codigo_periodo else [],
        "detalle_periodo": detalle_periodo,
        "detalle_periodos": detalle_periodo,
        "tipo_periodo": tipo_periodo,
        "es_homologacion": es_homologacion,
        "fecha_inicio": _contract_value(getattr(row, "fecha_inicio", None)),
        "fecha_fin": _contract_value(getattr(row, "fecha_fin", None)),
        "paralelo": _clean(row.paralelo),
        "cod_jornada": _int(row.cod_jornada),
        "jornada": _clean(getattr(row, "jornada", "")) or (
            f"Jornada {_clean(row.cod_jornada)}" if _clean(row.cod_jornada) else ""
        ),
        "semestre": _int(getattr(row, "semestre", None)),
        "unidad_curricular": _clean(getattr(row, "unidad_curricular", "")),
        "periodo_orden": _int(getattr(row, "periodo_orden", None)) or _int(codigo_periodo) or 0,
        "period_count": 1,
        "total_estudiantes": _int(row.total_estudiantes) or 0,
        "estado_moodle_doc": bool(_int(row.estado_moodle_doc)),
    }


def _group_teacher_courses(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    period_groups: dict[tuple[str, str, str, str, str, bool], dict[str, Any]] = {}
    for item in items:
        common_code = _clean(item.get("cod_materia") or item.get("codigo_materia"))
        key = (
            _clean(item.get("cod_anio_basica")),
            common_code,
            _clean(item.get("codigo_periodo")),
            _clean(item.get("paralelo")).upper(),
            _clean(item.get("cod_jornada")),
            bool(item.get("es_homologacion")),
        )
        bucket = period_groups.get(key)
        if not bucket:
            bucket = item.copy()
            bucket["codigo_materia"] = common_code
            bucket["cod_materia"] = common_code
            bucket["cod_anio_basicas"] = []
            bucket["codigo_materias"] = []
            bucket["_nombre_carreras"] = []
            bucket["total_estudiantes"] = 0
            period_groups[key] = bucket
        career_code = _clean(item.get("cod_anio_basica"))
        if career_code and career_code not in bucket["cod_anio_basicas"]:
            bucket["cod_anio_basicas"].append(career_code)
        internal_code_values = item.get("codigo_materias") if isinstance(item.get("codigo_materias"), list) else [item.get("codigo_materia")]
        for internal_code in internal_code_values:
            internal_code = _clean(internal_code)
            if internal_code and internal_code not in bucket["codigo_materias"]:
                bucket["codigo_materias"].append(internal_code)
        career_name = _clean(item.get("nombre_carrera"))
        if career_name and career_name not in bucket["_nombre_carreras"]:
            bucket["_nombre_carreras"].append(career_name)
        bucket["total_estudiantes"] = max(
            _int(bucket.get("total_estudiantes")) or 0,
            _int(item.get("total_estudiantes")) or 0,
        )

    normalized_items: list[dict[str, Any]] = []
    for bucket in period_groups.values():
        career_names = bucket.pop("_nombre_carreras", [])
        bucket["cod_anio_basica"] = ", ".join(bucket["cod_anio_basicas"])
        bucket["nombre_carrera"] = " / ".join(career_names) if len(career_names) <= 2 else f"{len(career_names)} carreras"
        normalized_items.append(bucket)

    assignments: list[dict[str, Any]] = []
    regular_groups: dict[tuple[str, str, str, str, str], list[dict[str, Any]]] = {}
    for item in normalized_items:
        if item.get("es_homologacion"):
            assignments.append(item)
            continue
        key = (
            _clean(item.get("cod_anio_basica")),
            _clean(item.get("cod_materia") or item.get("codigo_materia")),
            _clean(item.get("nombre_materia")),
            _clean(item.get("paralelo")).upper(),
            _clean(item.get("cod_jornada")),
        )
        regular_groups.setdefault(key, []).append(item)

    for courses in regular_groups.values():
        sorted_courses = sorted(
            courses,
            key=lambda item: (_int(item.get("periodo_orden")) or 0, _int(item.get("codigo_periodo")) or 0),
            reverse=True,
        )
        for index in range(0, len(sorted_courses), 2):
            chunk = sorted_courses[index:index + 2]
            base = chunk[0].copy()
            period_codes = [_clean(item.get("codigo_periodo")) for item in chunk if _clean(item.get("codigo_periodo"))]
            period_names = [_clean(item.get("detalle_periodo")) or _clean(item.get("codigo_periodo")) for item in chunk]
            base["codigo_periodos"] = period_codes
            base["codigo_periodo"] = period_codes[0] if period_codes else ""
            base["detalle_periodos"] = " / ".join(period_names)
            base["detalle_periodo"] = base["detalle_periodos"]
            base["period_count"] = len(chunk)
            base["total_estudiantes"] = sum(_int(item.get("total_estudiantes")) or 0 for item in chunk)
            assignments.append(base)

    # El codigo comun de PENSUM identifica una sola asignatura aunque exista en
    # varias carreras, periodos o paralelos. Se conserva cada alcance exacto en
    # ``asignaciones`` para consultar y actualizar la matricula correcta.
    subject_groups: dict[str, dict[str, Any]] = {}
    for assignment in assignments:
        common_code = _clean(assignment.get("cod_materia") or assignment.get("codigo_materia"))
        subject_key = common_code.upper()
        bucket = subject_groups.get(subject_key)
        if not bucket:
            bucket = assignment.copy()
            bucket["codigo_materia"] = common_code
            bucket["cod_materia"] = common_code
            bucket["asignaciones"] = []
            bucket["alcances_periodo"] = []
            bucket["cod_anio_basicas"] = []
            bucket["nombre_carreras"] = []
            bucket["codigo_materias"] = []
            bucket["codigo_periodos"] = []
            bucket["detalle_periodos_lista"] = []
            bucket["total_estudiantes"] = 0
            bucket["regular_count"] = 0
            bucket["homologation_count"] = 0
            subject_groups[subject_key] = bucket

        scope = assignment.copy()
        bucket["asignaciones"].append(scope)

        for career_code in assignment.get("cod_anio_basicas") or [assignment.get("cod_anio_basica")]:
            career_code = _clean(career_code)
            if career_code and career_code not in bucket["cod_anio_basicas"]:
                bucket["cod_anio_basicas"].append(career_code)

        career_name = _clean(assignment.get("nombre_carrera"))
        if career_name and career_name not in bucket["nombre_carreras"]:
            bucket["nombre_carreras"].append(career_name)

        for internal_code in assignment.get("codigo_materias") or []:
            internal_code = _clean(internal_code)
            if internal_code and internal_code not in bucket["codigo_materias"]:
                bucket["codigo_materias"].append(internal_code)

        period_codes = assignment.get("codigo_periodos") or [assignment.get("codigo_periodo")]
        for period_code in period_codes:
            period_code = _clean(period_code)
            if period_code and period_code not in bucket["codigo_periodos"]:
                bucket["codigo_periodos"].append(period_code)

        period_labels = [
            label.strip()
            for label in _clean(assignment.get("detalle_periodos") or assignment.get("detalle_periodo")).split(" / ")
            if label.strip()
        ]
        for period_label in period_labels:
            if period_label not in bucket["detalle_periodos_lista"]:
                bucket["detalle_periodos_lista"].append(period_label)

        bucket["total_estudiantes"] += _int(assignment.get("total_estudiantes")) or 0
        if assignment.get("es_homologacion"):
            bucket["homologation_count"] += 1
        else:
            bucket["regular_count"] += 1

    # Preserve the exact period/career scopes so clients can present the same
    # common subject code in separate period blocks without duplicating careers.
    for exact_scope in normalized_items:
        subject_key = _clean(exact_scope.get("cod_materia") or exact_scope.get("codigo_materia")).upper()
        bucket = subject_groups.get(subject_key)
        if bucket is not None:
            bucket["alcances_periodo"].append(exact_scope.copy())

    grouped: list[dict[str, Any]] = []
    for bucket in subject_groups.values():
        career_names = bucket.pop("nombre_carreras", [])
        period_labels = bucket.pop("detalle_periodos_lista", [])
        assignments_for_subject = bucket.get("asignaciones") or []
        regular_count = _int(bucket.get("regular_count")) or 0
        homologation_count = _int(bucket.get("homologation_count")) or 0

        bucket["cod_anio_basica"] = bucket["cod_anio_basicas"][0] if len(bucket["cod_anio_basicas"]) == 1 else ""
        bucket["nombre_carrera"] = (
            " / ".join(career_names)
            if len(career_names) <= 2
            else f"{len(career_names)} carreras"
        )
        bucket["codigo_periodo"] = bucket["codigo_periodos"][0] if len(bucket["codigo_periodos"]) == 1 else ""
        bucket["detalle_periodos"] = (
            " / ".join(period_labels)
            if len(period_labels) <= 2
            else f"{len(period_labels)} periodos"
        )
        bucket["detalle_periodo"] = bucket["detalle_periodos"]
        bucket["period_count"] = len(bucket["codigo_periodos"])
        bucket["assignment_count"] = len(assignments_for_subject)
        bucket["tiene_regular"] = regular_count > 0
        bucket["tiene_homologacion"] = homologation_count > 0
        bucket["es_homologacion"] = homologation_count > 0 and regular_count == 0
        bucket["tipo_periodo"] = "MIXTO" if regular_count and homologation_count else ("H" if homologation_count else "R")

        parallels = {_clean(scope.get("paralelo")) for scope in assignments_for_subject if _clean(scope.get("paralelo"))}
        journeys = {_clean(scope.get("jornada")) for scope in assignments_for_subject if _clean(scope.get("jornada"))}
        journey_codes = {
            _int(scope.get("cod_jornada"))
            for scope in assignments_for_subject
            if _int(scope.get("cod_jornada")) is not None
        }
        bucket["paralelo"] = next(iter(parallels)) if len(parallels) == 1 else "Varios"
        bucket["jornada"] = next(iter(journeys)) if len(journeys) == 1 else "Varias jornadas"
        bucket["cod_jornada"] = next(iter(journey_codes)) if len(journey_codes) == 1 else None
        grouped.append(bucket)

    return sorted(
        grouped,
        key=lambda item: (
            _clean(item.get("nombre_materia")),
            _clean(item.get("cod_materia") or item.get("codigo_materia")),
        ),
    )


@router.get("/teacher/me")
def teacher_profile(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
) -> dict[str, Any]:
    codigo_doc = _teacher_code(current_user)
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                """
                SELECT TOP (1)
                    TRY_CONVERT(varchar(50), d.codigo_doc) AS codigo_doc,
                    TRY_CONVERT(nvarchar(100), d.cedula_doc) AS cedula,
                    TRY_CONVERT(nvarchar(4000), d.apellidos_nombre) AS docente,
                    TRY_CONVERT(nvarchar(255), d.correo) AS correo,
                    TRY_CONVERT(nvarchar(255), d.correop) AS correo_personal,
                    TRY_CONVERT(nvarchar(100), d.telefono) AS telefono,
                    TRY_CONVERT(nvarchar(100), d.movil) AS movil,
                    TRY_CONVERT(nvarchar(255), d.TipoDocente) AS tipo_docente,
                    TRY_CONVERT(nvarchar(4000), d.Perfil) AS perfil
                FROM dbo.DATOSDOCENTE d
                WHERE TRY_CONVERT(int, d.codigo_doc) = ?
                """,
                codigo_doc,
            )
            row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="No se encontro el docente vinculado a la sesion")
        return {
            "teacher": {
                "codigo_doc": _clean(row.codigo_doc),
                "cedula": _clean(row.cedula),
                "docente": _clean(row.docente),
                "correo": _clean(row.correo),
                "correo_personal": _clean(row.correo_personal),
                "telefono": _clean(row.telefono),
                "movil": _clean(row.movil),
                "tipo_docente": _clean(row.tipo_docente),
                "perfil": _clean(row.perfil),
            }
        }
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error consultando perfil docente: {exc}") from exc


def _contract_value(value: Any) -> Any:
    if isinstance(value, Decimal):
        return float(value)
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    return value


def _teacher_contract_identity(current_user: SessionUser) -> dict[str, Any]:
    profile = teacher_profile(current_user).get("teacher", {})
    cedula = _clean(profile.get("cedula") or current_user.cedula)
    if not cedula:
        raise HTTPException(status_code=403, detail="La sesión no tiene una identificación docente vinculada")
    return {
        "codigo_doc": _clean(profile.get("codigo_doc") or current_user.codigo_doc),
        "cedula": cedula,
        "nombre": _clean(profile.get("docente") or current_user.nombres or current_user.login),
        "correo": _clean(profile.get("correo") or current_user.login),
        "correo_personal": _clean(profile.get("correo_personal")),
        "tipo_docente": _clean(profile.get("tipo_docente")),
        "perfil": _clean(profile.get("perfil")),
    }


def _ensure_teacher_contract_document_schema(cursor: pyodbc.Cursor) -> None:
    cursor.execute(
        """
        IF OBJECT_ID(N'rrhh.ContratoDocenteDocumento', N'U') IS NULL
        BEGIN
            CREATE TABLE rrhh.ContratoDocenteDocumento
            (
                ContratoDocumentoId BIGINT IDENTITY(1,1) NOT NULL
                    CONSTRAINT PK_ContratoDocenteDocumento PRIMARY KEY,
                ContratoDocenteId BIGINT NOT NULL,
                TipoDocumento VARCHAR(20) NOT NULL,
                ModalidadAcademica VARCHAR(20) NOT NULL,
                NombreArchivo NVARCHAR(260) NOT NULL,
                RutaInterna NVARCHAR(1000) NOT NULL,
                MimeType VARCHAR(100) NOT NULL
                    CONSTRAINT DF_ContratoDocenteDocumento_MimeType DEFAULT ('application/pdf'),
                TamanoBytes BIGINT NOT NULL,
                HashSha256 CHAR(64) NOT NULL,
                EsVigente BIT NOT NULL
                    CONSTRAINT DF_ContratoDocenteDocumento_EsVigente DEFAULT (1),
                UsuarioCarga NVARCHAR(256) NULL,
                FechaCarga DATETIME2 NOT NULL
                    CONSTRAINT DF_ContratoDocenteDocumento_FechaCarga DEFAULT (SYSDATETIME()),
                FirmanteDocumento NVARCHAR(300) NULL,
                FechaFirma DATETIME2 NULL,
                FirmaMotivo NVARCHAR(500) NULL,
                CONSTRAINT FK_ContratoDocenteDocumento_Contrato
                    FOREIGN KEY (ContratoDocenteId)
                    REFERENCES rrhh.ContratoDocente(ContratoDocenteId),
                CONSTRAINT CK_ContratoDocenteDocumento_Tipo
                    CHECK (TipoDocumento IN ('ORIGINAL', 'FIRMADO')),
                CONSTRAINT CK_ContratoDocenteDocumento_Modalidad
                    CHECK (ModalidadAcademica IN ('REGULAR', 'HOMOLOGACION')),
                CONSTRAINT CK_ContratoDocenteDocumento_Tamano
                    CHECK (TamanoBytes > 0)
            );
        END;

        IF NOT EXISTS
        (
            SELECT 1 FROM sys.indexes
            WHERE object_id = OBJECT_ID(N'rrhh.ContratoDocenteDocumento')
              AND name = N'IX_ContratoDocenteDocumento_ContratoVigente'
        )
        BEGIN
            CREATE INDEX IX_ContratoDocenteDocumento_ContratoVigente
                ON rrhh.ContratoDocenteDocumento
                (ContratoDocenteId, TipoDocumento, EsVigente, ContratoDocumentoId DESC);
        END;
        """
    )


def _sync_finance_teacher(cursor: pyodbc.Cursor, identity: dict[str, Any]) -> int:
    cursor.execute(
        """
        SELECT TOP (1) DocenteId
        FROM core.Docente
        WHERE LTRIM(RTRIM(NumeroIdentificacion)) = ?
        ORDER BY DocenteId DESC
        """,
        identity["cedula"],
    )
    row = cursor.fetchone()
    if row:
        teacher_id = int(row.DocenteId)
        cursor.execute(
            """
            UPDATE core.Docente
            SET CodigoDocente = TRY_CONVERT(decimal(18,0), ?),
                NombreCompleto = ?,
                Correo = NULLIF(?, N''),
                TipoDocente = NULLIF(?, N''),
                FuenteOrigen = 'INTECBDD',
                FechaSincronizacion = SYSDATETIME()
            WHERE DocenteId = ?
            """,
            identity.get("codigo_doc"),
            identity["nombre"],
            identity.get("correo"),
            identity.get("tipo_docente"),
            teacher_id,
        )
        return teacher_id

    cursor.execute(
        """
        SET NOCOUNT ON;
        DECLARE @DocenteCreado TABLE (DocenteId BIGINT NOT NULL);

        INSERT INTO core.Docente
        (
            CodigoDocente, NumeroIdentificacion, NombreCompleto, Correo,
            TipoDocente, FuenteOrigen, FechaSincronizacion
        )
        OUTPUT INSERTED.DocenteId INTO @DocenteCreado (DocenteId)
        VALUES
        (
            TRY_CONVERT(decimal(18,0), ?), ?, ?, NULLIF(?, N''),
            NULLIF(?, N''), 'INTECBDD', SYSDATETIME()
        );

        SELECT DocenteId FROM @DocenteCreado;
        """,
        identity.get("codigo_doc"),
        identity["cedula"],
        identity["nombre"],
        identity.get("correo"),
        identity.get("tipo_docente"),
    )
    created = cursor.fetchone()
    if not created:
        raise HTTPException(status_code=500, detail="No se pudo sincronizar el docente en la base financiera")
    return int(created.DocenteId)


def _teacher_contract_assignment(
    current_user: SessionUser,
    *,
    cod_anio_basica: int,
    codigo_periodo: int,
    codigo_materia: str,
    paralelo: str,
    cod_jornada: int | None,
) -> dict[str, Any]:
    codigo_doc = _teacher_code(current_user)
    subject_code = _clean(codigo_materia).upper()
    parallel = _clean(paralelo).upper()
    if not subject_code or not parallel:
        raise HTTPException(status_code=400, detail="Seleccione una asignación académica válida")
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                """
                SELECT TOP (1)
                    TRY_CONVERT(nvarchar(100), cxd.cod_Anio_Basica) AS codigo_carrera,
                    TRY_CONVERT(nvarchar(500), c.Nombre_Basica) AS nombre_carrera,
                    TRY_CONVERT(nvarchar(100), cxd.codigo_materia) AS codigo_materia_interno,
                    COALESCE(
                        NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), p.cod_materia))), N''),
                        TRY_CONVERT(nvarchar(100), cxd.codigo_materia)
                    ) AS codigo_materia,
                    TRY_CONVERT(nvarchar(500), p.Nomb_Materia) AS nombre_materia,
                    TRY_CONVERT(nvarchar(100), cxd.codigo_periodo) AS codigo_periodo,
                    TRY_CONVERT(nvarchar(500), pe.Detalle_Periodo) AS detalle_periodo,
                    TRY_CONVERT(nvarchar(100), pe.TipoMatricula) AS tipo_periodo,
                    TRY_CONVERT(nvarchar(40), cxd.Paralelo) AS paralelo,
                    TRY_CONVERT(int, cxd.Cod_Jornada) AS cod_jornada,
                    TRY_CONVERT(nvarchar(100), j.DetalleJ) AS jornada,
                    TRY_CONVERT(decimal(18,2), p.Horas) AS horas_planificadas
                FROM dbo.CARRERAXDOCENTE cxd
                LEFT JOIN dbo.CARRERAS c
                  ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                LEFT JOIN dbo.PENSUM p
                  ON TRY_CONVERT(int, p.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                 AND TRY_CONVERT(int, p.codigo_materia) = TRY_CONVERT(int, cxd.codigo_materia)
                LEFT JOIN dbo.PERIODO pe
                  ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cxd.codigo_periodo)
                LEFT JOIN dbo.JORNADA j
                  ON TRY_CONVERT(int, j.NumJ) = TRY_CONVERT(int, cxd.Cod_Jornada)
                WHERE TRY_CONVERT(int, cxd.codigo_doc) = ?
                  AND TRY_CONVERT(int, cxd.cod_Anio_Basica) = ?
                  AND TRY_CONVERT(int, cxd.codigo_periodo) = ?
                  AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(40), cxd.Paralelo)))) = ?
                  AND (? IS NULL OR TRY_CONVERT(int, cxd.Cod_Jornada) = ?)
                  AND
                  (
                      UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), cxd.codigo_materia)))) = ?
                      OR UPPER(LTRIM(RTRIM(COALESCE(
                          NULLIF(TRY_CONVERT(nvarchar(100), p.cod_materia), N''),
                          TRY_CONVERT(nvarchar(100), cxd.codigo_materia)
                      )))) = ?
                  )
                ORDER BY TRY_CONVERT(int, cxd.codigo_materia)
                """,
                codigo_doc,
                cod_anio_basica,
                codigo_periodo,
                parallel,
                cod_jornada,
                cod_jornada,
                subject_code,
                subject_code,
            )
            row = cursor.fetchone()
    except pyodbc.Error as exc:
        logger.exception("No se pudo validar la asignación del contrato docente")
        raise HTTPException(status_code=503, detail="No se pudo validar la asignación académica seleccionada") from exc

    if not row:
        raise HTTPException(
            status_code=403,
            detail="El curso seleccionado no pertenece al docente autenticado",
        )
    is_homologation = _is_homologation_type(row.tipo_periodo, row.detalle_periodo)
    return {
        "codigo_carrera": _clean(row.codigo_carrera),
        "nombre_carrera": _clean(row.nombre_carrera),
        "codigo_materia_interno": _clean(row.codigo_materia_interno),
        "codigo_materia": _clean(row.codigo_materia),
        "nombre_materia": _clean(row.nombre_materia),
        "codigo_periodo": _clean(row.codigo_periodo),
        "detalle_periodo": _clean(row.detalle_periodo),
        "tipo_periodo": _clean(row.tipo_periodo),
        "modalidad_academica": "HOMOLOGACION" if is_homologation else "REGULAR",
        "paralelo": _clean(row.paralelo),
        "cod_jornada": _int(row.cod_jornada),
        "jornada": _clean(row.jornada),
        "horas_planificadas": _number(row.horas_planificadas) or 0,
    }


def _validate_teacher_contract_pdf(filename: str | None, content: bytes) -> str:
    original_name = Path(_clean(filename)).name
    if not original_name.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="El contrato debe ser un archivo PDF")
    if not content:
        raise HTTPException(status_code=400, detail="El archivo PDF está vacío")
    if len(content) > _TEACHER_CONTRACT_MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="El contrato PDF debe pesar máximo 15 MB")
    if not content.startswith(b"%PDF-"):
        raise HTTPException(status_code=400, detail="El archivo seleccionado no contiene un PDF válido")
    try:
        IncrementalPdfFileWriter(BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="No se pudo abrir la estructura del PDF del contrato") from exc
    return original_name[:260]


def _validate_signed_teacher_document_pdf(content: bytes, label: str) -> None:
    if not content:
        raise HTTPException(status_code=400, detail=f"El PDF firmado de {label} está vacío")
    if len(content) > _SIGNED_TEACHER_DOCUMENT_MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"El PDF firmado de {label} excede 100 MB")
    if not content.startswith(b"%PDF-"):
        raise HTTPException(status_code=400, detail=f"El archivo de {label} no contiene un PDF válido")


def _signed_teacher_documents_archive(
    compliance_pdf: bytes,
    grades_pdf: bytes,
    contract_pdf: bytes,
) -> bytes:
    documents = (
        ("informe-cumplimiento-firmado.pdf", compliance_pdf, "informe de cumplimiento", "FirmaDocente"),
        (
            "reporte-notas-secretaria-firmado.pdf",
            grades_pdf,
            "reporte de notas",
            "FirmaDocenteReporteNotas",
        ),
        (
            "contrato-docente-firmado.pdf",
            contract_pdf,
            "contrato docente",
            "FirmaDocenteContratoInforme",
        ),
    )
    output = BytesIO()
    with ZipFile(output, "w", compression=ZIP_DEFLATED) as archive:
        for filename, content, label, signature_field in documents:
            _validate_signed_teacher_document_pdf(content, label)
            _assert_pdf_signature_field(content, signature_field)
            archive.writestr(filename, content)
    return output.getvalue()


def _teacher_signed_documents_folder(
    identity: dict[str, Any],
    *,
    subject_code: str = "",
    subject_name: str = "",
    period_codes: list[str] | None = None,
    signed_at: datetime | None = None,
) -> str:
    teacher_document = re.sub(r"\D+", "", _clean(identity.get("cedula")))
    teacher_document = teacher_document or _clean(identity.get("codigo_doc")) or "SIN-ID"
    teacher_name = graph_safe_folder_part(
        _clean(identity.get("nombre")),
        f"DOCENTE {teacher_document}",
        max_length=max(30, 100 - len(teacher_document) - 3),
    )
    teacher_folder = graph_safe_folder_part(
        f"{teacher_name} - {teacher_document}",
        teacher_document,
        max_length=110,
    )

    normalized_subject_code = _clean(subject_code)
    normalized_subject_name = _clean(subject_name)
    subject_label = normalized_subject_name
    if (
        normalized_subject_code
        and normalized_subject_code.upper() not in normalized_subject_name.upper()
    ):
        subject_label = (
            f"{normalized_subject_name} - {normalized_subject_code}"
            if normalized_subject_name
            else normalized_subject_code
        )
    subject_folder = graph_safe_folder_part(subject_label, "ASIGNATURA", max_length=110)

    unique_periods: list[str] = []
    for period_code in period_codes or []:
        normalized_period = _clean(period_code)
        if normalized_period and normalized_period not in unique_periods:
            unique_periods.append(normalized_period)
    period_folder = graph_safe_folder_part(" - ".join(unique_periods), "SIN PERIODO", max_length=110)
    operation_time = signed_at or datetime.now(timezone.utc)
    operation_folder = operation_time.astimezone(timezone.utc).strftime("FIRMA %Y%m%d-%H%M%S-%f UTC")

    return "/".join(
        [
            _TEACHER_DOCUMENT_ONEDRIVE_ROOT,
            teacher_folder,
            "DOCUMENTOS FIRMADOS",
            subject_folder,
            period_folder,
            operation_folder,
        ]
    )


def _store_signed_teacher_documents_onedrive(
    *,
    identity: dict[str, Any],
    compliance_pdf: bytes,
    grades_pdf: bytes,
    contract_pdf: bytes,
    archive_bytes: bytes,
    subject_code: str = "",
    subject_name: str = "",
    period_codes: list[str] | None = None,
) -> dict[str, Any]:
    folder_path = _teacher_signed_documents_folder(
        identity,
        subject_code=subject_code,
        subject_name=subject_name,
        period_codes=period_codes,
    )
    folder = ensure_graph_document_folder(folder_path)
    documents = (
        ("informe-cumplimiento-firmado.pdf", compliance_pdf, "application/pdf"),
        ("reporte-notas-secretaria-firmado.pdf", grades_pdf, "application/pdf"),
        ("contrato-docente-firmado.pdf", contract_pdf, "application/pdf"),
        ("documentos-docente-firmados.zip", archive_bytes, "application/zip"),
    )
    uploaded_items: list[dict[str, Any]] = []
    try:
        for filename, content, content_type in documents:
            item = upload_graph_document_bytes(
                f"{folder_path}/{filename}",
                content,
                content_type,
            )
            if not _clean(item.get("id")):
                raise RuntimeError(f"Microsoft Graph no confirmó el archivo {filename}.")
            uploaded_items.append(item)
    except Exception:
        for item in reversed(uploaded_items):
            try:
                delete_graph_document_item(_clean(item.get("id")))
            except Exception:
                logger.warning(
                    "No se pudo revertir el documento parcial %s en OneDrive.",
                    _clean(item.get("id")),
                    exc_info=True,
                )
        try:
            delete_graph_document_item(_clean(folder.get("id")))
        except Exception:
            logger.warning(
                "No se pudo eliminar la carpeta de operación incompleta %s en OneDrive.",
                _clean(folder.get("id")),
                exc_info=True,
            )
        raise

    return {
        "folder_path": folder_path,
        "folder": folder,
        "items": uploaded_items,
    }


def _contract_text_key(value: Any) -> str:
    normalized = unicodedata.normalize("NFD", _clean(value)).upper()
    return re.sub(r"[^A-Z0-9]", "", "".join(char for char in normalized if not unicodedata.combining(char)))


def _contract_pdf_lines(value: str) -> list[str]:
    normalized = unicodedata.normalize("NFC", value or "").replace("\x00", " ").replace("\u00a0", " ")
    return [re.sub(r"\s+", " ", line).strip() for line in normalized.splitlines() if line.strip()]


def _extract_teacher_contract_pdf_text(content: bytes) -> str:
    try:
        import pypdfium2 as pdfium

        document = pdfium.PdfDocument(content)
        try:
            if len(document) < 1:
                raise HTTPException(status_code=400, detail="El contrato PDF no contiene páginas")
            if len(document) > _TEACHER_CONTRACT_MAX_PAGES:
                raise HTTPException(
                    status_code=400,
                    detail=f"El contrato supera el límite de {_TEACHER_CONTRACT_MAX_PAGES} páginas",
                )
            text_parts: list[str] = []
            for page_index in range(len(document)):
                page = document[page_index]
                try:
                    text_page = page.get_textpage()
                    try:
                        text_parts.append(text_page.get_text_range())
                    finally:
                        text_page.close()
                finally:
                    page.close()
        finally:
            document.close()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail="No se pudo analizar el texto del contrato PDF") from exc

    extracted = "\n".join(text_parts).strip()
    if not extracted:
        raise HTTPException(
            status_code=400,
            detail="El contrato no contiene texto seleccionable para reconocer número, docente y fechas",
        )
    return extracted


def _contract_date_from_parts(day: str, month: str, year: str) -> date | None:
    month_key = _contract_text_key(month)
    month_number = _CONTRACT_SPANISH_MONTHS.get(month_key)
    if month_number is None and month_key.isdigit():
        month_number = int(month_key)
    try:
        return date(int(year), int(month_number or 0), int(day))
    except ValueError:
        return None


def _contract_labeled_date(text: str, labels: str) -> date | None:
    match = re.search(
        rf"(?:{labels})(?:\s+(?:EL|DEL|DE))?\s*[:.-]?\s*(\d{{1,2}})\s*[/.-]\s*([A-ZÁÉÍÓÚÑ]{{2,12}}|\d{{1,2}})\s*[/.-]\s*(\d{{4}})",
        text,
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    return _contract_date_from_parts(match.group(1), match.group(2), match.group(3))


def _contract_decimal(value: str) -> float | None:
    cleaned = re.sub(r"[^0-9.,]", "", value or "")
    if not cleaned:
        return None
    if "," in cleaned and "." in cleaned:
        if cleaned.rfind(",") > cleaned.rfind("."):
            cleaned = cleaned.replace(".", "").replace(",", ".")
        else:
            cleaned = cleaned.replace(",", "")
    elif "," in cleaned:
        integer, decimal_part = cleaned.rsplit(",", 1)
        cleaned = f"{integer.replace(',', '')}.{decimal_part}" if len(decimal_part) <= 2 else cleaned.replace(",", "")
    try:
        return round(float(cleaned), 2)
    except ValueError:
        return None


def _parse_teacher_contract_text(text: str) -> dict[str, Any]:
    lines = _contract_pdf_lines(text)
    flat_text = " ".join(lines)
    upper_text = flat_text.upper()

    contract_number = ""
    for line in lines:
        number_match = re.search(r"\bCONTRATO\s*:\s*(.+)$", line, flags=re.IGNORECASE)
        if number_match:
            contract_number = re.sub(
                r"^(?:N(?:O|º|°)?\.?\s*)",
                "",
                number_match.group(1).strip(),
                flags=re.IGNORECASE,
            ).strip(" .")
            break
    if not contract_number:
        number_match = re.search(
            r"\bCONTRATO\s*:\s*(?:N(?:O|º|°)?\.?\s*)?(.+?)(?=\s+CONTRATO\s+DE\b|\s+COMPARECIENTES\b)",
            flat_text,
            flags=re.IGNORECASE,
        )
        contract_number = number_match.group(1).strip(" .") if number_match else ""

    identity_match = re.search(
        r"(?:C[ÉE]DULA(?:\s+DE\s+CIUDADAN[IÍ]A)?|CIUDADAN[IÍ]A|IDENTIFICACI[ÓO]N)[^0-9]{0,45}([0-9]{10,13})",
        flat_text,
        flags=re.IGNORECASE,
    )
    subject_match = re.search(r"\b(VGA(?:-[A-Z0-9]+){2,5})\b", upper_text)
    start_date = _contract_labeled_date(flat_text, r"INICIA|INICIO|FECHA\s+DE\s+INICIO")
    end_date = _contract_labeled_date(flat_text, r"TERMINA|FINALIZA|FECHA\s+DE\s+FIN|FECHA\s+FINAL")

    value_match = re.search(
        r"VALOR\s+TOTAL.{0,300}?(?:US\$|USD|D[ÓO]LARES?)\s*[:.]?\s*([0-9][0-9.,]*)",
        flat_text,
        flags=re.IGNORECASE,
    )
    if not value_match:
        value_match = re.search(r"\(\s*US\$\s*([0-9][0-9.,]*)\s*\)", flat_text, flags=re.IGNORECASE)
    value_total = _contract_decimal(value_match.group(1)) if value_match else None
    modality_source = f"{contract_number} {' '.join(lines[:12])}"
    modality = "HOMOLOGACION" if re.search(r"\bHOMO(?:LOGACI[ÓO]N|LOGADO|LOGADA)?\b", modality_source, re.IGNORECASE) else "REGULAR"

    analysis = {
        "numero_contrato": contract_number,
        "cedula": identity_match.group(1) if identity_match else "",
        "codigo_materia": subject_match.group(1) if subject_match else "",
        "modalidad_academica": modality,
        "fecha_inicio": start_date.isoformat() if start_date else "",
        "fecha_fin": end_date.isoformat() if end_date else "",
        "valor_total": value_total,
    }
    missing_labels = {
        "numero_contrato": "número de contrato",
        "cedula": "cédula del docente",
        "codigo_materia": "código de materia",
        "fecha_inicio": "fecha de inicio",
        "fecha_fin": "fecha de finalización",
    }
    analysis["campos_detectados"] = [key for key in missing_labels if analysis.get(key)]
    analysis["advertencias"] = [
        f"No se reconoció {label}; complete o verifique el dato antes de guardar."
        for key, label in missing_labels.items()
        if not analysis.get(key)
    ]
    return analysis


def _validate_teacher_contract_analysis(
    analysis: dict[str, Any],
    identity: dict[str, Any],
    assignment: dict[str, Any] | None = None,
) -> None:
    detected_identity = _contract_text_key(analysis.get("cedula"))
    teacher_identity = _contract_text_key(identity.get("cedula"))
    if detected_identity and detected_identity != teacher_identity:
        raise HTTPException(
            status_code=400,
            detail="La cédula declarada en el contrato no corresponde al docente autenticado",
        )
    if assignment is None:
        return
    detected_subject = _contract_text_key(analysis.get("codigo_materia"))
    accepted_subjects = {
        _contract_text_key(assignment.get("codigo_materia")),
        _contract_text_key(assignment.get("codigo_materia_interno")),
    }
    accepted_subjects.discard("")
    if detected_subject and detected_subject not in accepted_subjects:
        raise HTTPException(
            status_code=400,
            detail="El código de materia del contrato no corresponde al curso seleccionado",
        )
    if analysis.get("modalidad_academica") != assignment.get("modalidad_academica"):
        raise HTTPException(
            status_code=400,
            detail="La modalidad reconocida en el contrato no corresponde al periodo seleccionado",
        )


def _teacher_contract_document_relative_path(
    cedula: str,
    contract_id: int,
    document_type: str,
    filename: str,
) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    safe_name = _safe_filename(Path(filename).stem)[:100]
    return Path(_safe_filename(cedula)) / str(contract_id) / document_type.lower() / f"{timestamp}-{safe_name}.pdf"


def _teacher_contract_document_path(relative_path: str | Path) -> Path:
    root = _TEACHER_CONTRACT_STORAGE_ROOT.resolve()
    candidate = (root / Path(relative_path)).resolve()
    try:
        candidate.relative_to(root)
    except ValueError as exc:
        raise HTTPException(status_code=500, detail="La ruta interna del contrato no es válida") from exc
    return candidate


def _teacher_contract_owned(
    cursor: pyodbc.Cursor,
    *,
    contract_id: int,
    teacher_id: int,
) -> Any:
    cursor.execute(
        """
        SELECT TOP (1)
            c.ContratoDocenteId AS contrato_id,
            c.NumeroContrato AS numero_contrato,
            c.RutaContratoFirmado AS ruta_contrato_firmado
        FROM rrhh.ContratoDocente c
        WHERE c.ContratoDocenteId = ? AND c.DocenteId = ?
        """,
        contract_id,
        teacher_id,
    )
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="No se encontró el contrato para el docente autenticado")
    return row


def _teacher_contract_document(
    cursor: pyodbc.Cursor,
    *,
    contract_id: int,
    version: str,
) -> Any:
    requested_type = "ORIGINAL" if version == "original" else "FIRMADO" if version == "signed" else None
    cursor.execute(
        """
        SELECT TOP (1)
            ContratoDocumentoId AS documento_id,
            TipoDocumento AS tipo_documento,
            ModalidadAcademica AS modalidad_academica,
            NombreArchivo AS nombre_archivo,
            RutaInterna AS ruta_interna,
            MimeType AS mime_type,
            TamanoBytes AS tamano_bytes,
            HashSha256 AS hash_sha256,
            FechaCarga AS fecha_carga,
            FirmanteDocumento AS firmante_documento,
            FechaFirma AS fecha_firma
        FROM rrhh.ContratoDocenteDocumento
        WHERE ContratoDocenteId = ?
          AND EsVigente = 1
          AND (? IS NULL OR TipoDocumento = ?)
        ORDER BY
            CASE WHEN TipoDocumento = 'FIRMADO' THEN 0 ELSE 1 END,
            ContratoDocumentoId DESC
        """,
        contract_id,
        requested_type,
        requested_type,
    )
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="El contrato todavía no tiene un documento disponible")
    return row


@router.get("/teacher/contracts")
def teacher_contracts(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
) -> dict[str, Any]:
    identity = _teacher_contract_identity(current_user)

    try:
        with get_finance_connection() as conn:
            cursor = conn.cursor()
            _ensure_teacher_contract_document_schema(cursor)
            teacher_id = _sync_finance_teacher(cursor, identity)
            conn.commit()
            cursor.execute(
                """
                SELECT TOP (1)
                    DocenteId AS docente_id,
                    NumeroIdentificacion AS cedula,
                    NombreCompleto AS nombre,
                    Correo AS correo,
                    TipoDocente AS tipo_docente,
                    RelacionLaboral AS relacion_laboral,
                    TiempoDedicacion AS tiempo_dedicacion
                FROM core.Docente
                WHERE DocenteId = ?
                ORDER BY FechaSincronizacion DESC
                """,
                teacher_id,
            )
            teacher_row = cursor.fetchone()
            if not teacher_row:
                raise HTTPException(status_code=500, detail="No se pudo recuperar el docente sincronizado")

            cursor.execute(
                """
                SELECT
                    c.ContratoDocenteId AS contrato_id,
                    c.NumeroContrato AS numero_contrato,
                    tc.Codigo AS tipo_codigo,
                    tc.Nombre AS tipo_nombre,
                    ec.Codigo AS estado_codigo,
                    ec.Nombre AS estado_nombre,
                    c.CodigoPeriodo AS codigo_periodo,
                    c.FechaInicio AS fecha_inicio,
                    c.FechaFin AS fecha_fin,
                    c.ValorHoraClase AS valor_hora_clase,
                    c.ValorMensual AS valor_mensual,
                    c.ValorTotalContrato AS valor_total_contrato,
                    c.ResponsableContratacion AS responsable_contratacion,
                    c.Observacion AS observacion,
                    c.RutaContratoFirmado AS ruta_contrato_firmado,
                    COALESCE(doc_firmado.ModalidadAcademica, doc_original.ModalidadAcademica) AS modalidad_academica,
                    doc_original.NombreArchivo AS nombre_documento_original,
                    doc_original.FechaCarga AS fecha_documento_original,
                    doc_firmado.NombreArchivo AS nombre_documento_firmado,
                    doc_firmado.FechaFirma AS fecha_documento_firmado,
                    cc.ContratoClaseId AS clase_id,
                    cc.CodigoCarrera AS codigo_carrera,
                    cc.NombreCarrera AS nombre_carrera,
                    cc.CodigoMateria AS codigo_materia,
                    cc.NombreMateria AS nombre_materia,
                    cc.CodigoPeriodo AS clase_periodo,
                    cc.Paralelo AS paralelo,
                    cc.Jornada AS jornada,
                    cc.HorasPlanificadas AS horas_planificadas,
                    cc.HorasEjecutadas AS horas_ejecutadas,
                    cc.ValorHora AS clase_valor_hora,
                    cc.ValorTotalPlanificado AS valor_total_planificado,
                    cc.EstadoClaseCodigo AS estado_clase,
                    cc.Observacion AS clase_observacion
                FROM rrhh.ContratoDocente c
                INNER JOIN cat.TipoContratoDocente tc
                    ON tc.TipoContratoDocenteId = c.TipoContratoDocenteId
                INNER JOIN cat.EstadoContratoDocente ec
                    ON ec.EstadoContratoDocenteId = c.EstadoContratoDocenteId
                LEFT JOIN rrhh.ContratoDocenteClase cc
                    ON cc.ContratoDocenteId = c.ContratoDocenteId
                OUTER APPLY
                (
                    SELECT TOP (1)
                        d.ModalidadAcademica,
                        d.NombreArchivo,
                        d.FechaCarga
                    FROM rrhh.ContratoDocenteDocumento d
                    WHERE d.ContratoDocenteId = c.ContratoDocenteId
                      AND d.TipoDocumento = 'ORIGINAL'
                      AND d.EsVigente = 1
                    ORDER BY d.ContratoDocumentoId DESC
                ) doc_original
                OUTER APPLY
                (
                    SELECT TOP (1)
                        d.ModalidadAcademica,
                        d.NombreArchivo,
                        d.FechaFirma
                    FROM rrhh.ContratoDocenteDocumento d
                    WHERE d.ContratoDocenteId = c.ContratoDocenteId
                      AND d.TipoDocumento = 'FIRMADO'
                      AND d.EsVigente = 1
                    ORDER BY d.ContratoDocumentoId DESC
                ) doc_firmado
                WHERE c.DocenteId = ?
                ORDER BY c.FechaInicio DESC, c.ContratoDocenteId DESC, cc.NombreMateria
                """,
                teacher_row.docente_id,
            )
            rows = cursor.fetchall()

        contracts_by_id: dict[int, dict[str, Any]] = {}
        for row in rows:
            contract_id = int(row.contrato_id)
            contract = contracts_by_id.get(contract_id)
            if contract is None:
                contract = {
                    "contrato_id": contract_id,
                    "numero_contrato": _clean(row.numero_contrato),
                    "tipo_codigo": _clean(row.tipo_codigo),
                    "tipo_nombre": _clean(row.tipo_nombre),
                    "estado_codigo": _clean(row.estado_codigo),
                    "estado_nombre": _clean(row.estado_nombre),
                    "codigo_periodo": _clean(row.codigo_periodo),
                    "fecha_inicio": _contract_value(row.fecha_inicio),
                    "fecha_fin": _contract_value(row.fecha_fin),
                    "valor_hora_clase": _contract_value(row.valor_hora_clase),
                    "valor_mensual": _contract_value(row.valor_mensual),
                    "valor_total_contrato": _contract_value(row.valor_total_contrato),
                    "responsable_contratacion": _clean(row.responsable_contratacion),
                    "observacion": _clean(row.observacion),
                    "ruta_contrato_firmado": _clean(row.ruta_contrato_firmado),
                    "modalidad_academica": _clean(row.modalidad_academica),
                    "tiene_documento_original": bool(_clean(row.nombre_documento_original)),
                    "nombre_documento_original": _clean(row.nombre_documento_original),
                    "fecha_documento_original": _contract_value(row.fecha_documento_original),
                    "tiene_documento_firmado": bool(_clean(row.nombre_documento_firmado)),
                    "nombre_documento_firmado": _clean(row.nombre_documento_firmado),
                    "fecha_documento_firmado": _contract_value(row.fecha_documento_firmado),
                    "clases": [],
                }
                contracts_by_id[contract_id] = contract
            if row.clase_id is not None:
                contract["clases"].append({
                    "clase_id": int(row.clase_id),
                    "codigo_carrera": _clean(row.codigo_carrera),
                    "nombre_carrera": _clean(row.nombre_carrera),
                    "codigo_materia": _clean(row.codigo_materia),
                    "nombre_materia": _clean(row.nombre_materia),
                    "codigo_periodo": _clean(row.clase_periodo),
                    "paralelo": _clean(row.paralelo),
                    "jornada": _clean(row.jornada),
                    "horas_planificadas": _contract_value(row.horas_planificadas),
                    "horas_ejecutadas": _contract_value(row.horas_ejecutadas),
                    "valor_hora": _contract_value(row.clase_valor_hora),
                    "valor_total_planificado": _contract_value(row.valor_total_planificado),
                    "estado": _clean(row.estado_clase),
                    "observacion": _clean(row.clase_observacion),
                })

        return {
            "teacher": {
                "docente_id": int(teacher_row.docente_id),
                "cedula": _clean(teacher_row.cedula),
                "nombre": _clean(teacher_row.nombre),
                "correo": _clean(teacher_row.correo),
                "tipo_docente": _clean(teacher_row.tipo_docente),
                "relacion_laboral": _clean(teacher_row.relacion_laboral),
                "tiempo_dedicacion": _clean(teacher_row.tiempo_dedicacion),
            },
            "contracts": list(contracts_by_id.values()),
        }
    except (pyodbc.Error, RuntimeError) as exc:
        logger.exception("No se pudieron consultar los contratos del docente")
        raise HTTPException(
            status_code=503,
            detail="No se pudo consultar la información contractual en este momento.",
        ) from exc


@router.post("/teacher/contracts/analyze-document")
async def teacher_analyze_contract_document(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    contrato: Annotated[UploadFile, File()],
) -> dict[str, Any]:
    pdf_bytes = await contrato.read()
    original_name = _validate_teacher_contract_pdf(contrato.filename, pdf_bytes)
    identity = _teacher_contract_identity(current_user)
    analysis = _parse_teacher_contract_text(_extract_teacher_contract_pdf_text(pdf_bytes))
    _validate_teacher_contract_analysis(analysis, identity)
    return {
        "ok": True,
        "nombre_archivo": original_name,
        "docente_coincide": bool(analysis.get("cedula")),
        **analysis,
    }


@router.post("/teacher/contracts/sign-uploaded")
async def teacher_sign_uploaded_contract_document(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    contrato: Annotated[UploadFile, File()],
    certificado: Annotated[UploadFile, File()],
    contrasena_certificado: Annotated[str, Form(min_length=1, max_length=256)],
    firma_motivo: Annotated[str, Form(max_length=200)] = "Aceptación y firma de contrato docente",
    firma_ubicacion: Annotated[str, Form(max_length=120)] = "Quito, Ecuador",
    firma_contacto: Annotated[str, Form(max_length=200)] = "",
) -> StreamingResponse:
    contract_bytes = await contrato.read()
    original_name = _validate_teacher_contract_pdf(contrato.filename, contract_bytes)
    identity = _teacher_contract_identity(current_user)
    analysis = _parse_teacher_contract_text(_extract_teacher_contract_pdf_text(contract_bytes))
    _validate_teacher_contract_analysis(analysis, identity)

    certificate_name = _clean(certificado.filename).lower()
    if not certificate_name.endswith((".p12", ".pfx")):
        raise HTTPException(status_code=400, detail="Seleccione un certificado con extensión .p12 o .pfx")
    certificate_bytes = await certificado.read()
    if not certificate_bytes:
        raise HTTPException(status_code=400, detail="El archivo de certificado está vacío")
    if len(certificate_bytes) > _TEACHER_CONTRACT_CERTIFICATE_MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="El archivo .p12 debe pesar máximo 2 MB")

    signed_pdf = await _sign_pdf_with_pkcs12(
        pdf_bytes=contract_bytes,
        pkcs12_bytes=certificate_bytes,
        password=contrasena_certificado,
        current_user=current_user,
        reason=firma_motivo,
        location=firma_ubicacion,
        contact=firma_contacto,
        signature_box=(335, 42, 565, 112),
        field_name="FirmaDocenteContratoInforme",
        readable_field_name="Firma electrónica del docente contratista",
    )
    filename = f"{_safe_filename(Path(original_name).stem)}-firmado-docente.pdf"
    return StreamingResponse(
        BytesIO(signed_pdf),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Cache-Control": "no-store, private",
            "X-Content-Type-Options": "nosniff",
        },
    )


@router.post("/teacher/contracts/document")
async def teacher_upload_contract_document(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    numero_contrato: Annotated[str, Form(min_length=1, max_length=100)],
    cod_anio_basica: Annotated[int, Form(gt=0)],
    codigo_periodo: Annotated[int, Form(gt=0)],
    codigo_materia: Annotated[str, Form(min_length=1, max_length=100)],
    paralelo: Annotated[str, Form(min_length=1, max_length=20)],
    modalidad_academica: Annotated[Literal["REGULAR", "HOMOLOGACION"], Form()],
    fecha_inicio: Annotated[str, Form(min_length=10, max_length=10)],
    fecha_fin: Annotated[str, Form(min_length=10, max_length=10)],
    horas_planificadas: Annotated[float, Form(ge=0, le=10000)],
    valor_hora: Annotated[float, Form(ge=0, le=100000)],
    valor_total: Annotated[float, Form(ge=0, le=10000000)],
    contrato: Annotated[UploadFile, File()],
    cod_jornada: Annotated[int | None, Form()] = None,
    responsable_contratacion: Annotated[str, Form(max_length=200)] = "",
    observacion: Annotated[str, Form(max_length=1000)] = "",
) -> dict[str, Any]:
    try:
        start_date = date.fromisoformat(fecha_inicio)
        end_date = date.fromisoformat(fecha_fin)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Las fechas del contrato no tienen un formato válido") from exc
    if end_date < start_date:
        raise HTTPException(status_code=400, detail="La fecha final no puede ser anterior a la fecha inicial")

    identity = _teacher_contract_identity(current_user)
    assignment = _teacher_contract_assignment(
        current_user,
        cod_anio_basica=cod_anio_basica,
        codigo_periodo=codigo_periodo,
        codigo_materia=codigo_materia,
        paralelo=paralelo,
        cod_jornada=cod_jornada,
    )
    if assignment["modalidad_academica"] != modalidad_academica:
        expected = "homologación" if assignment["modalidad_academica"] == "HOMOLOGACION" else "regular"
        raise HTTPException(
            status_code=400,
            detail=f"La modalidad del contrato debe ser {expected} según el periodo seleccionado",
        )

    pdf_bytes = await contrato.read()
    original_name = _validate_teacher_contract_pdf(contrato.filename, pdf_bytes)
    analysis = _parse_teacher_contract_text(_extract_teacher_contract_pdf_text(pdf_bytes))
    _validate_teacher_contract_analysis(analysis, identity, assignment)
    if analysis.get("numero_contrato") and _contract_text_key(analysis["numero_contrato"]) != _contract_text_key(numero_contrato):
        raise HTTPException(status_code=400, detail="El número ingresado no coincide con el contrato PDF")
    if analysis.get("fecha_inicio") and analysis["fecha_inicio"] != start_date.isoformat():
        raise HTTPException(status_code=400, detail="La fecha inicial ingresada no coincide con el contrato PDF")
    if analysis.get("fecha_fin") and analysis["fecha_fin"] != end_date.isoformat():
        raise HTTPException(status_code=400, detail="La fecha final ingresada no coincide con el contrato PDF")
    contract_number = _clean(numero_contrato)
    created_path: Path | None = None
    try:
        with get_finance_connection() as conn:
            cursor = conn.cursor()
            _ensure_teacher_contract_document_schema(cursor)
            teacher_id = _sync_finance_teacher(cursor, identity)
            cursor.execute(
                """
                SELECT TOP (1)
                    c.ContratoDocenteId AS contrato_id,
                    c.DocenteId AS docente_id,
                    CASE WHEN EXISTS
                    (
                        SELECT 1 FROM rrhh.ContratoDocenteDocumento d
                        WHERE d.ContratoDocenteId = c.ContratoDocenteId
                          AND d.TipoDocumento = 'FIRMADO'
                          AND d.EsVigente = 1
                    ) THEN 1 ELSE 0 END AS tiene_firma
                FROM rrhh.ContratoDocente c
                WHERE LTRIM(RTRIM(c.NumeroContrato)) = ?
                ORDER BY c.ContratoDocenteId DESC
                """,
                contract_number,
            )
            existing = cursor.fetchone()
            if existing and int(existing.docente_id) != teacher_id:
                raise HTTPException(status_code=409, detail="El número de contrato ya pertenece a otro docente")
            if existing and bool(existing.tiene_firma):
                raise HTTPException(
                    status_code=409,
                    detail="El contrato ya fue firmado y no puede reemplazarse sin una reapertura administrativa",
                )

            cursor.execute(
                """
                SELECT TOP (1) TipoContratoDocenteId
                FROM cat.TipoContratoDocente
                WHERE Codigo = 'POR_HORAS'
                """
            )
            type_row = cursor.fetchone()
            cursor.execute(
                """
                SELECT TOP (1) EstadoContratoDocenteId
                FROM cat.EstadoContratoDocente
                WHERE Codigo = 'BORRADOR'
                """
            )
            state_row = cursor.fetchone()
            if not type_row or not state_row:
                raise HTTPException(status_code=500, detail="Faltan catálogos contractuales en la base financiera")

            if existing:
                contract_id = int(existing.contrato_id)
                cursor.execute(
                    """
                    UPDATE rrhh.ContratoDocente
                    SET TipoContratoDocenteId = ?,
                        EstadoContratoDocenteId = ?,
                        CodigoPeriodo = ?,
                        FechaInicio = ?,
                        FechaFin = ?,
                        ValorHoraClase = ?,
                        ValorTotalContrato = ?,
                        ResponsableContratacion = NULLIF(?, N''),
                        Observacion = NULLIF(?, N''),
                        RutaContratoFirmado = NULL
                    WHERE ContratoDocenteId = ? AND DocenteId = ?
                    """,
                    int(type_row.TipoContratoDocenteId),
                    int(state_row.EstadoContratoDocenteId),
                    assignment["codigo_periodo"],
                    start_date,
                    end_date,
                    valor_hora,
                    valor_total,
                    _clean(responsable_contratacion),
                    _clean(observacion),
                    contract_id,
                    teacher_id,
                )
                cursor.execute("DELETE FROM rrhh.ContratoDocenteClase WHERE ContratoDocenteId = ?", contract_id)
            else:
                cursor.execute(
                    """
                    SET NOCOUNT ON;
                    DECLARE @ContratoCreado TABLE (ContratoDocenteId BIGINT NOT NULL);

                    INSERT INTO rrhh.ContratoDocente
                    (
                        DocenteId, TipoContratoDocenteId, EstadoContratoDocenteId,
                        NumeroContrato, CodigoPeriodo, FechaInicio, FechaFin,
                        ValorHoraClase, ValorTotalContrato, ResponsableContratacion,
                        Observacion, UsuarioCreacion
                    )
                    OUTPUT INSERTED.ContratoDocenteId
                        INTO @ContratoCreado (ContratoDocenteId)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, NULLIF(?, N''), NULLIF(?, N''), ?)

                    SELECT ContratoDocenteId FROM @ContratoCreado;
                    """,
                    teacher_id,
                    int(type_row.TipoContratoDocenteId),
                    int(state_row.EstadoContratoDocenteId),
                    contract_number,
                    assignment["codigo_periodo"],
                    start_date,
                    end_date,
                    valor_hora,
                    valor_total,
                    _clean(responsable_contratacion),
                    _clean(observacion),
                    _clean(current_user.login),
                )
                inserted = cursor.fetchone()
                if not inserted:
                    raise HTTPException(status_code=500, detail="No se pudo crear el registro contractual")
                contract_id = int(inserted.ContratoDocenteId)

            planned_hours = horas_planificadas or float(assignment["horas_planificadas"] or 0)
            cursor.execute(
                """
                INSERT INTO rrhh.ContratoDocenteClase
                (
                    ContratoDocenteId, CodigoCarrera, NombreCarrera,
                    CodigoMateria, NombreMateria, CodigoPeriodo, Paralelo, Jornada,
                    HorasPlanificadas, HorasEjecutadas, ValorHora,
                    EstadoClaseCodigo, Observacion
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 0, ?, 'PLANIFICADA', NULLIF(?, N''))
                """,
                contract_id,
                assignment["codigo_carrera"],
                assignment["nombre_carrera"],
                assignment["codigo_materia"],
                assignment["nombre_materia"],
                assignment["codigo_periodo"],
                assignment["paralelo"],
                assignment["jornada"],
                planned_hours,
                valor_hora,
                _clean(observacion),
            )

            relative_path = _teacher_contract_document_relative_path(
                identity["cedula"],
                contract_id,
                "ORIGINAL",
                original_name,
            )
            created_path = _teacher_contract_document_path(relative_path)
            created_path.parent.mkdir(parents=True, exist_ok=True)
            created_path.write_bytes(pdf_bytes)
            cursor.execute(
                """
                UPDATE rrhh.ContratoDocenteDocumento
                SET EsVigente = 0
                WHERE ContratoDocenteId = ? AND TipoDocumento = 'ORIGINAL' AND EsVigente = 1;

                INSERT INTO rrhh.ContratoDocenteDocumento
                (
                    ContratoDocenteId, TipoDocumento, ModalidadAcademica,
                    NombreArchivo, RutaInterna, MimeType, TamanoBytes,
                    HashSha256, EsVigente, UsuarioCarga
                )
                VALUES (?, 'ORIGINAL', ?, ?, ?, 'application/pdf', ?, ?, 1, ?);
                """,
                contract_id,
                contract_id,
                modalidad_academica,
                original_name,
                str(relative_path).replace("\\", "/"),
                len(pdf_bytes),
                sha256(pdf_bytes).hexdigest(),
                _clean(current_user.login),
            )
            conn.commit()
        return {
            "ok": True,
            "contrato_id": contract_id,
            "message": "Contrato adjuntado y vinculado al curso del docente.",
        }
    except HTTPException:
        if created_path and created_path.exists():
            created_path.unlink(missing_ok=True)
        raise
    except (OSError, pyodbc.Error, RuntimeError) as exc:
        if created_path and created_path.exists():
            created_path.unlink(missing_ok=True)
        logger.exception("No se pudo almacenar el contrato docente")
        raise HTTPException(status_code=503, detail="No se pudo guardar el contrato docente") from exc


@router.get("/teacher/contracts/{contract_id}/document")
def teacher_contract_document(
    contract_id: int,
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    version: Annotated[Literal["current", "original", "signed"], Query()] = "current",
    download: Annotated[bool, Query()] = False,
) -> StreamingResponse:
    identity = _teacher_contract_identity(current_user)
    try:
        with get_finance_connection() as conn:
            cursor = conn.cursor()
            _ensure_teacher_contract_document_schema(cursor)
            teacher_id = _sync_finance_teacher(cursor, identity)
            _teacher_contract_owned(cursor, contract_id=contract_id, teacher_id=teacher_id)
            document = _teacher_contract_document(cursor, contract_id=contract_id, version=version)
            conn.commit()
        path = _teacher_contract_document_path(_clean(document.ruta_interna))
        if not path.is_file():
            raise HTTPException(status_code=404, detail="El archivo físico del contrato no está disponible")
        content = path.read_bytes()
    except HTTPException:
        raise
    except (OSError, pyodbc.Error, RuntimeError) as exc:
        logger.exception("No se pudo abrir el documento contractual")
        raise HTTPException(status_code=503, detail="No se pudo abrir el documento contractual") from exc

    disposition = "attachment" if download else "inline"
    filename = f"contrato-{contract_id}-{_safe_filename(document.tipo_documento)}.pdf"
    return StreamingResponse(
        BytesIO(content),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'{disposition}; filename="{filename}"',
            "Cache-Control": "no-store, private",
            "X-Content-Type-Options": "nosniff",
        },
    )


@router.post("/teacher/contracts/{contract_id}/sign")
async def teacher_sign_contract_document(
    contract_id: int,
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    certificado: Annotated[UploadFile, File()],
    contrasena_certificado: Annotated[str, Form(min_length=1, max_length=256)],
    firma_motivo: Annotated[str, Form(max_length=200)] = "Aceptación y firma de contrato docente",
    firma_ubicacion: Annotated[str, Form(max_length=120)] = "Quito, Ecuador",
    firma_contacto: Annotated[str, Form(max_length=200)] = "",
) -> StreamingResponse:
    certificate_name = _clean(certificado.filename).lower()
    if not certificate_name.endswith((".p12", ".pfx")):
        raise HTTPException(status_code=400, detail="Seleccione un certificado con extensión .p12 o .pfx")
    certificate_bytes = await certificado.read()
    if not certificate_bytes:
        raise HTTPException(status_code=400, detail="El archivo de certificado está vacío")
    if len(certificate_bytes) > _TEACHER_CONTRACT_CERTIFICATE_MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="El archivo .p12 debe pesar máximo 2 MB")

    identity = _teacher_contract_identity(current_user)
    created_path: Path | None = None
    try:
        with get_finance_connection() as conn:
            cursor = conn.cursor()
            _ensure_teacher_contract_document_schema(cursor)
            teacher_id = _sync_finance_teacher(cursor, identity)
            contract = _teacher_contract_owned(cursor, contract_id=contract_id, teacher_id=teacher_id)
            cursor.execute(
                """
                SELECT TOP (1) 1
                FROM rrhh.ContratoDocenteDocumento
                WHERE ContratoDocenteId = ? AND TipoDocumento = 'FIRMADO' AND EsVigente = 1
                """,
                contract_id,
            )
            if cursor.fetchone():
                raise HTTPException(
                    status_code=409,
                    detail="El contrato ya tiene una firma docente vigente",
                )
            source = _teacher_contract_document(cursor, contract_id=contract_id, version="original")
            conn.commit()

        source_path = _teacher_contract_document_path(_clean(source.ruta_interna))
        if not source_path.is_file():
            raise HTTPException(status_code=404, detail="No se encontró el PDF original del contrato")
        source_pdf = source_path.read_bytes()
        signed_pdf = await _sign_pdf_with_pkcs12(
            pdf_bytes=source_pdf,
            pkcs12_bytes=certificate_bytes,
            password=contrasena_certificado,
            current_user=current_user,
            reason=firma_motivo,
            location=firma_ubicacion,
            contact=firma_contacto,
            signature_box=(335, 42, 565, 112),
            field_name=f"FirmaDocenteContrato_{contract_id}",
            readable_field_name="Firma electrónica del docente contratista",
        )
        signed_name = f"{Path(_clean(source.nombre_archivo)).stem}-firmado-docente.pdf"
        relative_path = _teacher_contract_document_relative_path(
            identity["cedula"],
            contract_id,
            "FIRMADO",
            signed_name,
        )
        created_path = _teacher_contract_document_path(relative_path)
        created_path.parent.mkdir(parents=True, exist_ok=True)
        created_path.write_bytes(signed_pdf)

        with get_finance_connection() as conn:
            cursor = conn.cursor()
            _ensure_teacher_contract_document_schema(cursor)
            teacher_id = _sync_finance_teacher(cursor, identity)
            _teacher_contract_owned(cursor, contract_id=contract_id, teacher_id=teacher_id)
            cursor.execute(
                """
                SELECT TOP (1) EstadoContratoDocenteId
                FROM cat.EstadoContratoDocente
                WHERE Codigo = 'VIGENTE'
                """
            )
            state_row = cursor.fetchone()
            if not state_row:
                raise HTTPException(status_code=500, detail="No existe el estado contractual VIGENTE")
            cursor.execute(
                """
                UPDATE rrhh.ContratoDocenteDocumento
                SET EsVigente = 0
                WHERE ContratoDocenteId = ? AND TipoDocumento = 'FIRMADO' AND EsVigente = 1;

                INSERT INTO rrhh.ContratoDocenteDocumento
                (
                    ContratoDocenteId, TipoDocumento, ModalidadAcademica,
                    NombreArchivo, RutaInterna, MimeType, TamanoBytes,
                    HashSha256, EsVigente, UsuarioCarga, FirmanteDocumento,
                    FechaFirma, FirmaMotivo
                )
                VALUES (?, 'FIRMADO', ?, ?, ?, 'application/pdf', ?, ?, 1, ?, ?, SYSDATETIME(), ?);

                UPDATE rrhh.ContratoDocente
                SET RutaContratoFirmado = ?, EstadoContratoDocenteId = ?
                WHERE ContratoDocenteId = ? AND DocenteId = ?;
                """,
                contract_id,
                contract_id,
                _clean(source.modalidad_academica),
                signed_name[:260],
                str(relative_path).replace("\\", "/"),
                len(signed_pdf),
                sha256(signed_pdf).hexdigest(),
                _clean(current_user.login),
                identity["nombre"][:300],
                _clean(firma_motivo)[:500],
                str(relative_path).replace("\\", "/"),
                int(state_row.EstadoContratoDocenteId),
                contract_id,
                teacher_id,
            )
            conn.commit()
    except HTTPException:
        if created_path and created_path.exists():
            created_path.unlink(missing_ok=True)
        raise
    except (OSError, pyodbc.Error, RuntimeError) as exc:
        if created_path and created_path.exists():
            created_path.unlink(missing_ok=True)
        logger.exception("No se pudo firmar y registrar el contrato docente")
        raise HTTPException(status_code=503, detail="No se pudo registrar la firma del contrato") from exc

    filename = f"contrato-{contract_id}-firmado-docente.pdf"
    return StreamingResponse(
        BytesIO(signed_pdf),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Cache-Control": "no-store, private",
        },
    )


@router.get("/teacher/courses")
def teacher_courses(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
) -> dict[str, Any]:
    codigo_doc = _teacher_code(current_user)
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                """
                SELECT TOP (1000)
                    TRY_CONVERT(varchar(50), cxd.codigo_doc) AS codigo_doc,
                    TRY_CONVERT(varchar(50), cxd.cod_Anio_Basica) AS cod_anio_basica,
                    TRY_CONVERT(nvarchar(4000), c.Nombre_Basica) AS nombre_carrera,
                    TRY_CONVERT(varchar(50), cxd.codigo_materia) AS codigo_materia,
                    COALESCE(
                        NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), p.cod_materia))), N''),
                        TRY_CONVERT(nvarchar(100), p.codigo_materia),
                        TRY_CONVERT(nvarchar(100), cxd.codigo_materia)
                    ) AS cod_materia,
                    TRY_CONVERT(nvarchar(4000), p.Nomb_Materia) AS nombre_materia,
                    TRY_CONVERT(varchar(50), cxd.codigo_periodo) AS codigo_periodo,
                    TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo) AS detalle_periodo,
                    TRY_CONVERT(nvarchar(100), pe.TipoMatricula) AS tipo_periodo,
                    TRY_CONVERT(date, pe.fechain) AS fecha_inicio,
                    TRY_CONVERT(date, pe.fechafin) AS fecha_fin,
                    COALESCE(TRY_CONVERT(int, pe.Orden), TRY_CONVERT(int, pe.cod_periodo)) AS periodo_orden,
                    TRY_CONVERT(nvarchar(50), cxd.Paralelo) AS paralelo,
                    TRY_CONVERT(int, cxd.Cod_Jornada) AS cod_jornada,
                    TRY_CONVERT(nvarchar(255), j.DetalleJ) AS jornada,
                    TRY_CONVERT(int, p.Semestre) AS semestre,
                    TRY_CONVERT(nvarchar(255), p.Unidad_Organiza) AS unidad_curricular,
                    TRY_CONVERT(int, cxd.estadoMoodleDoc) AS estado_moodle_doc,
                    stats.total_estudiantes
                FROM dbo.CARRERAXDOCENTE cxd
                LEFT JOIN dbo.CARRERAS c
                  ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                LEFT JOIN dbo.PERIODO pe
                  ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cxd.codigo_periodo)
                LEFT JOIN dbo.PENSUM p
                  ON TRY_CONVERT(int, p.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                 AND TRY_CONVERT(int, p.codigo_materia) = TRY_CONVERT(int, cxd.codigo_materia)
                LEFT JOIN dbo.JORNADA j
                  ON TRY_CONVERT(int, j.NumJ) = TRY_CONVERT(int, cxd.Cod_Jornada)
                OUTER APPLY (
                    SELECT COUNT(DISTINCT TRY_CONVERT(int, cxe.codigo_estud)) AS total_estudiantes
                    FROM dbo.CARRERAXESTUD cxe
                    INNER JOIN dbo.DATOS_ESTUD de_active
                      ON TRY_CONVERT(int, de_active.codigo_estud) = TRY_CONVERT(int, cxe.codigo_estud)
                    LEFT JOIN dbo.PENSUM pxe
                      ON TRY_CONVERT(int, pxe.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
                     AND TRY_CONVERT(int, pxe.codigo_materia) = TRY_CONVERT(int, cxe.codigo_materia)
                    WHERE TRY_CONVERT(int, cxe.codigo_periodo) = TRY_CONVERT(int, cxd.codigo_periodo)
                      AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), de_active.Estado))))
                          IN (N'A', N'ACTIVO', N'ACTIVA')
                      AND TRY_CONVERT(int, cxe.cod_anio_Basica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                      AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxe.paralelo)))) =
                          UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxd.Paralelo))))
                      AND UPPER(LTRIM(RTRIM(COALESCE(
                            NULLIF(TRY_CONVERT(nvarchar(100), pxe.cod_materia), N''),
                            TRY_CONVERT(nvarchar(100), pxe.codigo_materia),
                            TRY_CONVERT(nvarchar(100), cxe.codigo_materia),
                            N''
                      )))) = UPPER(LTRIM(RTRIM(COALESCE(
                            NULLIF(TRY_CONVERT(nvarchar(100), p.cod_materia), N''),
                            TRY_CONVERT(nvarchar(100), p.codigo_materia),
                            TRY_CONVERT(nvarchar(100), cxd.codigo_materia),
                            N''
                      ))))
                ) stats
                WHERE TRY_CONVERT(int, cxd.codigo_doc) = ?
                ORDER BY
                    COALESCE(TRY_CONVERT(int, pe.Orden), TRY_CONVERT(int, pe.cod_periodo)) DESC,
                    TRY_CONVERT(nvarchar(4000), c.Nombre_Basica),
                    TRY_CONVERT(nvarchar(4000), p.Nomb_Materia),
                    TRY_CONVERT(nvarchar(50), cxd.Paralelo)
                """,
                codigo_doc,
            )
            items = _group_teacher_courses([_course_item(row) for row in cursor.fetchall()])
        return {"total": len(items), "items": items}
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error consultando materias del docente: {exc}") from exc


def _teacher_student_item(row: Any) -> dict[str, Any]:
    item = _record_item(row)
    item.update(
        {
            "cedula": _clean(row.cedula),
            "nombre_estudiante": _clean(row.nombre_estudiante),
            "correo_personal": _clean(row.correo_personal),
            "correo_intec": _clean(row.correo_intec),
        }
    )
    return item


def _teacher_course_students_for_report(
    current_user: SessionUser,
    period_codes: list[int],
    subject_filter: str,
    parallel: str,
    cod_anio_basica: int | None = None,
    cod_jornada: int | None = None,
) -> list[dict[str, Any]]:
    students_by_key: dict[str, dict[str, Any]] = {}
    if parallel in {"*", "TODOS", "VARIOS"}:
        catalog = teacher_courses(current_user=current_user)
        subject = next(
            (
                item
                for item in catalog.get("items") or []
                if subject_filter
                in {
                    _clean(item.get("cod_materia") or item.get("codigo_materia")).upper(),
                    *{
                        _clean(code).upper()
                        for code in item.get("codigo_materias") or []
                        if _clean(code)
                    },
                }
            ),
            None,
        )
        if not subject:
            return []

        selected_periods = set(period_codes)
        scopes = subject.get("asignaciones") or [subject]
        for kind in ("R", "H"):
            available_codes = {
                period_code
                for scope in scopes
                if ("H" if scope.get("es_homologacion") else "R") == kind
                for period_code in (
                    _int(value)
                    for value in scope.get("codigo_periodos") or [scope.get("codigo_periodo")]
                )
                if period_code is not None and period_code in selected_periods
            }
            chunk_size = 2 if kind == "R" else 1
            ordered_codes = sorted(available_codes, reverse=True)
            for index in range(0, len(ordered_codes), chunk_size):
                payload = teacher_subject_students(
                    current_user=current_user,
                    codigo_materia=subject_filter,
                    tipo_periodo=kind,
                    codigo_periodo=ordered_codes[index:index + chunk_size],
                )
                for item in payload.get("items") or []:
                    key = "|".join(
                        [
                            _clean(item.get("codigo_estud")),
                            _clean(item.get("codigo_periodo")),
                            _clean(item.get("cod_anio_basica")),
                            _clean(item.get("codigo_materia")),
                            _clean(item.get("paralelo")),
                            _clean(item.get("num_matricula")),
                            _clean(item.get("num_grupo")),
                        ]
                    )
                    students_by_key[key] = item
        return list(students_by_key.values())

    chunks = [period_codes[index:index + 2] for index in range(0, len(period_codes), 2)]
    for chunk in chunks:
        payload = teacher_course_students(
            current_user=current_user,
            codigo_periodo=chunk,
            codigo_materia=subject_filter,
            paralelo=parallel,
            cod_anio_basica=cod_anio_basica,
            cod_jornada=cod_jornada,
        )
        for item in payload.get("items") or []:
            key = "|".join(
                [
                    _clean(item.get("codigo_estud")),
                    _clean(item.get("codigo_periodo")),
                    _clean(item.get("cod_anio_basica")),
                    _clean(item.get("codigo_materia")),
                    _clean(item.get("paralelo")),
                    _clean(item.get("num_matricula")),
                    _clean(item.get("num_grupo")),
                ]
            )
            students_by_key[key] = item
    return list(students_by_key.values())


@router.get("/teacher/course-students")
def teacher_course_students(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    codigo_periodo: Annotated[list[int], Query()],
    codigo_materia: Annotated[str, Query()],
    paralelo: Annotated[str, Query(min_length=1)],
    cod_anio_basica: Annotated[int | None, Query()] = None,
    cod_jornada: Annotated[int | None, Query()] = None,
    buscar: Annotated[str | None, Query(max_length=160)] = None,
) -> dict[str, Any]:
    codigo_doc = _teacher_code(current_user)
    parallel = paralelo.strip().upper()
    subject_filter = _clean(codigo_materia).upper()
    search_term = _clean(buscar)
    search_like = f"%{search_term}%"
    period_codes = list(dict.fromkeys(codigo_periodo))
    if not period_codes:
        raise HTTPException(status_code=400, detail="Debe seleccionar al menos un periodo")
    if not subject_filter:
        raise HTTPException(status_code=400, detail="Debe seleccionar una materia")
    if cod_anio_basica is None:
        raise HTTPException(status_code=400, detail="Debe seleccionar la carrera del curso")
    if len(period_codes) > 2:
        raise HTTPException(status_code=400, detail="Solo se pueden consultar hasta 2 periodos regulares unidos")
    period_placeholders = ", ".join("?" for _ in period_codes)
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                f"""
                WITH teacher_assignment AS (
                    SELECT DISTINCT
                        cxd.codigo_periodo,
                        cxd.Paralelo,
                        cxd.Cod_Jornada,
                        COALESCE(
                            NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), pta.cod_materia))), N''),
                            TRY_CONVERT(nvarchar(100), pta.codigo_materia),
                            TRY_CONVERT(nvarchar(100), cxd.codigo_materia)
                        ) AS common_subject_code
                    FROM dbo.CARRERAXDOCENTE cxd
                    LEFT JOIN dbo.PENSUM pta
                      ON TRY_CONVERT(int, pta.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                     AND TRY_CONVERT(int, pta.codigo_materia) = TRY_CONVERT(int, cxd.codigo_materia)
                    WHERE TRY_CONVERT(int, cxd.codigo_doc) = ?
                      AND (? IS NULL OR TRY_CONVERT(int, cxd.cod_Anio_Basica) = ?)
                      AND (? IS NULL OR TRY_CONVERT(int, cxd.Cod_Jornada) = ?)
                      AND (
                            TRY_CONVERT(nvarchar(100), cxd.codigo_materia) = ?
                            OR UPPER(LTRIM(RTRIM(COALESCE(
                                NULLIF(TRY_CONVERT(nvarchar(100), pta.cod_materia), N''),
                                TRY_CONVERT(nvarchar(100), pta.codigo_materia),
                                TRY_CONVERT(nvarchar(100), cxd.codigo_materia),
                                N''
                            )))) = ?
                      )
                      AND TRY_CONVERT(int, cxd.codigo_periodo) IN ({period_placeholders})
                      AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxd.Paralelo)))) = ?
                )
                SELECT TOP (1000)
                    TRY_CONVERT(varchar(50), de.codigo_estud) AS codigo_estud,
                    TRY_CONVERT(nvarchar(100), de.Cedula_Est) AS cedula,
                    TRY_CONVERT(nvarchar(4000), de.Apellidos_nombre) AS nombre_estudiante,
                    TRY_CONVERT(nvarchar(255), de.correo) AS correo_personal,
                    COALESCE(
                        NULLIF(TRY_CONVERT(nvarchar(255), ce.CorreoIntec), N''),
                        TRY_CONVERT(nvarchar(255), de.correointec)
                    ) AS correo_intec,
                    TRY_CONVERT(varchar(50), cxe.cod_anio_Basica) AS cod_anio_basica,
                    TRY_CONVERT(nvarchar(4000), c.Nombre_Basica) AS nombre_carrera,
                    TRY_CONVERT(varchar(50), cxe.codigo_periodo) AS codigo_periodo,
                    TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo) AS detalle_periodo,
                    TRY_CONVERT(int, pe.anio) AS anio_periodo,
                    TRY_CONVERT(varchar(50), cxe.codigo_materia) AS codigo_materia,
                    TRY_CONVERT(varchar(100), p.cod_materia) AS cod_materia,
                    TRY_CONVERT(nvarchar(4000), p.Nomb_Materia) AS nombre_materia,
                    TRY_CONVERT(int, p.Semestre) AS semestre,
                    TRY_CONVERT(int, p.NumMalla) AS num_malla,
                    TRY_CONVERT(float, p.Horas) AS horas,
                    TRY_CONVERT(int, p.Orden) AS orden,
                    TRY_CONVERT(float, COALESCE(NULLIF(cxe.Num_Creditos, 0), p.Creditos)) AS creditos,
                    TRY_CONVERT(nvarchar(50), cxe.paralelo) AS paralelo,
                    TRY_CONVERT(int, cxe.NumGrupo) AS num_grupo,
                    TRY_CONVERT(varchar(50), cxe.Num_Matricula) AS num_matricula,
                    cxe.Fecha_Matricula AS fecha_matricula,
                    TRY_CONVERT(nvarchar(20), cxe.TipoMatricula) AS tipo_matricula,
                    TRY_CONVERT(float, cxe.teoriaHomo) AS teoria_homo,
                    TRY_CONVERT(float, cxe.practicahomo) AS practica_homo,
                    TRY_CONVERT(float, cxe.P1Tareas) AS p1_tareas,
                    TRY_CONVERT(float, cxe.P1Proyectos) AS p1_proyectos,
                    TRY_CONVERT(float, cxe.P1Examen) AS p1_examen,
                    TRY_CONVERT(float, cxe.promP1) AS prom_p1,
                    TRY_CONVERT(float, cxe.P2Tareas) AS p2_tareas,
                    TRY_CONVERT(float, cxe.P2Proyectos) AS p2_proyectos,
                    TRY_CONVERT(float, cxe.P2Examen) AS p2_examen,
                    TRY_CONVERT(float, cxe.promP2) AS prom_p2,
                    TRY_CONVERT(float, cxe.P3Tareas) AS p3_tareas,
                    TRY_CONVERT(float, cxe.P3Proyectos) AS p3_proyectos,
                    TRY_CONVERT(float, cxe.P3Examen) AS p3_examen,
                    TRY_CONVERT(float, cxe.promP3) AS prom_p3,
                    TRY_CONVERT(float, cxe.Promedio) AS promedio,
                    TRY_CONVERT(float, cxe.Asistencia) AS asistencia,
                    TRY_CONVERT(float, cxe.Recuperacion) AS recuperacion,
                    TRY_CONVERT(float, cxe.PromedioFinal) AS promedio_final_raw,
                    COALESCE(
                        TRY_CONVERT(float, cxe.PromedioFinal),
                        CASE
                            WHEN (
                                    UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(50), cxe.TipoMatricula), N'')))) = N'H'
                                 OR UPPER(COALESCE(TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo), N'')) LIKE N'%HOMO%'
                                 )
                             AND TRY_CONVERT(float, cxe.teoriaHomo) IS NOT NULL
                             AND TRY_CONVERT(float, cxe.practicahomo) IS NOT NULL
                            THEN (TRY_CONVERT(float, cxe.teoriaHomo) * 0.4) + (TRY_CONVERT(float, cxe.practicahomo) * 0.6)
                        END,
                        CASE
                            WHEN TRY_CONVERT(float, cxe.promP1) IS NOT NULL
                             AND TRY_CONVERT(float, cxe.promP2) IS NOT NULL
                             AND TRY_CONVERT(float, cxe.promP3) IS NOT NULL
                            THEN (TRY_CONVERT(float, cxe.promP1) + TRY_CONVERT(float, cxe.promP2) + TRY_CONVERT(float, cxe.promP3)) / 3
                        END,
                        TRY_CONVERT(float, cxe.Promedio),
                        TRY_CONVERT(float, cxe.PromedioAux)
                    ) AS nota_final,
                    COALESCE(TRY_CONVERT(float, pe.NotaAprobar), 7) AS nota_aprobar,
                    CASE
                        WHEN UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(50), cxe.caprueba), N'')))) LIKE N'A%' THEN 1
                        WHEN COALESCE(
                                TRY_CONVERT(float, cxe.PromedioFinal),
                                CASE
                                    WHEN (
                                            UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(50), cxe.TipoMatricula), N'')))) = N'H'
                                         OR UPPER(COALESCE(TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo), N'')) LIKE N'%HOMO%'
                                         )
                                     AND TRY_CONVERT(float, cxe.teoriaHomo) IS NOT NULL
                                     AND TRY_CONVERT(float, cxe.practicahomo) IS NOT NULL
                                    THEN (TRY_CONVERT(float, cxe.teoriaHomo) * 0.4) + (TRY_CONVERT(float, cxe.practicahomo) * 0.6)
                                END,
                                CASE
                                    WHEN TRY_CONVERT(float, cxe.promP1) IS NOT NULL
                                     AND TRY_CONVERT(float, cxe.promP2) IS NOT NULL
                                     AND TRY_CONVERT(float, cxe.promP3) IS NOT NULL
                                    THEN (TRY_CONVERT(float, cxe.promP1) + TRY_CONVERT(float, cxe.promP2) + TRY_CONVERT(float, cxe.promP3)) / 3
                                END,
                                TRY_CONVERT(float, cxe.Promedio),
                                TRY_CONVERT(float, cxe.PromedioAux)
                             ) >= COALESCE(TRY_CONVERT(float, pe.NotaAprobar), 7)
                        THEN 1
                        ELSE 0
                    END AS aprobada,
                    TRY_CONVERT(nvarchar(max), cxe.observaciones) AS observaciones,
                    TRY_CONVERT(nvarchar(255), cxe.seguimiento) AS seguimiento
                FROM dbo.CARRERAXESTUD cxe
                INNER JOIN dbo.DATOS_ESTUD de
                  ON TRY_CONVERT(int, de.codigo_estud) = TRY_CONVERT(int, cxe.codigo_estud)
                LEFT JOIN dbo.CorreosEstudIntec ce
                  ON TRY_CONVERT(int, ce.codestud) = TRY_CONVERT(int, de.codigo_estud)
                LEFT JOIN dbo.CARRERAS c
                  ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
                LEFT JOIN dbo.PERIODO pe
                  ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cxe.codigo_periodo)
                LEFT JOIN dbo.PENSUM p
                  ON TRY_CONVERT(int, p.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
                 AND TRY_CONVERT(int, p.codigo_materia) = TRY_CONVERT(int, cxe.codigo_materia)
                WHERE (? IS NULL OR TRY_CONVERT(int, cxe.cod_anio_Basica) = ?)
                  AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), de.Estado))))
                      IN (N'A', N'ACTIVO', N'ACTIVA')
                  AND (
                        ? = N''
                        OR TRY_CONVERT(nvarchar(4000), de.Apellidos_nombre)
                            COLLATE Latin1_General_100_CI_AI LIKE ?
                        OR TRY_CONVERT(nvarchar(100), de.Cedula_Est)
                            COLLATE Latin1_General_100_CI_AI LIKE ?
                        OR TRY_CONVERT(nvarchar(50), de.codigo_estud)
                            COLLATE Latin1_General_100_CI_AI LIKE ?
                  )
                  AND EXISTS (
                      SELECT 1
                      FROM teacher_assignment ta
                      WHERE TRY_CONVERT(int, ta.codigo_periodo) = TRY_CONVERT(int, cxe.codigo_periodo)
                        AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), ta.Paralelo)))) =
                            UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxe.paralelo))))
                        AND UPPER(LTRIM(RTRIM(COALESCE(ta.common_subject_code, N'')))) =
                            UPPER(LTRIM(RTRIM(COALESCE(
                                NULLIF(TRY_CONVERT(nvarchar(100), p.cod_materia), N''),
                                TRY_CONVERT(nvarchar(100), p.codigo_materia),
                                TRY_CONVERT(nvarchar(100), cxe.codigo_materia),
                                N''
                            ))))
                  )
                ORDER BY
                    TRY_CONVERT(int, cxe.codigo_periodo) DESC,
                    TRY_CONVERT(nvarchar(4000), de.Apellidos_nombre)
                """,
                codigo_doc,
                cod_anio_basica,
                cod_anio_basica,
                cod_jornada,
                cod_jornada,
                subject_filter,
                subject_filter,
                *period_codes,
                parallel,
                cod_anio_basica,
                cod_anio_basica,
                search_term,
                search_like,
                search_like,
                search_like,
            )
            rows = cursor.fetchall()
        return {"total": len(rows), "items": [_teacher_student_item(row) for row in rows]}
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error consultando estudiantes del curso: {exc}") from exc


@router.get("/teacher/subject-students")
def teacher_subject_students(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    codigo_materia: Annotated[str, Query(min_length=1)],
    tipo_periodo: Annotated[Literal["R", "H"], Query()],
    codigo_periodo: Annotated[list[int], Query(min_length=1)],
) -> dict[str, Any]:
    """Return every exact enrolment assigned to a teacher for one common subject code."""
    selected_periods = list(dict.fromkeys(codigo_periodo))
    if not selected_periods:
        raise HTTPException(status_code=400, detail="Debe seleccionar al menos un periodo academico.")
    period_limit = 2 if tipo_periodo == "R" else 1
    if len(selected_periods) > period_limit:
        detail = (
            "Solo se pueden unir hasta dos periodos regulares."
            if tipo_periodo == "R"
            else "La homologacion debe consultarse en un solo periodo independiente."
        )
        raise HTTPException(status_code=400, detail=detail)

    subject_filter = _clean(codigo_materia).upper()
    catalog = teacher_courses(current_user=current_user)
    subject = next(
        (
            item
            for item in catalog.get("items") or []
            if subject_filter
            in {
                _clean(item.get("cod_materia") or item.get("codigo_materia")).upper(),
                *{
                    _clean(code).upper()
                    for code in item.get("codigo_materias") or []
                    if _clean(code)
                },
            }
        ),
        None,
    )
    if not subject:
        raise HTTPException(status_code=404, detail="La materia no esta asignada al docente autenticado")

    scopes = [
        scope
        for scope in subject.get("asignaciones") or [subject]
        if ("H" if scope.get("es_homologacion") else "R") == tipo_periodo
    ]
    if not scopes:
        return {
            "total": 0,
            "items": [],
            "codigo_materia": _clean(subject.get("cod_materia") or subject.get("codigo_materia")),
            "tipo_periodo": tipo_periodo,
            "codigo_periodos": selected_periods,
            "asignaciones_consultadas": 0,
        }

    available_periods = {
        period_code
        for scope in scopes
        for period_code in (
            _int(value)
            for value in scope.get("codigo_periodos") or [scope.get("codigo_periodo")]
        )
        if period_code is not None
    }
    missing_periods = [period_code for period_code in selected_periods if period_code not in available_periods]
    if missing_periods:
        raise HTTPException(
            status_code=404,
            detail="Uno o mas periodos no pertenecen a la materia y tipo seleccionados.",
        )

    students_by_key: dict[str, dict[str, Any]] = {}
    consulted_scopes = 0
    for scope in scopes:
        period_codes = [
            code
            for code in (_int(value) for value in scope.get("codigo_periodos") or [scope.get("codigo_periodo")])
            if code is not None and code in selected_periods
        ]
        career_code = _int(scope.get("cod_anio_basica"))
        parallel = _clean(scope.get("paralelo"))
        if not period_codes or career_code is None or not parallel:
            continue
        consulted_scopes += 1
        payload = teacher_course_students(
            current_user=current_user,
            codigo_periodo=period_codes,
            codigo_materia=_clean(subject.get("cod_materia") or subject.get("codigo_materia")),
            paralelo=parallel,
            cod_anio_basica=career_code,
            cod_jornada=_int(scope.get("cod_jornada")),
        )
        for item in payload.get("items") or []:
            key = "|".join(
                [
                    _clean(item.get("codigo_estud")),
                    _clean(item.get("codigo_periodo")),
                    _clean(item.get("cod_anio_basica")),
                    _clean(item.get("codigo_materia")),
                    _clean(item.get("paralelo")),
                    _clean(item.get("num_matricula")),
                    _clean(item.get("num_grupo")),
                ]
            )
            students_by_key[key] = item

    students = sorted(
        students_by_key.values(),
        key=lambda item: (
            _clean(item.get("nombre_carrera")),
            -(_int(item.get("codigo_periodo")) or 0),
            _clean(item.get("nombre_estudiante")),
        ),
    )
    return {
        "total": len(students),
        "items": students,
        "codigo_materia": _clean(subject.get("cod_materia") or subject.get("codigo_materia")),
        "tipo_periodo": tipo_periodo,
        "codigo_periodos": selected_periods,
        "asignaciones_consultadas": consulted_scopes,
    }


def _admin_grade_completion(item: dict[str, Any]) -> str:
    if item.get("promedio_final") is not None:
        return "COMPLETA"

    if item.get("es_homologacion"):
        grade_fields = ("teoria_homo", "practica_homo")
    else:
        grade_fields = (
            "p1_tareas",
            "p1_proyectos",
            "p1_examen",
            "prom_p1",
            "p2_tareas",
            "p2_proyectos",
            "p2_examen",
            "prom_p2",
            "p3_tareas",
            "p3_proyectos",
            "p3_examen",
            "prom_p3",
            "recuperacion",
        )
    return "EN_PROCESO" if any(item.get(field) is not None for field in grade_fields) else "SIN_CALIFICAR"


def _courses_with_enrolled_students(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [item for item in items if (_int(item.get("total_estudiantes")) or 0) > 0]


def _admin_grade_course_key(item: dict[str, Any] | AdminGradeCourseSelectionPayload) -> tuple[str, str, str, str, str]:
    source = item.model_dump() if isinstance(item, AdminGradeCourseSelectionPayload) else item
    return (
        _clean(source.get("codigo_periodo")),
        _clean(source.get("cod_anio_basica")),
        _clean(source.get("cod_materia") or source.get("codigo_materia")).upper(),
        _clean(source.get("paralelo")).upper(),
        _clean(source.get("cod_jornada")),
    )


def _admin_grade_period_type(item: dict[str, Any]) -> Literal["R", "H"]:
    return "H" if item.get("es_homologacion") or _is_homologation_type(
        item.get("tipo_periodo"),
        item.get("detalle_periodo"),
    ) else "R"


def _resolve_admin_grade_course_selections(
    available_courses: list[dict[str, Any]],
    selections: list[AdminGradeCourseSelectionPayload],
) -> list[dict[str, Any]]:
    if not selections:
        raise HTTPException(status_code=400, detail="Debe seleccionar al menos un periodo")
    if len(selections) > 3:
        raise HTTPException(status_code=400, detail="Solo se pueden seleccionar hasta 3 periodos")

    requested_period_codes = [_clean(selection.codigo_periodo) for selection in selections]
    if len(set(requested_period_codes)) != len(requested_period_codes):
        raise HTTPException(
            status_code=400,
            detail="Cada selección debe corresponder a un periodo diferente",
        )

    requested_keys = [_admin_grade_course_key(selection) for selection in selections]
    if len(set(requested_keys)) != len(requested_keys):
        raise HTTPException(status_code=400, detail="No se puede seleccionar el mismo periodo y curso más de una vez")

    available_by_key = {_admin_grade_course_key(course): course for course in available_courses}
    requested_courses: list[dict[str, Any]] = []
    for key in requested_keys:
        course = available_by_key.get(key)
        if course is None:
            raise HTTPException(
                status_code=400,
                detail="Uno de los periodos no corresponde a una asignación activa del docente con estudiantes",
            )
        requested_courses.append(course)

    subject_codes = {
        _clean(course.get("cod_materia") or course.get("codigo_materia")).upper()
        for course in requested_courses
    }
    if len(subject_codes) != 1:
        raise HTTPException(status_code=400, detail="Los periodos deben pertenecer a una sola asignatura")

    period_types = {_admin_grade_period_type(course) for course in requested_courses}
    if len(period_types) != 1:
        raise HTTPException(
            status_code=400,
            detail="No se pueden mezclar periodos regulares y de homologación en la misma consulta",
        )

    selected_subject = next(iter(subject_codes))
    selected_type = next(iter(period_types))
    selected_courses: list[dict[str, Any]] = []
    selected_keys: set[tuple[str, str, str, str, str]] = set()
    for period_code in requested_period_codes:
        for course in available_courses:
            if _clean(course.get("codigo_periodo")) != period_code:
                continue
            if _clean(course.get("cod_materia") or course.get("codigo_materia")).upper() != selected_subject:
                continue
            if _admin_grade_period_type(course) != selected_type:
                continue
            key = _admin_grade_course_key(course)
            if key in selected_keys:
                continue
            selected_keys.add(key)
            selected_courses.append(course)

    if not selected_courses:
        raise HTTPException(status_code=400, detail="Los periodos seleccionados no tienen cursos disponibles")
    return selected_courses


def _admin_grade_students_response(source_items: list[dict[str, Any]]) -> dict[str, Any]:
    items = []
    summary = {"completas": 0, "en_proceso": 0, "sin_calificar": 0, "aprobados": 0, "reprobados": 0}
    for source_item in source_items:
        item = dict(source_item)
        completion = _admin_grade_completion(item)
        item["estado_registro"] = completion
        if completion == "COMPLETA":
            summary["completas"] += 1
        elif completion == "EN_PROCESO":
            summary["en_proceso"] += 1
        else:
            summary["sin_calificar"] += 1
        if item.get("estado_nota") == "APROBADO":
            summary["aprobados"] += 1
        elif item.get("estado_nota") == "REPROBADO":
            summary["reprobados"] += 1
        items.append(item)
    return {"total": len(items), "items": items, "summary": summary}


@router.get("/admin/grades/teachers")
def admin_grade_teachers(
    current_user: Annotated[SessionUser, Depends(_GRADES_ADMIN_ACCESS)],
    buscar: Annotated[str, Query(max_length=150)] = "",
    estado: Annotated[str, Query(max_length=30)] = "",
    limit: Annotated[int, Query(ge=1, le=2000)] = 1000,
) -> dict[str, Any]:
    del current_user
    cleaned_search = buscar.strip()
    cleaned_state = estado.strip().upper()
    like = f"%{cleaned_search}%"
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                f"""
                SELECT TOP ({limit})
                    TRY_CONVERT(varchar(50), d.codigo_doc) AS codigo_doc,
                    LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), d.cedula_doc))) AS cedula,
                    LTRIM(RTRIM(TRY_CONVERT(nvarchar(4000), d.apellidos_nombre))) AS docente,
                    COALESCE(
                        NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(255), d.correo))), N''),
                        NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(255), d.correop))), N''),
                        NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(255), usuario_docente.login))), N'')
                    ) AS correo,
                    LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), usuario_docente.Estado))) AS estado,
                    COALESCE(
                        NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(255), est.ESTADO))), N''),
                        CASE
                            WHEN usuario_docente.Codigo_Usuario IS NULL THEN N'Sin usuario vinculado'
                            ELSE LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), usuario_docente.Estado)))
                        END
                    ) AS estado_nombre,
                    asignaciones.total_asignaciones,
                    asignaciones.total_asignaturas,
                    asignaciones.total_periodos
                FROM dbo.DATOSDOCENTE d
                OUTER APPLY (
                    SELECT TOP (1)
                        u.Codigo_Usuario,
                        u.login,
                        u.Estado
                    FROM dbo.USUARIOS u
                    WHERE TRY_CONVERT(int, u.Codigo_Usuario) = TRY_CONVERT(int, d.codigo_doc)
                       OR (
                            NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), d.cedula_doc))), N'') IS NOT NULL
                        AND LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), u.cedula))) =
                            LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), d.cedula_doc)))
                       )
                    ORDER BY
                        CASE WHEN COALESCE(TRY_CONVERT(int, u.tipo_usuario), 2) <> 1 THEN 0 ELSE 1 END,
                        CASE WHEN TRY_CONVERT(int, u.Codigo_Usuario) = TRY_CONVERT(int, d.codigo_doc) THEN 0 ELSE 1 END,
                        TRY_CONVERT(int, u.Codigo_Usuario)
                ) usuario_docente
                LEFT JOIN dbo.ESTADO est
                  ON UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), est.IDESTADO)))) =
                     UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), usuario_docente.Estado))))
                OUTER APPLY (
                    SELECT
                        COUNT(*) AS total_asignaciones,
                        COUNT(DISTINCT CONCAT(
                            TRY_CONVERT(nvarchar(50), cxd.cod_Anio_Basica), N'|',
                            TRY_CONVERT(nvarchar(50), cxd.codigo_materia)
                        )) AS total_asignaturas,
                        COUNT(DISTINCT TRY_CONVERT(nvarchar(50), cxd.codigo_periodo)) AS total_periodos
                    FROM dbo.CARRERAXDOCENTE cxd
                    WHERE TRY_CONVERT(int, cxd.codigo_doc) = TRY_CONVERT(int, d.codigo_doc)
                ) asignaciones
                WHERE asignaciones.total_asignaciones > 0
                  AND (
                        ? = N''
                     OR TRY_CONVERT(nvarchar(50), d.codigo_doc) LIKE ?
                     OR TRY_CONVERT(nvarchar(100), d.cedula_doc) LIKE ?
                     OR TRY_CONVERT(nvarchar(4000), d.apellidos_nombre) LIKE ?
                     OR TRY_CONVERT(nvarchar(255), d.correo) LIKE ?
                     OR TRY_CONVERT(nvarchar(255), d.correop) LIKE ?
                     OR TRY_CONVERT(nvarchar(255), usuario_docente.login) LIKE ?
                  )
                  AND (
                        ? = N''
                     OR (? = N'SIN_USUARIO' AND usuario_docente.Codigo_Usuario IS NULL)
                     OR UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), usuario_docente.Estado)))) = ?
                  )
                ORDER BY LTRIM(RTRIM(TRY_CONVERT(nvarchar(4000), d.apellidos_nombre)))
                """,
                cleaned_search,
                like,
                like,
                like,
                like,
                like,
                like,
                cleaned_state,
                cleaned_state,
                cleaned_state,
            )
            rows = cursor.fetchall()
        items = [
            {
                "codigo_doc": _clean(row.codigo_doc),
                "cedula": _clean(row.cedula),
                "docente": _clean(row.docente),
                "correo": _clean(row.correo),
                "estado": _clean(row.estado),
                "estado_nombre": _clean(row.estado_nombre),
                "total_asignaciones": _int(row.total_asignaciones) or 0,
                "total_asignaturas": _int(row.total_asignaturas) or 0,
                "total_periodos": _int(row.total_periodos) or 0,
            }
            for row in rows
        ]
        return {"total": len(items), "items": items}
    except pyodbc.Error as exc:
        logger.exception("No se pudieron consultar los docentes para calificaciones")
        raise HTTPException(status_code=500, detail="No se pudo consultar el listado de docentes") from exc


@router.get("/admin/grades/teachers/{codigo_doc}/courses")
def admin_grade_teacher_courses(
    codigo_doc: int,
    current_user: Annotated[SessionUser, Depends(_GRADES_ADMIN_ACCESS)],
) -> dict[str, Any]:
    del current_user
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                """
                SELECT TOP (3000)
                    TRY_CONVERT(varchar(50), cxd.codigo_doc) AS codigo_doc,
                    TRY_CONVERT(varchar(50), cxd.cod_Anio_Basica) AS cod_anio_basica,
                    TRY_CONVERT(nvarchar(4000), c.Nombre_Basica) AS nombre_carrera,
                    TRY_CONVERT(varchar(50), cxd.codigo_materia) AS codigo_materia,
                    COALESCE(
                        NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), p.cod_materia))), N''),
                        TRY_CONVERT(nvarchar(100), p.codigo_materia),
                        TRY_CONVERT(nvarchar(100), cxd.codigo_materia)
                    ) AS cod_materia,
                    TRY_CONVERT(nvarchar(4000), p.Nomb_Materia) AS nombre_materia,
                    TRY_CONVERT(varchar(50), cxd.codigo_periodo) AS codigo_periodo,
                    TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo) AS detalle_periodo,
                    TRY_CONVERT(nvarchar(100), pe.TipoMatricula) AS tipo_periodo,
                    TRY_CONVERT(date, pe.fechain) AS fecha_inicio,
                    TRY_CONVERT(date, pe.fechafin) AS fecha_fin,
                    COALESCE(TRY_CONVERT(int, pe.Orden), TRY_CONVERT(int, pe.cod_periodo)) AS periodo_orden,
                    TRY_CONVERT(nvarchar(50), cxd.Paralelo) AS paralelo,
                    TRY_CONVERT(int, cxd.Cod_Jornada) AS cod_jornada,
                    TRY_CONVERT(nvarchar(255), j.DetalleJ) AS jornada,
                    TRY_CONVERT(int, p.Semestre) AS semestre,
                    TRY_CONVERT(nvarchar(255), p.Unidad_Organiza) AS unidad_curricular,
                    TRY_CONVERT(int, cxd.estadoMoodleDoc) AS estado_moodle_doc,
                    stats.total_estudiantes
                FROM dbo.CARRERAXDOCENTE cxd
                LEFT JOIN dbo.CARRERAS c
                  ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                LEFT JOIN dbo.PERIODO pe
                  ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cxd.codigo_periodo)
                LEFT JOIN dbo.PENSUM p
                  ON TRY_CONVERT(int, p.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                 AND TRY_CONVERT(int, p.codigo_materia) = TRY_CONVERT(int, cxd.codigo_materia)
                LEFT JOIN dbo.JORNADA j
                  ON TRY_CONVERT(int, j.NumJ) = TRY_CONVERT(int, cxd.Cod_Jornada)
                OUTER APPLY (
                    SELECT COUNT(DISTINCT TRY_CONVERT(int, cxe.codigo_estud)) AS total_estudiantes
                    FROM dbo.CARRERAXESTUD cxe
                    INNER JOIN dbo.DATOS_ESTUD de_active
                      ON TRY_CONVERT(int, de_active.codigo_estud) = TRY_CONVERT(int, cxe.codigo_estud)
                    LEFT JOIN dbo.PENSUM pxe
                      ON TRY_CONVERT(int, pxe.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
                     AND TRY_CONVERT(int, pxe.codigo_materia) = TRY_CONVERT(int, cxe.codigo_materia)
                    WHERE TRY_CONVERT(int, cxe.codigo_periodo) = TRY_CONVERT(int, cxd.codigo_periodo)
                      AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), de_active.Estado))))
                          IN (N'A', N'ACTIVO', N'ACTIVA')
                      AND TRY_CONVERT(int, cxe.cod_anio_Basica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                      AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxe.paralelo)))) =
                          UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxd.Paralelo))))
                      AND UPPER(LTRIM(RTRIM(COALESCE(
                            NULLIF(TRY_CONVERT(nvarchar(100), pxe.cod_materia), N''),
                            TRY_CONVERT(nvarchar(100), pxe.codigo_materia),
                            TRY_CONVERT(nvarchar(100), cxe.codigo_materia),
                            N''
                      )))) = UPPER(LTRIM(RTRIM(COALESCE(
                            NULLIF(TRY_CONVERT(nvarchar(100), p.cod_materia), N''),
                            TRY_CONVERT(nvarchar(100), p.codigo_materia),
                            TRY_CONVERT(nvarchar(100), cxd.codigo_materia),
                            N''
                      ))))
                ) stats
                WHERE TRY_CONVERT(int, cxd.codigo_doc) = ?
                ORDER BY
                    COALESCE(TRY_CONVERT(int, pe.Orden), TRY_CONVERT(int, pe.cod_periodo)) DESC,
                    TRY_CONVERT(nvarchar(4000), c.Nombre_Basica),
                    TRY_CONVERT(nvarchar(4000), p.Nomb_Materia),
                    TRY_CONVERT(nvarchar(50), cxd.Paralelo)
                """,
                codigo_doc,
            )
            raw_items = [_course_item(row) for row in cursor.fetchall()]

        courses_by_key: dict[tuple[str, str, str, str, str], dict[str, Any]] = {}
        for item in raw_items:
            key = (
                _clean(item.get("cod_anio_basica")),
                _clean(item.get("cod_materia") or item.get("codigo_materia")).upper(),
                _clean(item.get("codigo_periodo")),
                _clean(item.get("paralelo")).upper(),
                _clean(item.get("cod_jornada")),
            )
            current = courses_by_key.get(key)
            if current is None or (_int(item.get("total_estudiantes")) or 0) > (_int(current.get("total_estudiantes")) or 0):
                courses_by_key[key] = item
        items = _courses_with_enrolled_students(list(courses_by_key.values()))
        return {"total": len(items), "items": items}
    except pyodbc.Error as exc:
        logger.exception("No se pudieron consultar las asignaturas del docente %s", codigo_doc)
        raise HTTPException(status_code=500, detail="No se pudieron consultar las asignaturas del docente") from exc


@router.get("/admin/grades/teachers/{codigo_doc}/students")
def admin_grade_teacher_students(
    codigo_doc: int,
    current_user: Annotated[SessionUser, Depends(_GRADES_ADMIN_ACCESS)],
    codigo_periodo: Annotated[list[int], Query()],
    codigo_materia: Annotated[str, Query()],
    paralelo: Annotated[str, Query(min_length=1)],
    cod_anio_basica: Annotated[int | None, Query()] = None,
    cod_jornada: Annotated[int | None, Query()] = None,
) -> dict[str, Any]:
    teacher_session = current_user.model_copy(update={"codigo_doc": codigo_doc})
    payload = teacher_course_students(
        current_user=teacher_session,
        codigo_periodo=codigo_periodo,
        codigo_materia=codigo_materia,
        paralelo=paralelo,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
    )
    return _admin_grade_students_response(payload.get("items") or [])


@router.post("/admin/grades/teachers/{codigo_doc}/students")
def admin_grade_teacher_students_batch(
    codigo_doc: int,
    payload: AdminGradeCourseBatchPayload,
    current_user: Annotated[SessionUser, Depends(_GRADES_ADMIN_ACCESS)],
) -> dict[str, Any]:
    available_payload = admin_grade_teacher_courses(codigo_doc=codigo_doc, current_user=current_user)
    selected_courses = _resolve_admin_grade_course_selections(
        available_payload.get("items") or [],
        payload.courses,
    )
    teacher_session = current_user.model_copy(update={"codigo_doc": codigo_doc})
    source_items: list[dict[str, Any]] = []
    for course in selected_courses:
        period_code = _int(course.get("codigo_periodo"))
        career_code = _int(course.get("cod_anio_basica"))
        if period_code is None or career_code is None:
            raise HTTPException(status_code=400, detail="La asignación seleccionada no tiene periodo o carrera válidos")
        course_payload = teacher_course_students(
            current_user=teacher_session,
            codigo_periodo=[period_code],
            codigo_materia=_clean(course.get("cod_materia") or course.get("codigo_materia")),
            paralelo=_clean(course.get("paralelo")),
            cod_anio_basica=career_code,
            cod_jornada=_int(course.get("cod_jornada")),
        )
        source_items.extend(course_payload.get("items") or [])

    response = _admin_grade_students_response(source_items)
    response["periodos_seleccionados"] = len({
        _clean(course.get("codigo_periodo")) for course in selected_courses
    })
    response["tipo_seleccion"] = _admin_grade_period_type(selected_courses[0])
    return response


def _legacy_grade_text(value: Any) -> str:
    number = _number(value)
    if number is None:
        return "-"
    return f"{number:.2f}".replace(".", ",")


def _teacher_course_report_meta(
    codigo_doc: int,
    period_codes: list[int],
    subject_filter: str,
    parallel: str,
    cod_anio_basica: int | None,
) -> dict[str, Any]:
    if not period_codes:
        return {}
    parallel_filter = None if parallel in {"*", "TODOS", "VARIOS"} else parallel
    period_placeholders = ", ".join("?" for _ in period_codes)
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                f"""
                SELECT TOP (50)
                    TRY_CONVERT(nvarchar(4000), c.Nombre_Basica) AS nombre_carrera,
                    TRY_CONVERT(varchar(50), cxd.codigo_materia) AS codigo_materia,
                    COALESCE(
                        NULLIF(LTRIM(RTRIM(TRY_CONVERT(nvarchar(100), p.cod_materia))), N''),
                        TRY_CONVERT(nvarchar(100), p.codigo_materia),
                        TRY_CONVERT(nvarchar(100), cxd.codigo_materia)
                    ) AS cod_materia,
                    TRY_CONVERT(nvarchar(4000), p.Nomb_Materia) AS nombre_materia,
                    TRY_CONVERT(varchar(50), cxd.codigo_periodo) AS codigo_periodo,
                    TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo) AS detalle_periodo,
                    TRY_CONVERT(nvarchar(100), pe.TipoMatricula) AS tipo_periodo,
                    TRY_CONVERT(nvarchar(50), cxd.Paralelo) AS paralelo,
                    TRY_CONVERT(int, cxd.Cod_Jornada) AS cod_jornada,
                    TRY_CONVERT(nvarchar(255), j.DetalleJ) AS jornada,
                    TRY_CONVERT(int, p.Semestre) AS semestre,
                    TRY_CONVERT(nvarchar(255), p.Unidad_Organiza) AS unidad_curricular,
                    TRY_CONVERT(float, p.Horas) AS horas
                FROM dbo.CARRERAXDOCENTE cxd
                LEFT JOIN dbo.CARRERAS c
                  ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                LEFT JOIN dbo.PERIODO pe
                  ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cxd.codigo_periodo)
                LEFT JOIN dbo.PENSUM p
                  ON TRY_CONVERT(int, p.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                 AND TRY_CONVERT(int, p.codigo_materia) = TRY_CONVERT(int, cxd.codigo_materia)
                LEFT JOIN dbo.JORNADA j
                  ON TRY_CONVERT(int, j.NumJ) = TRY_CONVERT(int, cxd.Cod_Jornada)
                WHERE TRY_CONVERT(int, cxd.codigo_doc) = ?
                  AND (? IS NULL OR TRY_CONVERT(int, cxd.cod_Anio_Basica) = ?)
                  AND (
                        TRY_CONVERT(nvarchar(100), cxd.codigo_materia) = ?
                        OR UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(100), p.cod_materia), N'')))) = ?
                  )
                  AND TRY_CONVERT(int, cxd.codigo_periodo) IN ({period_placeholders})
                  AND (
                        ? IS NULL
                        OR UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxd.Paralelo)))) = ?
                  )
                ORDER BY TRY_CONVERT(int, cxd.codigo_periodo) DESC
                """,
                codigo_doc,
                cod_anio_basica,
                cod_anio_basica,
                subject_filter,
                subject_filter,
                *period_codes,
                parallel_filter,
                parallel_filter,
            )
            rows = cursor.fetchall()
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error consultando datos del reporte docente: {exc}") from exc

    if not rows:
        return {}
    careers: list[str] = []
    periods: list[str] = []
    parallels: list[str] = []
    journeys: list[str] = []
    first = rows[0]
    for row in rows:
        career = _clean(row.nombre_carrera)
        if career and career not in careers:
            careers.append(career)
        period = _clean(row.detalle_periodo) or _clean(row.codigo_periodo)
        if period and period not in periods:
            periods.append(period)
        row_parallel = _clean(row.paralelo)
        if row_parallel and row_parallel not in parallels:
            parallels.append(row_parallel)
        journey = _clean(row.jornada)
        if journey and journey not in journeys:
            journeys.append(journey)
    return {
        "nombre_carrera": " / ".join(careers) if len(careers) <= 2 else f"{len(careers)} carreras",
        "detalle_periodo": " / ".join(periods),
        "codigo_materia": _clean(first.codigo_materia),
        "cod_materia": _clean(first.cod_materia),
        "nombre_materia": _clean(first.nombre_materia),
        "paralelo": parallels[0] if len(parallels) == 1 else ("Varios" if parallels else ""),
        "cod_jornada": _clean(first.cod_jornada),
        "jornada": (
            journeys[0]
            if len(journeys) == 1
            else "Varias jornadas"
            if journeys
            else (f"Jornada {_clean(first.cod_jornada)}" if _clean(first.cod_jornada) else "")
        ),
        "semestre": _int(first.semestre),
        "unidad_curricular": _clean(first.unidad_curricular),
        "horas": _number(first.horas),
        "es_homologacion": _is_homologation_type(first.tipo_periodo, first.detalle_periodo),
    }


def _teacher_report_paragraph(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(_pdf_text(value), style)


def _teacher_notes_report_pdf(
    teacher: dict[str, Any],
    meta: dict[str, Any],
    students: list[dict[str, Any]],
) -> bytes:
    red = colors.HexColor("#931913")
    light_blue = colors.HexColor("#EAF5F8")
    blue = colors.HexColor("#8DBBC7")
    gray = colors.HexColor("#777777")
    dark = colors.HexColor("#111A3A")
    border = colors.HexColor("#BFC7CC")

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TeacherReportTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontSize=13,
            leading=15,
            textColor=dark,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TeacherReportCenter",
            parent=styles["BodyText"],
            alignment=TA_CENTER,
            fontSize=8,
            leading=9.5,
            textColor=dark,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TeacherReportMeta",
            parent=styles["BodyText"],
            fontSize=7.4,
            leading=9,
            textColor=dark,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TeacherReportCell",
            parent=styles["BodyText"],
            fontSize=5.5,
            leading=6.35,
            textColor=dark,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TeacherReportCellBold",
            parent=styles["TeacherReportCell"],
            fontName="Helvetica-Bold",
            alignment=TA_CENTER,
            textColor=dark,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TeacherReportHeaderWhite",
            parent=styles["TeacherReportCellBold"],
            textColor=colors.white,
            fontSize=5.6,
            leading=6.25,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TeacherReportHeaderLight",
            parent=styles["TeacherReportCellBold"],
            textColor=dark,
            fontSize=5.4,
            leading=6.2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TeacherReportTiny",
            parent=styles["BodyText"],
            fontSize=6.3,
            leading=7.5,
            textColor=gray,
        )
    )

    period_label = _clean(meta.get("detalle_periodo")) or "-"
    subject_name = _clean(meta.get("nombre_materia")) or _clean(meta.get("codigo_materia")) or "-"
    is_homologation = bool(meta.get("es_homologacion")) or any(item.get("es_homologacion") for item in students)
    logo = _template_logo(3.45 * cm)

    def _status(value: Any) -> str:
        number = _number(value)
        if number is None:
            return "Pendiente"
        return "Aprobado" if number >= 7 else "Reprobado"

    story: list[Any] = []

    header_table = Table(
        [
            [
                logo,
                [
                    Paragraph(
                        "INSTITUTO SUPERIOR TECNOLÓGICO DE TÉCNICAS EMPRESARIALES Y DEL CONOCIMIENTO",
                        styles["TeacherReportCenter"],
                    ),
                    Paragraph("Reporte de notas por docente", styles["TeacherReportTitle"]),
                    Paragraph(f"Período académico: {period_label}", styles["TeacherReportCenter"]),
                ],
                Paragraph(
                    f"<b>Emitido:</b><br/>{datetime.now().strftime('%d/%m/%Y %H:%M')}",
                    styles["TeacherReportTiny"],
                ),
            ]
        ],
        colWidths=[4.1 * cm, 17.2 * cm, 6.8 * cm],
    )
    header_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (0, 0), "LEFT"),
                ("ALIGN", (1, 0), (1, 0), "CENTER"),
                ("ALIGN", (2, 0), (2, 0), "RIGHT"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LINEBELOW", (0, 0), (-1, -1), 1.2, red),
            ]
        )
    )
    story.extend([header_table, Spacer(1, 0.18 * cm)])

    jornada_label = _clean(meta.get("jornada")) or (
        f"Jornada {_clean(meta.get('cod_jornada'))}" if _clean(meta.get("cod_jornada")) else "-"
    )
    meta_rows = [
        [
            Paragraph(f"<b>Carrera:</b> {_pdf_text(meta.get('nombre_carrera'))}", styles["TeacherReportMeta"]),
            Paragraph(f"<b>Paralelo:</b> {_pdf_text(meta.get('paralelo'))}", styles["TeacherReportMeta"]),
            Paragraph(f"<b>Jornada:</b> {_pdf_text(jornada_label)}", styles["TeacherReportMeta"]),
        ],
        [
            Paragraph(f"<b>Docente:</b> {_pdf_text(teacher.get('docente'))}", styles["TeacherReportMeta"]),
            Paragraph(f"<b>Asignatura:</b> {_pdf_text(subject_name)}", styles["TeacherReportMeta"]),
            Paragraph(f"<b>Código:</b> {_pdf_text(meta.get('cod_materia') or meta.get('codigo_materia'))}", styles["TeacherReportMeta"]),
        ],
        [
            Paragraph(f"<b>Semestre:</b> {_pdf_text(meta.get('semestre'))}", styles["TeacherReportMeta"]),
            Paragraph(f"<b>Horas:</b> {_legacy_grade_text(meta.get('horas')) if _number(meta.get('horas')) is not None else '-'}", styles["TeacherReportMeta"]),
            Paragraph(f"<b>Estudiantes:</b> {len(students)}", styles["TeacherReportMeta"]),
        ],
    ]
    meta_table = Table(meta_rows, colWidths=[10.0 * cm, 10.0 * cm, 8.1 * cm])
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), light_blue),
                ("BOX", (0, 0), (-1, -1), 0.45, blue),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D6E3E8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([meta_table, Spacer(1, 0.16 * cm)])

    if is_homologation:
        headers = [
            "No.",
            "CARRERA",
            "CEDULA",
            "APELLIDOS Y NOMBRES",
            "TEORIA 40%",
            "PRACTICA 60%",
            "Promedio Final",
            "Estado",
        ]
        col_widths = [0.8 * cm, 4.1 * cm, 2.35 * cm, 9.1 * cm, 2.35 * cm, 2.35 * cm, 2.2 * cm, 2.5 * cm]
        table_rows = [
            [
                _teacher_report_paragraph(index, styles["TeacherReportCell"]),
                _teacher_report_paragraph(item.get("nombre_carrera"), styles["TeacherReportCell"]),
                _teacher_report_paragraph(item.get("cedula"), styles["TeacherReportCell"]),
                _teacher_report_paragraph(item.get("nombre_estudiante"), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("teoria_homo")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("practica_homo")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("promedio_final")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_status(item.get("promedio_final")), styles["TeacherReportCell"]),
            ]
            for index, item in enumerate(students, start=1)
        ]
        table_data: list[list[Any]] = [
            [Paragraph(f"<b>{escape(header)}</b>", styles["TeacherReportHeaderWhite"]) for header in headers]
        ]
        repeat_rows = 1
        table_style_commands: list[tuple[Any, ...]] = [
            ("BACKGROUND", (0, 0), (-1, 0), red),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ]
    else:
        header_group = [
            "No.",
            "CARRERA",
            "CEDULA",
            "APELLIDOS Y NOMBRES",
            "Parcial 1",
            "",
            "",
            "",
            "Parcial 2",
            "",
            "",
            "",
            "Parcial 3",
            "",
            "",
            "",
            "Prom.",
            "Recup.",
            "Final",
            "Estado",
        ]
        header_detail = [
            "",
            "",
            "",
            "",
            "Tareas 30%",
            "Proy. 30%",
            "Examen 40%",
            "Prom.",
            "Tareas 30%",
            "Proy. 30%",
            "Examen 40%",
            "Prom.",
            "Tareas 30%",
            "Proy. 30%",
            "Examen 40%",
            "Prom.",
            "",
            "",
            "",
            "",
        ]
        col_widths = [
            0.62 * cm,
            2.75 * cm,
            2.0 * cm,
            5.35 * cm,
            0.88 * cm,
            0.88 * cm,
            0.88 * cm,
            0.95 * cm,
            0.88 * cm,
            0.88 * cm,
            0.88 * cm,
            0.95 * cm,
            0.88 * cm,
            0.88 * cm,
            0.88 * cm,
            0.95 * cm,
            0.95 * cm,
            0.95 * cm,
            0.95 * cm,
            1.45 * cm,
        ]
        table_rows = [
            [
                _teacher_report_paragraph(index, styles["TeacherReportCell"]),
                _teacher_report_paragraph(item.get("nombre_carrera"), styles["TeacherReportCell"]),
                _teacher_report_paragraph(item.get("cedula"), styles["TeacherReportCell"]),
                _teacher_report_paragraph(item.get("nombre_estudiante"), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p1_tareas")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p1_proyectos")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p1_examen")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("prom_p1")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p2_tareas")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p2_proyectos")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p2_examen")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("prom_p2")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p3_tareas")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p3_proyectos")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("p3_examen")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("prom_p3")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("promedio")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("recuperacion")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_legacy_grade_text(item.get("promedio_final")), styles["TeacherReportCell"]),
                _teacher_report_paragraph(_status(item.get("promedio_final")), styles["TeacherReportCell"]),
            ]
            for index, item in enumerate(students, start=1)
        ]
        table_data = [
            [
                Paragraph(f"<b>{escape(header)}</b>", styles["TeacherReportHeaderWhite" if header else "TeacherReportHeaderLight"])
                for header in header_group
            ],
            [
                Paragraph(f"<b>{escape(header)}</b>", styles["TeacherReportHeaderLight"])
                for header in header_detail
            ],
        ]
        repeat_rows = 2
        table_style_commands = [
            ("BACKGROUND", (0, 0), (-1, 0), red),
            ("BACKGROUND", (0, 1), (-1, 1), light_blue),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("SPAN", (0, 0), (0, 1)),
            ("SPAN", (1, 0), (1, 1)),
            ("SPAN", (2, 0), (2, 1)),
            ("SPAN", (3, 0), (3, 1)),
            ("SPAN", (4, 0), (7, 0)),
            ("SPAN", (8, 0), (11, 0)),
            ("SPAN", (12, 0), (15, 0)),
            ("SPAN", (16, 0), (16, 1)),
            ("SPAN", (17, 0), (17, 1)),
            ("SPAN", (18, 0), (18, 1)),
            ("SPAN", (19, 0), (19, 1)),
        ]

    table_data.extend(table_rows)
    if len(table_rows) == 0:
        table_data.append([Paragraph("Sin estudiantes matriculados.", styles["TeacherReportCell"])] + [""] * (len(col_widths) - 1))

    grade_table = Table(table_data, colWidths=col_widths, repeatRows=repeat_rows)
    grade_table.setStyle(
        TableStyle(
            [
                *table_style_commands,
                ("GRID", (0, 0), (-1, -1), 0.25, border),
                ("FONTNAME", (0, 0), (-1, repeat_rows - 1), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("ALIGN", (1, repeat_rows), (3, -1), "LEFT"),
                ("ROWBACKGROUNDS", (0, repeat_rows), (-1, -1), [colors.white, colors.HexColor("#F7FAFB")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 2.4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2.4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(grade_table)

    finals = [_number(item.get("promedio_final")) for item in students if _number(item.get("promedio_final")) is not None]
    average = round(sum(finals) / len(finals), 2) if finals else None
    story.extend(
        [
            Spacer(1, 0.18 * cm),
            Table(
                [[Paragraph(f"<b>Promedio general del curso:</b> {_legacy_grade_text(average)}", styles["TeacherReportMeta"])]],
                colWidths=[28.1 * cm],
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FCFD")),
                        ("BOX", (0, 0), (-1, -1), 0.35, blue),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                ),
            ),
            Spacer(1, 0.85 * cm),
        ]
    )

    signature_style = styles["TeacherReportCenter"]
    signatures = Table(
        [
            ["____________________________", "____________________________", "____________________________"],
            [
                Paragraph("Secretaria Academica INTEC", signature_style),
                Paragraph("Firma del docente", signature_style),
                Paragraph("Coordinacion Academica", signature_style),
            ],
            [
                "",
                Paragraph(f"CI: {_pdf_text(teacher.get('cedula'))}", signature_style),
                "",
            ],
        ],
        colWidths=[6.4 * cm, 6.4 * cm, 6.4 * cm],
    )
    signatures.setStyle(
        TableStyle(
            [
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    story.append(signatures)

    def draw_page(canvas: Any, _doc: Any) -> None:
        page_width, page_height = landscape(A4)
        canvas.saveState()
        canvas.setStrokeColor(red)
        canvas.setLineWidth(2.0)
        canvas.line(0.55 * cm, page_height - 0.45 * cm, page_width - 0.55 * cm, page_height - 0.45 * cm)
        canvas.setFont("Helvetica", 6.3)
        canvas.setFillColor(gray)
        canvas.drawString(0.6 * cm, 0.42 * cm, "Reporte academico INTEC")
        canvas.drawRightString(page_width - 0.6 * cm, 0.42 * cm, f"Pagina {canvas.getPageNumber()}")
        canvas.restoreState()

    output = BytesIO()
    SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=0.55 * cm,
        leftMargin=0.55 * cm,
        topMargin=0.72 * cm,
        bottomMargin=0.7 * cm,
        title="Notas Por Docente",
    ).build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    output.seek(0)
    return output.getvalue()


def _student_grade_report_rows(students: list[dict[str, Any]]) -> list[dict[str, Any]]:
    groups: dict[str, dict[str, Any]] = {}
    for item in students:
        key = "|".join(
            [
                _clean(item.get("codigo_estud")),
                _clean(item.get("codigo_periodo")),
                _clean(item.get("cod_anio_basica")),
            ]
        )
        if key.strip("|"):
            groups[key] = item

    result: list[dict[str, Any]] = []
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            for source in groups.values():
                cursor.execute(
                    """
                    SELECT
                        TRY_CONVERT(varchar(50), de.codigo_estud) AS codigo_estud,
                        TRY_CONVERT(nvarchar(100), de.Cedula_Est) AS cedula,
                        TRY_CONVERT(nvarchar(4000), de.Apellidos_nombre) AS nombre_estudiante,
                        TRY_CONVERT(varchar(50), cxe.cod_anio_Basica) AS cod_anio_basica,
                        TRY_CONVERT(nvarchar(4000), c.Nombre_Basica) AS nombre_carrera,
                        TRY_CONVERT(varchar(50), cxe.codigo_periodo) AS codigo_periodo,
                        TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo) AS detalle_periodo,
                        TRY_CONVERT(nvarchar(100), pe.TipoMatricula) AS tipo_periodo,
                        TRY_CONVERT(varchar(50), cxe.codigo_materia) AS codigo_materia,
                        TRY_CONVERT(varchar(100), p.cod_materia) AS cod_materia,
                        TRY_CONVERT(nvarchar(4000), p.Nomb_Materia) AS nombre_materia,
                        TRY_CONVERT(float, COALESCE(NULLIF(cxe.Num_Creditos, 0), p.Creditos)) AS creditos,
                        TRY_CONVERT(float, cxe.teoriaHomo) AS teoria_homo,
                        TRY_CONVERT(float, cxe.practicahomo) AS practica_homo,
                        TRY_CONVERT(float, cxe.P1Tareas) AS p1_tareas,
                        TRY_CONVERT(float, cxe.P1Proyectos) AS p1_proyectos,
                        TRY_CONVERT(float, cxe.P1Examen) AS p1_examen,
                        TRY_CONVERT(float, cxe.promP1) AS prom_p1,
                        TRY_CONVERT(float, cxe.P2Tareas) AS p2_tareas,
                        TRY_CONVERT(float, cxe.P2Proyectos) AS p2_proyectos,
                        TRY_CONVERT(float, cxe.P2Examen) AS p2_examen,
                        TRY_CONVERT(float, cxe.promP2) AS prom_p2,
                        TRY_CONVERT(float, cxe.P3Tareas) AS p3_tareas,
                        TRY_CONVERT(float, cxe.P3Proyectos) AS p3_proyectos,
                        TRY_CONVERT(float, cxe.P3Examen) AS p3_examen,
                        TRY_CONVERT(float, cxe.promP3) AS prom_p3,
                        TRY_CONVERT(float, cxe.Promedio) AS promedio,
                        TRY_CONVERT(float, cxe.Recuperacion) AS recuperacion,
                        COALESCE(
                            TRY_CONVERT(float, cxe.PromedioFinal),
                            CASE
                                WHEN (
                                        UPPER(LTRIM(RTRIM(COALESCE(TRY_CONVERT(nvarchar(50), cxe.TipoMatricula), N'')))) = N'H'
                                     OR UPPER(COALESCE(TRY_CONVERT(nvarchar(4000), pe.Detalle_Periodo), N'')) LIKE N'%HOMO%'
                                     )
                                 AND TRY_CONVERT(float, cxe.teoriaHomo) IS NOT NULL
                                 AND TRY_CONVERT(float, cxe.practicahomo) IS NOT NULL
                                THEN (TRY_CONVERT(float, cxe.teoriaHomo) * 0.4) + (TRY_CONVERT(float, cxe.practicahomo) * 0.6)
                            END,
                            CASE
                                WHEN TRY_CONVERT(float, cxe.promP1) IS NOT NULL
                                 AND TRY_CONVERT(float, cxe.promP2) IS NOT NULL
                                 AND TRY_CONVERT(float, cxe.promP3) IS NOT NULL
                                THEN (TRY_CONVERT(float, cxe.promP1) + TRY_CONVERT(float, cxe.promP2) + TRY_CONVERT(float, cxe.promP3)) / 3
                            END,
                            TRY_CONVERT(float, cxe.Promedio)
                        ) AS promedio_final,
                        COALESCE(TRY_CONVERT(float, pe.NotaAprobar), 7) AS nota_aprobar,
                        TRY_CONVERT(nvarchar(50), cxe.TipoMatricula) AS tipo_matricula,
                        TRY_CONVERT(nvarchar(50), cxe.paralelo) AS paralelo,
                        TRY_CONVERT(int, p.Orden) AS orden_materia
                    FROM dbo.CARRERAXESTUD cxe
                    INNER JOIN dbo.DATOS_ESTUD de
                      ON TRY_CONVERT(int, de.codigo_estud) = TRY_CONVERT(int, cxe.codigo_estud)
                    LEFT JOIN dbo.CARRERAS c
                      ON TRY_CONVERT(int, c.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
                    LEFT JOIN dbo.PERIODO pe
                      ON TRY_CONVERT(int, pe.cod_periodo) = TRY_CONVERT(int, cxe.codigo_periodo)
                    LEFT JOIN dbo.PENSUM p
                      ON TRY_CONVERT(int, p.Cod_AnioBasica) = TRY_CONVERT(int, cxe.cod_anio_Basica)
                     AND TRY_CONVERT(int, p.codigo_materia) = TRY_CONVERT(int, cxe.codigo_materia)
                    WHERE TRY_CONVERT(int, cxe.codigo_estud) = ?
                      AND TRY_CONVERT(int, cxe.codigo_periodo) = ?
                      AND TRY_CONVERT(int, cxe.cod_anio_Basica) = ?
                    ORDER BY TRY_CONVERT(int, p.Orden), TRY_CONVERT(nvarchar(4000), p.Nomb_Materia)
                    """,
                    _int(source.get("codigo_estud")),
                    _int(source.get("codigo_periodo")),
                    _int(source.get("cod_anio_basica")),
                )
                columns = [column[0] for column in cursor.description or []]
                rows = [{column: getattr(row, column) for column in columns} for row in cursor.fetchall()]
                if rows:
                    result.extend(rows)
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error consultando reporte de notas del estudiante: {exc}") from exc
    return result


def _student_secretaria_notes_pdf(
    profile: dict[str, Any],
    items: list[dict[str, Any]],
    report_type: str,
    period_label: str,
) -> bytes:
    output = BytesIO()
    canvas = Canvas(output, pagesize=landscape(A4))
    width, height = landscape(A4)
    generated_at = datetime.now()
    hour = generated_at.hour % 12 or 12
    display_date = f"{generated_at.day} de {generated_at.strftime('%B')} de {generated_at.year}"

    def clean(value: Any, fallback: str = "") -> str:
        text = _clean(value)
        return text or fallback

    def fit_text(text: Any, max_chars: int) -> str:
        value = clean(text, "-")
        if len(value) <= max_chars:
            return value
        return value[: max(max_chars - 1, 1)].rstrip() + "…"

    def grade(value: Any, empty: str = "") -> str:
        number = _number(value)
        if number is None:
            return empty
        return f"{number:.2f}"

    def credit(value: Any) -> str:
        number = _number(value)
        if number is None:
            return ""
        return f"{number:.2f}".replace(".", ",")

    def final_status(item: dict[str, Any]) -> str:
        return _grade_result(item.get("promedio_final"))

    def row_values(item: dict[str, Any]) -> list[str]:
        is_homo = _is_homologation_type(
            item.get("tipo_matricula"),
            item.get("detalle_periodo") or item.get("ultimo_periodo"),
            item.get("esquema_calificacion"),
        )
        partial_values = [
            grade(item.get("p1_tareas")),
            grade(item.get("p1_proyectos")),
            grade(item.get("p1_examen")),
            grade(item.get("prom_p1")),
            grade(item.get("p2_tareas")),
            grade(item.get("p2_proyectos")),
            grade(item.get("p2_examen")),
            grade(item.get("prom_p2")),
            grade(item.get("p3_tareas")),
            grade(item.get("p3_proyectos")),
            grade(item.get("p3_examen")),
            grade(item.get("prom_p3")),
        ]
        if is_homo and all(value == "-" for value in partial_values):
            partial_values = [
                grade(item.get("teoria_homo")),
                grade(item.get("practica_homo")),
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
            ]
        subject_name = item.get("nombre_materia") or item.get("codigo_materia")
        subject_max_chars = 50 if _practice_requirement_code(item) else 36
        return [
            fit_text(subject_name, subject_max_chars),
            credit(item.get("creditos")),
            *partial_values,
            grade(item.get("recuperacion")),
            grade(item.get("promedio_final")),
            final_status(item),
        ]

    def period_key(item: dict[str, Any]) -> tuple[int, str]:
        code = _int(item.get("codigo_periodo"))
        label = clean(item.get("detalle_periodo") or item.get("ultimo_periodo") or period_label or "Malla academica general")
        return (code or 999999, label)

    groups: list[tuple[str, list[dict[str, Any]]]] = []
    grouped: dict[tuple[int, str], list[dict[str, Any]]] = {}
    for item in items:
        grouped.setdefault(period_key(item), []).append(item)
    for (_code, label), group_items in sorted(grouped.items(), key=lambda pair: pair[0]):
        groups.append((label, group_items))
    if not groups:
        groups = [(period_label or "Malla academica general", [])]

    def draw_logo() -> None:
        if not _LOGO_PATH.exists():
            canvas.setFont("Helvetica-Bold", 38)
            canvas.setFillColor(colors.HexColor("#808285"))
            canvas.drawString(36, height - 48, "intec")
            return
        drawing = svg2rlg(str(_LOGO_PATH))
        if not drawing:
            return
        target_width = 130
        scale = target_width / float(drawing.width or target_width)
        canvas.saveState()
        canvas.translate(36, height - 58)
        canvas.scale(scale, scale)
        renderPDF.draw(drawing, canvas, 0, 0)
        canvas.restoreState()

    def draw_header(page_number: int) -> None:
        canvas.setFillColor(colors.black)
        canvas.rect(0, height - 5, width, 5, stroke=0, fill=1)
        if page_number == 1:
            draw_logo()
            canvas.setFont("Helvetica-Bold", 18)
            canvas.drawCentredString(width / 2, height - 52, "RECORD ACADÉMICO")
        canvas.setFont("Helvetica-Bold", 7.6)
        career = next((clean(item.get("nombre_carrera")) for item in items if clean(item.get("nombre_carrera"))), "")
        canvas.drawString(54, height - 72, "Estudiante:")
        canvas.drawString(114, height - 72, fit_text(profile.get("nombre_estudiante"), 42))
        canvas.drawString(378, height - 72, "Cédula:")
        canvas.setFont("Helvetica", 7.6)
        canvas.drawString(422, height - 72, fit_text(profile.get("cedula"), 18))
        canvas.setFont("Helvetica-Bold", 7.6)
        canvas.drawString(74, height - 90, "Carrera:")
        canvas.setFont("Helvetica", 7.6)
        canvas.drawString(114, height - 90, fit_text(career, 48))
        canvas.setFont("Helvetica-Bold", 7.6)
        canvas.drawString(62, height - 108, "Modalidad:")
        canvas.setFont("Helvetica", 7.6)
        canvas.drawString(114, height - 108, "En linea")

    x_positions = [30, 228, 282, 306, 330, 354, 402, 426, 450, 474, 522, 546, 570, 594, 642, 684, 740]
    col_widths = [194, 38, 23, 23, 23, 42, 23, 23, 23, 42, 23, 23, 23, 42, 34, 40, 62]
    row_height = 16.2
    header_y = height - 126
    first_row_y = height - 155
    bottom_y = 76

    def center_text(text: str, x: float, y: float, w: float, font: str = "Helvetica", size: float = 7) -> None:
        canvas.setFont(font, size)
        canvas.drawCentredString(x + (w / 2), y, text)

    def draw_table_header() -> None:
        canvas.setFillColor(colors.black)
        canvas.setFont("Helvetica-Bold", 7.2)
        canvas.drawString(54, header_y + 1, "Asignatura")
        center_text("Créditos", x_positions[1], header_y + 1, col_widths[1], "Helvetica-Bold", 7.2)
        center_text("PARCIAL 1", x_positions[2], header_y + 10, sum(col_widths[2:6]), "Helvetica-Bold", 7.2)
        center_text("PARCIAL 2", x_positions[6], header_y + 10, sum(col_widths[6:10]), "Helvetica-Bold", 7.2)
        center_text("PARCIAL 3", x_positions[10], header_y + 10, sum(col_widths[10:14]), "Helvetica-Bold", 7.2)
        center_text("Recup.", x_positions[14], header_y + 1, col_widths[14], "Helvetica-Bold", 6.6)
        center_text("Promedio", x_positions[15], header_y + 9, col_widths[15], "Helvetica-Bold", 6.6)
        center_text("final", x_positions[15], header_y - 1, col_widths[15], "Helvetica-Bold", 6.6)
        center_text("Estado", x_positions[16], header_y + 1, col_widths[16], "Helvetica-Bold", 7.2)
        labels = ["N 1", "N 2", "N3", "PROM 1", "N 1", "N 2", "N3", "PROM 2", "N 1", "N 2", "N3", "PROM 3"]
        for index, label in enumerate(labels, start=2):
            center_text(label, x_positions[index], header_y - 7, col_widths[index], "Helvetica-Bold", 6.3)

    def draw_dotted_line(y: float) -> None:
        canvas.saveState()
        canvas.setDash(1, 3)
        canvas.setStrokeColor(colors.black)
        canvas.setLineWidth(0.45)
        canvas.line(0, y - 5, width, y - 5)
        canvas.restoreState()

    def group_stats(group_items: list[dict[str, Any]]) -> tuple[str, str]:
        finals = [_number(item.get("promedio_final")) for item in group_items if _number(item.get("promedio_final")) is not None]
        credits = [_number(item.get("creditos")) for item in group_items if _number(item.get("creditos")) is not None]
        average = sum(finals) / len(finals) if finals else None
        total_credits = sum(credits) if credits else None
        return credit(total_credits), grade(average)

    def draw_row(item: dict[str, Any], y: float) -> None:
        values = row_values(item)
        canvas.setFont("Helvetica", 6.7)
        canvas.drawString(x_positions[0], y, values[0])
        canvas.setFont("Helvetica", 6.5)
        for index, value in enumerate(values[1:], start=1):
            if index == 16:
                canvas.drawString(x_positions[index], y, fit_text(value, 10))
            else:
                canvas.drawRightString(x_positions[index] + col_widths[index] - 2, y, value)
        draw_dotted_line(y)

    def draw_group_footer(group_items: list[dict[str, Any]], y: float) -> float:
        total_credits, average = group_stats(group_items)
        canvas.setFont("Helvetica-Bold", 7.2)
        canvas.drawString(162, y, "Total Créditos:")
        canvas.drawRightString(258, y, total_credits)
        canvas.setFont("Helvetica", 7.2)
        canvas.drawString(648, y, "Promedio")
        canvas.drawRightString(742, y, average)
        return y - 18

    def draw_final_footer(y: float) -> None:
        footer_y = min(y, 204)
        canvas.setFillColor(colors.black)
        canvas.setFont("Helvetica", 7.2)
        canvas.drawString(32, footer_y - 20, "NOTA:  *  Información basada en los soportes de los archivos y registros académicos que reposan en el Departamento de Secretaría General, cualquier alteración al texto del")
        canvas.drawString(32, footer_y - 32, "presente documento, como enmendadura, tachado, borrón o repisado entre otros lo inválida.")
        canvas.drawString(32, footer_y - 80, "* Este documento tiene una validez si tiene firma y sello del Instituto INTEC")
        canvas.drawRightString(width - 74, footer_y - 102, display_date)
        canvas.line(340, footer_y - 158, 535, footer_y - 158)
        canvas.setFont("Helvetica", 7.2)
        canvas.drawCentredString(437, footer_y - 180, "María Verónica Cevallos Calderón")
        canvas.setFont("Helvetica-Bold", 7.2)
        canvas.drawCentredString(437, footer_y - 198, "Vicerrectora General Académico")

    def new_page(page_number: int) -> float:
        if page_number > 1:
            canvas.showPage()
        draw_header(page_number)
        draw_table_header()
        return first_row_y

    page_number = 1
    y = new_page(page_number)
    for label, group_items in groups:
        if y < bottom_y + (row_height * 3):
            page_number += 1
            y = new_page(page_number)
        canvas.setFont("Helvetica-Bold", 7.2)
        canvas.drawString(36, y, fit_text(label, 78))
        y -= row_height
        if not group_items:
            draw_row({"nombre_materia": "No hay informacion para mostrar.", "estado_academico": "PENDIENTE"}, y)
            y -= row_height
        for item in group_items:
            if y < bottom_y:
                page_number += 1
                y = new_page(page_number)
            draw_row(item, y)
            y -= row_height
        if y < bottom_y:
            page_number += 1
            y = new_page(page_number)
        y = draw_group_footer(group_items, y)

    if y < 224:
        page_number += 1
        y = new_page(page_number)
    draw_final_footer(y)

    canvas.save()
    output.seek(0)
    return output.getvalue()


def _student_grade_report_pdf(
    teacher: dict[str, Any],
    meta: dict[str, Any],
    students: list[dict[str, Any]],
    include_teacher: bool = True,
) -> bytes:
    # La consulta ya esta limitada a una asignacion exacta. El PDF replica el
    # formato historico de Secretaria sin ampliar el historial del estudiante.
    rows = list(students)
    table_width = 539.0
    grid_color = colors.HexColor("#777777")
    text_color = colors.black

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="SecretaryLegacyInstitution",
            parent=styles["BodyText"],
            fontName="Times-Bold",
            fontSize=9.4,
            leading=10.2,
            alignment=TA_CENTER,
            textColor=text_color,
            spaceAfter=1.0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SecretaryLegacyTitle",
            parent=styles["BodyText"],
            fontName="Times-Bold",
            fontSize=9.6,
            leading=10.5,
            alignment=TA_CENTER,
            textColor=text_color,
            spaceAfter=1.0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SecretaryLegacyMeta",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=8.2,
            leading=9.4,
            alignment=TA_CENTER,
            textColor=text_color,
            spaceAfter=0.5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SecretaryLegacyHeader",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=5.05,
            leading=5.45,
            alignment=TA_CENTER,
            textColor=text_color,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SecretaryLegacyCell",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=5.0,
            leading=5.6,
            textColor=text_color,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SecretaryLegacyCellCenter",
            parent=styles["SecretaryLegacyCell"],
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SecretaryLegacySummary",
            parent=styles["BodyText"],
            fontName="Times-BoldItalic",
            fontSize=9.2,
            leading=10.2,
            alignment=TA_RIGHT,
            textColor=text_color,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SecretaryLegacySignature",
            parent=styles["BodyText"],
            fontName="Times-Italic",
            fontSize=8.8,
            leading=10.4,
            alignment=TA_CENTER,
            textColor=text_color,
        )
    )

    def grade(value: Any) -> str:
        number = _number(value)
        return "-" if number is None else f"{number:.2f}".replace(".", ",")

    def compact_number(value: Any) -> str:
        number = _number(value)
        if number is None:
            return "-"
        if number.is_integer():
            return str(int(number))
        return f"{number:.2f}".rstrip("0").rstrip(".").replace(".", ",")

    def partial_average(item: dict[str, Any], partial: int) -> float | None:
        stored = _number(item.get(f"prom_p{partial}"))
        if stored is not None:
            return stored
        return _weighted_regular_partial(
            item.get(f"p{partial}_tareas"),
            item.get(f"p{partial}_proyectos"),
            item.get(f"p{partial}_examen"),
        )

    def regular_average(item: dict[str, Any]) -> float | None:
        stored = _number(item.get("promedio"))
        if stored is not None:
            return stored
        partials = [partial_average(item, partial) for partial in (1, 2, 3)]
        if any(value is None for value in partials):
            return None
        return round(sum(value for value in partials if value is not None) / 3, 2)

    def final_average(item: dict[str, Any]) -> float | None:
        stored = _number(item.get("promedio_final"))
        return stored if stored is not None else regular_average(item)

    def cell(value: Any, *, centered: bool = False) -> Paragraph:
        style = styles["SecretaryLegacyCellCenter"] if centered else styles["SecretaryLegacyCell"]
        return Paragraph(_pdf_text(value), style)

    period_label = _clean(meta.get("detalle_periodo")) or "-"
    subject_name = _clean(meta.get("nombre_materia")) or _clean(meta.get("codigo_materia")) or "-"
    teacher_label = _pdf_text(teacher.get("docente")) if include_teacher else "-"
    jornada_label = _clean(meta.get("cod_jornada")) or _clean(meta.get("jornada")) or "-"
    is_homologation = bool(meta.get("es_homologacion")) or any(item.get("es_homologacion") for item in rows)

    logo = _SvgLogo(_LOGO_PATH, 2.75 * cm)
    logo.hAlign = "CENTER"
    story: list[Any] = [
        logo,
        Spacer(1, 0.35 * cm),
        Paragraph(
            "INSTITUTO SUPERIOR TECNOLÓGICO DE TÉCNICAS EMPRESARIALES Y DEL CONOCIMIENTO",
            styles["SecretaryLegacyInstitution"],
        ),
        Paragraph("Reporte de notas", styles["SecretaryLegacyTitle"]),
        Paragraph(f"<b>Periodo:</b>&nbsp;&nbsp;{_pdf_text(period_label)}", styles["SecretaryLegacyTitle"]),
        Paragraph(
            f"<b>Paralelo:</b>&nbsp;&nbsp;{_pdf_text(meta.get('paralelo'))}"
            f"&nbsp;&nbsp;&nbsp;&nbsp;<b>Jornada:</b>&nbsp;&nbsp;{_pdf_text(jornada_label)}",
            styles["SecretaryLegacyMeta"],
        ),
        Paragraph(
            f"<b>Docente:</b>&nbsp;&nbsp;{teacher_label}"
            f"&nbsp;&nbsp;&nbsp;&nbsp;<b>Asignatura:</b>&nbsp;{_pdf_text(subject_name)}"
            f"&nbsp;&nbsp;&nbsp;&nbsp;<b>Semestre:</b>&nbsp;&nbsp;{_pdf_text(meta.get('semestre'))}"
            f"&nbsp;&nbsp;&nbsp;&nbsp;<b>Horas:</b>&nbsp;&nbsp;{compact_number(meta.get('horas'))}",
            styles["SecretaryLegacyMeta"],
        ),
        Spacer(1, 0.18 * cm),
    ]

    if is_homologation:
        headers = [
            "No.",
            "CARRERA",
            "CEDULA",
            "APELLIDOS Y NOMBRES",
            "TEORÍA 40%",
            "PRÁCTICA 60%",
            "PROMEDIO FINAL",
        ]
        col_widths = [18, 78, 62, 205, 55, 55, 66]
        table_data: list[list[Any]] = [
            [Paragraph(escape(label), styles["SecretaryLegacyHeader"]) for label in headers]
        ]
        for index, item in enumerate(rows, start=1):
            table_data.append(
                [
                    cell(index, centered=True),
                    cell(item.get("nombre_carrera")),
                    cell(item.get("cedula"), centered=True),
                    cell(item.get("nombre_estudiante")),
                    cell(grade(item.get("teoria_homo")), centered=True),
                    cell(grade(item.get("practica_homo")), centered=True),
                    cell(grade(final_average(item)), centered=True),
                ]
            )
    else:
        headers = [
            "No.",
            "CARRERA",
            "CEDULA",
            "APELLIDOS Y NOMBRES",
            "NOTA<br/>1 P1<br/>30%",
            "NOTA<br/>2 P1<br/>30%",
            "NOTA<br/>3 P1<br/>40%",
            "Promedio<br/>Parcial 1",
            "NOTA<br/>1 P2<br/>30%",
            "NOTA<br/>2 P2<br/>30%",
            "NOTA<br/>3 P2<br/>40%",
            "Promedio<br/>Parcial 2",
            "NOTA<br/>1 P3<br/>30%",
            "NOTA<br/>2 P3<br/>30%",
            "NOTA<br/>3 P3<br/>40%",
            "Promedio<br/>Parcial 3",
            "Promedio",
            "Recuperación",
            "Promedio<br/>Final",
        ]
        col_widths = [11, 31, 35, 121, *([17.5, 17.5, 17.5, 28.5] * 3), 29, 40, 29]
        table_data = [[Paragraph(label, styles["SecretaryLegacyHeader"]) for label in headers]]
        for index, item in enumerate(rows, start=1):
            table_data.append(
                [
                    cell(index, centered=True),
                    cell(item.get("nombre_carrera")),
                    cell(item.get("cedula"), centered=True),
                    cell(item.get("nombre_estudiante")),
                    cell(grade(item.get("p1_tareas")), centered=True),
                    cell(grade(item.get("p1_proyectos")), centered=True),
                    cell(grade(item.get("p1_examen")), centered=True),
                    cell(grade(partial_average(item, 1)), centered=True),
                    cell(grade(item.get("p2_tareas")), centered=True),
                    cell(grade(item.get("p2_proyectos")), centered=True),
                    cell(grade(item.get("p2_examen")), centered=True),
                    cell(grade(partial_average(item, 2)), centered=True),
                    cell(grade(item.get("p3_tareas")), centered=True),
                    cell(grade(item.get("p3_proyectos")), centered=True),
                    cell(grade(item.get("p3_examen")), centered=True),
                    cell(grade(partial_average(item, 3)), centered=True),
                    cell(grade(regular_average(item)), centered=True),
                    cell(grade(item.get("recuperacion")), centered=True),
                    cell(grade(final_average(item)), centered=True),
                ]
            )

    if not rows:
        table_data.append(
            [Paragraph("Sin estudiantes matriculados para los filtros seleccionados.", styles["SecretaryLegacyCell"])]
            + [""] * (len(col_widths) - 1)
        )

    row_heights = [27] + ([25.5] * max(len(table_data) - 1, 1))
    grades_table = Table(table_data, colWidths=col_widths, rowHeights=row_heights, repeatRows=1)
    table_commands: list[tuple[Any, ...]] = [
        ("GRID", (0, 0), (-1, -1), 0.45, grid_color),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("ALIGN", (1, 1), (3, -1), "LEFT"),
        ("BACKGROUND", (0, 0), (-1, -1), colors.white),
        ("LEFTPADDING", (0, 0), (-1, -1), 1.1),
        ("RIGHTPADDING", (0, 0), (-1, -1), 1.1),
        ("TOPPADDING", (0, 0), (-1, -1), 1.2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1.2),
    ]
    if not rows:
        table_commands.append(("SPAN", (0, 1), (-1, 1)))
    grades_table.setStyle(TableStyle(table_commands))
    story.append(grades_table)

    finals = [final_average(item) for item in rows]
    completed_finals = [value for value in finals if value is not None]
    course_average = round(sum(completed_finals) / len(completed_finals), 2) if completed_finals else None
    summary = Table(
        [["", Paragraph("Promedio:", styles["SecretaryLegacySummary"]), Paragraph(grade(course_average), styles["SecretaryLegacyMeta"]) ]],
        colWidths=[313, 85, 141],
    )
    summary.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(summary)

    signatures = Table(
        [
            [
                Paragraph("Secretaría Académica INTEC", styles["SecretaryLegacySignature"]),
                Paragraph("Firma del docente", styles["SecretaryLegacySignature"]),
                Paragraph("Coordinación Académica", styles["SecretaryLegacySignature"]),
            ],
            [
                "",
                Paragraph(f"CI: {_pdf_text(teacher.get('cedula'))}", styles["SecretaryLegacySignature"]),
                "",
            ],
        ],
        colWidths=[4.8 * cm] * 3,
    )
    signatures.hAlign = "CENTER"
    signatures.setStyle(
        TableStyle(
            [
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    # Keep a stable blank area immediately above the teacher label. The
    # electronic signature is positioned in this area after the PDF is built.
    story.append(KeepTogether([Spacer(1, 2.9 * cm), signatures]))

    output = BytesIO()
    page_margin = (letter[0] - table_width) / 2
    SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=page_margin,
        leftMargin=page_margin,
        topMargin=0.55 * cm,
        bottomMargin=2.8 * cm,
        title="Notas Por Docente",
        author="Instituto Superior Tecnológico INTEC",
    ).build(story)
    output.seek(0)
    return output.getvalue()


def _teacher_compliance_report_pdf(
    teacher: dict[str, Any],
    meta: dict[str, Any],
    students: list[dict[str, Any]],
    report_format: dict[str, Any],
    params: dict[str, Any],
    evidence_images: list[dict[str, Any]] | None = None,
) -> bytes:
    return _teacher_compliance_model_pdf(teacher, meta, students, report_format, params, evidence_images)


def _teacher_compliance_model_pdf(
    teacher: dict[str, Any],
    meta: dict[str, Any],
    students: list[dict[str, Any]],
    report_format: dict[str, Any],
    params: dict[str, Any],
    evidence_images: list[dict[str, Any]] | None = None,
) -> bytes:
    output = BytesIO()
    canvas = Canvas(output, pagesize=A4)
    width, height = A4
    margin_x = 72
    body_x = 72
    body_right = width - 58
    content_width = body_right - body_x
    continuation_content_y = height - 150
    dark = colors.HexColor("#111111")

    def draw_page_background() -> None:
        if not _TEACHER_COMPLIANCE_BACKGROUND_PATH.exists():
            return
        canvas.saveState()
        canvas.drawImage(
            ImageReader(str(_TEACHER_COMPLIANCE_BACKGROUND_PATH)),
            0,
            0,
            width=width,
            height=height,
            preserveAspectRatio=False,
            mask="auto",
        )
        canvas.restoreState()

    def draw_logo() -> None:
        drawing = svg2rlg(str(_LOGO_PATH)) if _LOGO_PATH.exists() else None
        if drawing:
            scale = 190 / float(drawing.width or 190)
            canvas.saveState()
            canvas.translate(34, height - 104)
            canvas.scale(scale, scale)
            renderPDF.draw(drawing, canvas, 0, 0)
            canvas.restoreState()
        else:
            canvas.setFont("Helvetica-Bold", 58)
            canvas.setFillColor(colors.HexColor("#808285"))
            canvas.drawString(34, height - 86, "intec")

    def draw_header(page_num: int) -> None:
        if not _TEACHER_COMPLIANCE_BACKGROUND_PATH.exists():
            draw_logo()
        x = 242
        y = height - 123
        w = 322
        bottom_h = 38
        row_h = 24
        canvas.setStrokeColor(colors.black)
        canvas.setLineWidth(1)
        canvas.rect(x, y, w, bottom_h + row_h * 2, stroke=1, fill=0)
        canvas.line(x, y + bottom_h, x + w, y + bottom_h)
        canvas.line(x, y + bottom_h + row_h, x + w, y + bottom_h + row_h)
        canvas.line(x + 194, y, x + 194, y + bottom_h)
        canvas.setFillColor(dark)
        canvas.setFont("Times-Bold", 12)
        canvas.drawCentredString(x + w / 2, y + bottom_h + row_h + 7, "Instituto Superior Tecnológico INTEC")
        canvas.drawCentredString(x + w / 2, y + bottom_h + 7, "Vicerrectorado Académico")
        canvas.setFont("Times-Roman", 12)
        canvas.drawCentredString(x + 97, y + 21, "Informe de finalización de asignatura")
        canvas.drawCentredString(x + 97, y + 8, "para pago")
        canvas.setFont("Times-Roman", 14)
        canvas.drawCentredString(x + 258, y + 15, f"Página {page_num} de {total_pages}")

    def draw_watermark() -> None:
        canvas.saveState()
        canvas.setFillColor(colors.Color(0.55, 0.0, 0.0, alpha=0.13))
        canvas.setFont("Helvetica-Bold", 620)
        canvas.drawString(-150, -20, "e")
        canvas.restoreState()

    def draw_footer() -> None:
        canvas.saveState()
        canvas.setFillColor(gray)
        canvas.setFont("Helvetica", 13)
        text = canvas.beginText(width / 2 - 90, 34)
        text.setCharSpace(6)
        text.textLine("www.intec.edu.ec")
        canvas.drawText(text)
        canvas.restoreState()

    def start_page(page_num: int) -> None:
        draw_page_background()
        draw_header(page_num)
        if not _TEACHER_COMPLIANCE_BACKGROUND_PATH.exists():
            draw_watermark()
            draw_footer()
        canvas.setFillColor(dark)

    def new_page(page_num: int) -> None:
        canvas.showPage()
        start_page(page_num)

    def line(text: str, x: float, y: float, size: int = 11, bold: bool = False) -> float:
        canvas.setFillColor(dark)
        canvas.setFont("Times-Bold" if bold else "Times-Roman", size)
        canvas.drawString(x, y, text)
        return y - (size + 4)

    def wrapped(text: str, x: float, y: float, max_width: float, size: int = 11, bold: bool = False, leading: float = 13) -> float:
        canvas.setFont("Times-Bold" if bold else "Times-Roman", size)
        words = _clean(text).split()
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if canvas.stringWidth(candidate, "Times-Bold" if bold else "Times-Roman", size) <= max_width:
                current = candidate
            else:
                canvas.drawString(x, y, current)
                y -= leading
                current = word
        if current:
            canvas.drawString(x, y, current)
            y -= leading
        return y

    def highlight(text: str, x: float, y: float, size: int = 11) -> None:
        canvas.setFont("Times-Bold", size)
        tw = canvas.stringWidth(text, "Times-Bold", size)
        canvas.setFillColor(colors.yellow)
        canvas.rect(x - 1, y - 2, tw + 2, size + 2, stroke=0, fill=1)
        canvas.setFillColor(dark)
        canvas.drawString(x, y, text)

    def evidence_group(*terms: str) -> list[dict[str, Any]]:
        lowered_terms = [term.lower() for term in terms]
        return [
            item
            for item in (evidence_images or [])
            if any(term in _clean(item.get("label")).lower() for term in lowered_terms)
        ]

    def draw_image(content: bytes, x: float, y: float, max_w: float, max_h: float) -> float:
        try:
            image = PILImage.open(BytesIO(content))
            image.verify()
            image = PILImage.open(BytesIO(content))
        except Exception:
            return y
        iw, ih = image.size
        if iw <= 0 or ih <= 0:
            return y
        ratio = min(max_w / iw, max_h / ih, 1)
        draw_w = iw * ratio
        draw_h = ih * ratio
        canvas.drawImage(ImageReader(BytesIO(content)), x, y - draw_h, width=draw_w, height=draw_h, preserveAspectRatio=True, mask="auto")
        return y - draw_h - 12

    def draw_group(terms: tuple[str, ...], x: float, y: float, max_w: float, max_h_each: float, max_count: int | None = None) -> float:
        items = evidence_group(*terms)
        if max_count is not None:
            items = items[:max_count]
        for item in items:
            content = item.get("content")
            if content:
                y = draw_image(content, x, y, max_w, max_h_each)
        return y

    def draw_selected_students_table(x: float, y: float, max_w: float) -> float:
        row_h = 12
        headers = ["No.", "Estudiante", "Cédula", "Carrera", "Final"]
        col_widths = [22, 170, 70, 118, 46]
        table_w = min(sum(col_widths), max_w)
        max_rows = 7
        visible_students = students[:max_rows]
        total_rows = 1 + len(visible_students)
        table_h = total_rows * row_h
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#c7c7c7"))
        canvas.setLineWidth(0.45)
        canvas.setFillColor(colors.HexColor("#f2f2f2"))
        canvas.rect(x, y - row_h, table_w, row_h, stroke=1, fill=1)
        cursor_x = x
        canvas.setFillColor(dark)
        canvas.setFont("Times-Bold", 6.8)
        for header, col_w in zip(headers, col_widths):
            canvas.drawString(cursor_x + 3, y - 8, header)
            canvas.line(cursor_x, y, cursor_x, y - table_h)
            cursor_x += col_w
        canvas.line(x + table_w, y, x + table_w, y - table_h)
        canvas.line(x, y, x + table_w, y)
        canvas.line(x, y - row_h, x + table_w, y - row_h)
        canvas.setFont("Times-Roman", 6.6)
        current_y = y - row_h
        for index, item in enumerate(visible_students, start=1):
            next_y = current_y - row_h
            canvas.line(x, next_y, x + table_w, next_y)
            values = [
                str(index),
                _clean(item.get("nombre_estudiante")) or _clean(item.get("estudiante")) or "-",
                _clean(item.get("cedula")) or _clean(item.get("numero_identificacion")) or "-",
                _clean(item.get("nombre_carrera")) or "-",
                _legacy_grade_text(item.get("promedio_final")),
            ]
            cursor_x = x
            for value, col_w in zip(values, col_widths):
                text = value
                while canvas.stringWidth(text, "Times-Roman", 6.6) > col_w - 6 and len(text) > 4:
                    text = text[:-4].rstrip() + "..."
                canvas.drawString(cursor_x + 3, current_y - 8, text)
                canvas.line(cursor_x, current_y, cursor_x, next_y)
                cursor_x += col_w
            canvas.line(x + table_w, current_y, x + table_w, next_y)
            current_y = next_y
        canvas.restoreState()
        y_after = y - table_h - 8
        if len(students) > max_rows:
            canvas.setFont("Times-Italic", 7)
            canvas.setFillColor(dark)
            canvas.drawString(x, y_after, f"Se muestran {max_rows} de {len(students)} estudiante(s) seleccionados.")
            y_after -= 10
        return y_after

    def draw_teams_recordings_table(recordings: list[dict[str, Any]], x: float, y: float, max_w: float) -> float:
        row_h = 17
        headers = ["Fecha", "Nombre de la grabación", "Duración"]
        col_widths = [92, 290, 103]
        scale = min(1.0, max_w / sum(col_widths))
        col_widths = [value * scale for value in col_widths]
        table_w = sum(col_widths)
        table_h = (len(recordings) + 1) * row_h

        def fitted_text(value: Any, width_value: float, bold: bool = False) -> str:
            text = _clean(value) or "-"
            font_name = "Times-Bold" if bold else "Times-Roman"
            while canvas.stringWidth(text, font_name, 6.4) > width_value - 6 and len(text) > 4:
                text = text[:-4].rstrip() + "..."
            return text

        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#b9c6cc"))
        canvas.setLineWidth(0.45)
        canvas.setFillColor(colors.HexColor("#eaf2f5"))
        canvas.rect(x, y - row_h, table_w, row_h, stroke=1, fill=1)
        canvas.setFillColor(dark)
        canvas.setFont("Times-Bold", 6.4)
        cursor_x = x
        for header, col_w in zip(headers, col_widths):
            canvas.drawString(cursor_x + 3, y - 11, fitted_text(header, col_w, True))
            canvas.line(cursor_x, y, cursor_x, y - table_h)
            cursor_x += col_w
        canvas.line(x + table_w, y, x + table_w, y - table_h)
        canvas.line(x, y, x + table_w, y)
        canvas.line(x, y - row_h, x + table_w, y - row_h)

        current_y = y - row_h
        for index, item in enumerate(recordings):
            next_y = current_y - row_h
            if index % 2:
                canvas.setFillColor(colors.HexColor("#f8fbfc"))
                canvas.rect(x, next_y, table_w, row_h, stroke=0, fill=1)
            canvas.setStrokeColor(colors.HexColor("#b9c6cc"))
            canvas.line(x, next_y, x + table_w, next_y)
            values = [
                item.get("date"),
                item.get("name"),
                item.get("recording_duration") or item.get("call_duration"),
            ]
            cursor_x = x
            canvas.setFillColor(dark)
            canvas.setFont("Times-Roman", 6.4)
            for column_index, (value, col_w) in enumerate(zip(values, col_widths)):
                text = fitted_text(value, col_w)
                if column_index == 1 and _clean(item.get("web_url")):
                    canvas.setFillColor(colors.HexColor("#145da0"))
                    canvas.drawString(cursor_x + 3, current_y - 11, text)
                    canvas.linkURL(
                        _clean(item.get("web_url")),
                        (cursor_x + 2, next_y + 2, cursor_x + col_w - 2, current_y - 2),
                        relative=0,
                    )
                    canvas.setFillColor(dark)
                else:
                    canvas.drawString(cursor_x + 3, current_y - 11, text)
                canvas.line(cursor_x, current_y, cursor_x, next_y)
                cursor_x += col_w
            canvas.line(x + table_w, current_y, x + table_w, next_y)
            current_y = next_y
        canvas.restoreState()
        return y - table_h - 8

    def draw_grade_summary_table(x: float, y: float) -> float:
        headers = ["Nota máxima", "Nota mínima", "Estudiantes reprobados"]
        values = [
            _grade_text(max(grade_values) if grade_values else None),
            _grade_text(min(grade_values) if grade_values else None),
            str(failed),
        ]
        col_w = 118
        row_h = 16
        table_w = col_w * len(headers)
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#c7c7c7"))
        canvas.setLineWidth(0.5)
        canvas.setFillColor(colors.HexColor("#f2f2f2"))
        canvas.rect(x, y - row_h, table_w, row_h, stroke=1, fill=1)
        canvas.setFillColor(dark)
        canvas.setFont("Times-Bold", 7.4)
        for index, header in enumerate(headers):
            cell_x = x + (index * col_w)
            canvas.drawCentredString(cell_x + col_w / 2, y - 10.5, header)
            canvas.line(cell_x, y, cell_x, y - row_h * 2)
        canvas.line(x + table_w, y, x + table_w, y - row_h * 2)
        canvas.line(x, y, x + table_w, y)
        canvas.line(x, y - row_h, x + table_w, y - row_h)
        canvas.line(x, y - row_h * 2, x + table_w, y - row_h * 2)
        canvas.setFont("Times-Bold", 8)
        for index, value in enumerate(values):
            cell_x = x + (index * col_w)
            canvas.drawCentredString(cell_x + col_w / 2, y - row_h - 10.5, value)
        canvas.restoreState()
        return y - (row_h * 2) - 10

    teacher_name = _clean(teacher.get("docente"))
    name_parts = teacher_name.split()
    first_names = " ".join(name_parts[2:]) if len(name_parts) > 2 else teacher_name
    last_names = " ".join(name_parts[:2]) if len(name_parts) > 2 else "-"
    course_name = _clean(meta.get("nombre_materia")) or _clean(meta.get("cod_materia"))
    grade_values = [_number(item.get("promedio_final")) for item in students]
    grade_values = [value for value in grade_values if value is not None]
    failed = sum(1 for value in grade_values if value < 7)
    teams_recordings = [
        item
        for item in (params.get("teams_recordings") or [])
        if isinstance(item, dict) and _clean(item.get("name"))
    ]
    recordings_per_page = 18
    recording_chunks = [
        teams_recordings[index : index + recordings_per_page]
        for index in range(0, len(teams_recordings), recordings_per_page)
    ]
    grade_report_images = evidence_group("reporte de notas firmado")
    total_pages = 4 + len(recording_chunks) + len(grade_report_images)

    start_page(1)
    y = 662
    y = line("1.   DATOS DEL DOCENTE:", body_x, y, 12, True)
    y = line(f"Nombres del Docente: {first_names}", body_x + 38, y, 11)
    y = line(f"Apellidos del Docente: {last_names}", body_x + 38, y, 11)
    y = line(f"Cédula: {_clean(teacher.get('cedula'))}", body_x + 38, y, 11)
    y = line(f"Correo institucional: {_clean(teacher.get('correo')) or _clean(teacher.get('correo_personal'))}", body_x + 38, y, 11)
    y = line(f"Teléfono de contacto: {_clean(params.get('telefono')) or '-'}", body_x + 38, y, 11)
    y -= 14
    canvas.setFont("Times-Bold", 12)
    canvas.drawString(body_x, y, "2.   DATOS DE LA ASIGNATRURA:")
    canvas.setFont("Times-Roman", 11)
    canvas.drawString(body_x + 205, y, f"Asignatura: {course_name}")
    y -= 16
    y = line(f"Fecha de inicio: {_clean(params.get('fecha_inicio')) or '-'}", body_x + 38, y, 11)
    y = line(f"Fecha fin: {_clean(params.get('fecha_fin')) or '-'}", body_x + 38, y, 11)
    y = line(f"Número de estudiantes matriculados: {len(students)}", body_x + 38, y, 11)
    y = draw_selected_students_table(body_x + 38, y + 4, 426)
    y -= 14
    y = line("3.   REPORTE ACADÉMICO", body_x, y, 12, True)
    y -= 14
    canvas.setFont("Times-Bold", 11)
    canvas.drawString(body_x + 18, y, "3.1.")
    canvas.drawString(body_x + 56, y, "Cumplimiento del PEA Y silabo")
    highlight("(debidamente firmado)", body_x + 228, y, 11)
    y -= 24
    y = wrapped("Evidenciar (captura de pantalla) silabo y PEA cargado en el sistema de Aula virtuales, debidamente firmado electrónicamente.", body_x, y, content_width, 11)
    y = line("Ejemplo:", body_x, y - 4, 11)
    y = draw_group(("pea", "sílabo", "silabo"), body_x + 26, y, 468, 68, 1)
    y -= 4
    canvas.setFont("Times-Bold", 11)
    canvas.drawString(body_x + 18, y, "3.2.")
    canvas.drawString(body_x + 56, y, "Reporte de actualización del silabo")
    highlight("(Describir actualizaciones realizadas al sílabo", body_x + 260, y, 11)
    y -= 14
    highlight("y su justificativo)", body_x + 56, y, 11)
    y -= 28
    y = line(_clean(params.get("actualizaciones")) or "Sin cambios realizados.", body_x, y, 11, True)
    y -= 22
    y = line("3.3.     Reporte del aula virtual.", body_x + 18, y, 11, True)
    y -= 16
    y = wrapped("En el reporte consolidado evidencia en el sistema de aulas virtuales que se cargaron los siguientes recursos en material académico a través de capturas de pantalla:", body_x, y, content_width, 11)
    for item in [
        "Bibliografía del material académico",
        "Presentación PPT cargado como PDF por cada clase.",
        "Link de grabaciones de cada clase o tutoría impartida",
        "Simulador de examen (para los casos que aplique) y su banco de preguntas.",
    ]:
        y = line(f"•    {item}", body_x + 18, y - 1, 11)

    new_page(2)
    y = continuation_content_y
    for item in ["Evaluación(es) teórica(s)", "Componente(s) práctico(s)", "Evidencia de clases grabadas en TEAMS:"]:
        y = line(f"•    {item}", body_x + 18, y, 11)
    y = draw_group(("aula", "virtual", "recursos"), body_x + 18, y - 6, 494, 150, 3)

    if recording_chunks:
        for chunk_index, recording_chunk in enumerate(recording_chunks):
            new_page(3 + chunk_index)
            y = continuation_content_y
            team_name = _clean(recording_chunk[0].get("team_name")) if recording_chunk else ""
            y = line(
                f"Clases grabadas en TEAMS ({len(teams_recordings)} registro(s) obtenidos desde Microsoft Graph)",
                body_x + 18,
                y,
                10,
                True,
            )
            y = line(
                "Fecha, nombre y duración de cada grabación disponible en Microsoft Teams.",
                body_x + 18,
                y,
                8,
            )
            if team_name:
                y = wrapped(f"Equipo: {team_name}", body_x + 18, y, content_width, 9, False, 11)
            if chunk_index == 0 and _clean(params.get("observaciones")):
                y = wrapped(_clean(params.get("observaciones")), body_x + 18, y, content_width, 9, False, 11)
            y = draw_teams_recordings_table(recording_chunk, body_x + 18, y - 4, content_width)
            canvas.setFont("Times-Italic", 7)
            canvas.setFillColor(dark)
            canvas.drawString(
                body_x + 18,
                y,
                f"Bloque {chunk_index + 1} de {len(recording_chunks)}. Los nombres azules enlazan al archivo en Microsoft 365.",
            )
        new_page(3 + len(recording_chunks))
        y = continuation_content_y
    else:
        new_page(3)
        y = continuation_content_y
        y = line("Clases grabadas en TEAMS", body_x + 18, y, 10, True)
        y = wrapped(
            "Microsoft Graph no encontró grabaciones seleccionadas para la materia y el rango de fechas del informe.",
            body_x + 18,
            y,
            content_width,
            9,
            False,
            11,
        )
        y = draw_group(("teams", "clases"), body_x + 18, y - 4, 494, 105, 1)
        if _clean(params.get("observaciones")):
            y = wrapped(_clean(params.get("observaciones")), body_x + 18, y - 4, content_width, 9, False, 11)
        y -= 42
    y = line("3.4.     Reporte de Notas", body_x + 18, y, 11, True)
    y -= 14
    y = wrapped(
        "El sistema genera automáticamente el reporte detallado de notas en formato Secretaría con la misma "
        "asignatura, períodos y estudiantes seleccionados. Primero se firma electrónicamente el reporte de notas; "
        "después, cada página firmada se incorpora como imagen al presente informe. El PDF original firmado se "
        "mantiene como documento independiente para su descarga y validación:",
        body_x,
        y,
        content_width,
        11,
    )
    y -= 8
    y = draw_grade_summary_table(body_x + 24, y)

    for image_index, item in enumerate(grade_report_images):
        new_page(4 + len(recording_chunks) + image_index)
        y = continuation_content_y
        y = line(
            f"3.4.1.   Reporte de notas firmado - página {image_index + 1} de {len(grade_report_images)}",
            body_x + 18,
            y,
            10,
            True,
        )
        y = wrapped(
            "Imagen generada automáticamente desde el PDF de notas firmado electrónicamente por el docente.",
            body_x + 18,
            y,
            content_width,
            8,
            False,
            10,
        )
        content = item.get("content")
        if content:
            y = draw_image(content, body_x + 8, y - 6, content_width - 16, 520)

    new_page(4 + len(recording_chunks) + len(grade_report_images))
    y = 662
    y = line("3.5.     Anexos:", body_x + 18, y, 12, True)
    y -= 16
    y = wrapped("El presente informe debe ir acompañado de la siguiente documentación de respaldo:", body_x, y, content_width, 11)
    for item in [
        "Contrato firmado electrónicamente",
        "Reporte de notas firmado electrónicamente",
        "Factura electrónica, emitida de acuerdo al número de contrato y valor",
    ]:
        y = line(f"      -    {item}", body_x + 18, y - 2, 11)
    y -= 34
    y = line("Saludos cordiales,", body_x, y, 11)
    y -= 84
    y = line("Firma electrónica", body_x, y, 11)
    y = line(teacher_name.title(), body_x, y, 11)
    y = line("DOCENTE", body_x, y, 11)
    line(f"Cédula: {_clean(teacher.get('cedula'))}", body_x, y, 11)
    canvas.save()
    output.seek(0)
    return output.getvalue()

def _teacher_compliance_report_pdf_legacy(
    teacher: dict[str, Any],
    meta: dict[str, Any],
    students: list[dict[str, Any]],
    report_format: dict[str, Any],
    params: dict[str, Any],
    evidence_images: list[dict[str, Any]] | None = None,
) -> bytes:
    light_gray = colors.HexColor("#F4F4F4")
    dark = colors.HexColor("#111111")
    border = colors.HexColor("#BFC7CC")

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ComplianceTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=14,
            textColor=dark,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ComplianceSection",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=dark,
            spaceBefore=10,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ComplianceBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12.5,
            textColor=dark,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ComplianceJustify",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12.5,
            textColor=dark,
            alignment=TA_JUSTIFY,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ComplianceSmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=7.4,
            leading=9,
            textColor=dark,
        )
    )

    def p(value: Any, style: str = "ComplianceBody") -> Paragraph:
        return Paragraph(_pdf_text(value), styles[style])

    def bp(value: Any, style: str = "ComplianceBody") -> Paragraph:
        return Paragraph(f"<b>{_pdf_text(value)}</b>", styles[style])

    def section(title: Any, body: Any | None = None, body_style: str = "ComplianceBody") -> None:
        story.append(bp(title, "ComplianceBody"))
        if body:
            story.append(p(body, body_style))
            story.append(Spacer(1, 0.08 * cm))

    def grade_cell(value: Any) -> Paragraph:
        return Paragraph(_pdf_text(value), styles["ComplianceSmall"])

    def add_grades_annex() -> None:
        if not students:
            story.append(p("No existen estudiantes seleccionados para adjuntar calificaciones.", "ComplianceSmall"))
            return
        story.append(Spacer(1, 0.1 * cm))
        story.append(p("Cuadro de notas", "ComplianceBody"))
        story.append(
            p(
                (
                    f"Periodo: {_clean(meta.get('detalle_periodo')) or '-'} | "
                    f"Paralelo: {_clean(meta.get('paralelo')) or '-'} | "
                    f"Jornada: {_clean(meta.get('jornada')) or '-'} | "
                    f"Semestre: {_clean(meta.get('semestre')) or '-'} | "
                    f"Horas: {_grade_text(meta.get('horas'), 0)}"
                ),
                "ComplianceSmall",
            )
        )
        rows: list[list[Any]] = [[
            grade_cell("No."),
            grade_cell("Carrera"),
            grade_cell("Cédula"),
            grade_cell("Apellidos y nombres"),
            grade_cell("P1"),
            grade_cell("P2"),
            grade_cell("P3"),
            grade_cell("Promedio"),
            grade_cell("Recuperación"),
            grade_cell("Final"),
        ]]
        for index, item in enumerate(students, start=1):
            rows.append([
                grade_cell(index),
                grade_cell(item.get("nombre_carrera")),
                grade_cell(item.get("cedula")),
                grade_cell(item.get("nombre_estudiante")),
                grade_cell(_legacy_grade_text(item.get("prom_p1"))),
                grade_cell(_legacy_grade_text(item.get("prom_p2"))),
                grade_cell(_legacy_grade_text(item.get("prom_p3"))),
                grade_cell(_legacy_grade_text(item.get("promedio"))),
                grade_cell(_legacy_grade_text(item.get("recuperacion"))),
                grade_cell(_legacy_grade_text(item.get("promedio_final"))),
            ])
        table = Table(
            rows,
            repeatRows=1,
            colWidths=[0.7 * cm, 2.5 * cm, 2.0 * cm, 4.3 * cm, 1.15 * cm, 1.15 * cm, 1.15 * cm, 1.45 * cm, 1.65 * cm, 1.25 * cm],
        )
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), light_gray),
                    ("BOX", (0, 0), (-1, -1), 0.45, border),
                    ("INNERGRID", (0, 0), (-1, -1), 0.3, border),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (0, 0), (0, -1), "CENTER"),
                    ("ALIGN", (4, 1), (-1, -1), "CENTER"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 3),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(table)

    grade_values = [_number(item.get("promedio_final")) for item in students]
    grade_values = [value for value in grade_values if value is not None]
    failed = sum(1 for value in grade_values if value < 7)
    max_grade = max(grade_values) if grade_values else None
    min_grade = min(grade_values) if grade_values else None

    def evidence_group(*terms: str) -> list[dict[str, Any]]:
        lowered_terms = [term.lower() for term in terms]
        return [
            item
            for item in (evidence_images or [])
            if any(term in _clean(item.get("label")).lower() for term in lowered_terms)
        ]

    def add_evidence(*terms: str, always_example: bool = False) -> None:
        flowables = _image_evidence_flowables(evidence_group(*terms), styles)
        if flowables or always_example:
            story.append(p("Ejemplo:", "ComplianceBody"))
        if flowables:
            story.extend(flowables)

    teacher_name = _clean(teacher.get("docente"))
    name_parts = teacher_name.split()
    first_names = " ".join(name_parts[2:]) if len(name_parts) > 2 else teacher_name
    last_names = " ".join(name_parts[:2]) if len(name_parts) > 2 else "-"
    course_name = _clean(meta.get("nombre_materia")) or _clean(meta.get("cod_materia"))

    story: list[Any] = []
    story.append(bp("DATOS DEL DOCENTE:", "ComplianceBody"))
    story.append(p(f"Nombres del Docente: {first_names}   Apellidos del Docente: {last_names}", "ComplianceBody"))
    story.append(p(f"Cédula: {_clean(teacher.get('cedula'))}", "ComplianceBody"))
    story.append(p(f"Correo institucional: {_clean(teacher.get('correo')) or _clean(teacher.get('correo_personal'))}", "ComplianceBody"))
    story.append(p(f"Teléfono de contacto: {_clean(params.get('telefono')) or '-'}", "ComplianceBody"))
    story.append(Spacer(1, 0.28 * cm))
    story.append(Paragraph(f"<b>DATOS DE LA ASIGNATRURA:</b> Asignatura: {_pdf_text(course_name)}", styles["ComplianceBody"]))
    story.append(p(f"Fecha de inicio: {_clean(params.get('fecha_inicio')) or '-'}", "ComplianceBody"))
    story.append(p(f"Fecha fin: {_clean(params.get('fecha_fin')) or '-'}", "ComplianceBody"))
    story.append(p(f"Número de estudiantes matriculados: {len(students)}", "ComplianceBody"))
    initial_flowables = _image_evidence_flowables(evidence_group("datos", "matriculados", "inicial"), styles)
    if initial_flowables:
        story.append(Spacer(1, 0.12 * cm))
        story.extend(initial_flowables)
    story.append(Spacer(1, 0.35 * cm))
    story.append(bp("REPORTE ACADÉMICO", "ComplianceBody"))
    story.append(Spacer(1, 0.2 * cm))

    story.append(bp("Cumplimiento del PEA Y silabo (debidamente firmado)", "ComplianceBody"))
    story.append(p("Evidenciar (captura de pantalla) silabo y PEA cargado en el sistema de Aula virtuales, debidamente firmado electrónicamente.", "ComplianceBody"))
    add_evidence("pea", "sílabo", "silabo", always_example=True)

    story.append(Spacer(1, 0.16 * cm))
    story.append(bp("Reporte de actualización del silabo (Describir actualizaciones realizadas al sílabo y su justificativo)", "ComplianceBody"))
    story.append(bp(_clean(params.get("actualizaciones")) or "Sin cambios realizados.", "ComplianceBody"))

    story.append(Spacer(1, 0.16 * cm))
    story.append(bp("Reporte del aula virtual.", "ComplianceBody"))
    story.append(p("En el reporte consolidado evidencia en el sistema de aulas virtuales que se cargaron los siguientes recursos en material académico a través de capturas de pantalla:", "ComplianceJustify"))
    add_evidence("aula", "virtual", "recursos")
    for item in report_format.get("resources") or []:
        story.append(p(item, "ComplianceBody"))

    story.append(Spacer(1, 0.16 * cm))
    story.append(p("Evidencia de clases grabadas en TEAMS:", "ComplianceBody"))
    if _clean(params.get("observaciones")):
        story.append(p(_clean(params.get("observaciones")), "ComplianceBody"))
    add_evidence("teams", "clases")

    story.append(Spacer(1, 0.16 * cm))
    story.append(bp("Asistencias", "ComplianceBody"))
    add_evidence("asistencia")
    section(
        "Reporte de Notas",
        (
            "Indicar la nota máxima obtenida y la nota mínima obtenida y si existieron casos de estudiantes reprobados, "
            "junto con captura de pantalla del reporte de notas debidamente firmado electrónicamente y de las notas subidas en el sistema académico:"
        ),
        "ComplianceJustify",
    )
    add_evidence("notas", "reporte", always_example=True)
    story.append(
        p(
            (
                f"Resumen generado: Nota máxima: {_grade_text(max_grade)}. "
                f"Nota mínima: {_grade_text(min_grade)}. Estudiantes reprobados: {failed}."
            ),
            "ComplianceSmall",
        )
    )
    add_grades_annex()
    story.append(Spacer(1, 0.22 * cm))
    story.append(bp("Anexos:", "ComplianceBody"))
    story.append(p("El presente informe debe ir acompañado de la siguiente documentación de respaldo:", "ComplianceBody"))
    for item in report_format.get("annexes") or []:
        story.append(p(item, "ComplianceBody"))
    story.append(Spacer(1, 0.55 * cm))
    story.append(p("Saludos cordiales,", "ComplianceBody"))
    story.append(Spacer(1, 0.9 * cm))
    story.append(p("Firma electrónica", "ComplianceBody"))
    story.append(p(teacher_name, "ComplianceBody"))
    story.append(p("DOCENTE", "ComplianceBody"))
    story.append(p(f"Cédula: {_clean(teacher.get('cedula'))}", "ComplianceBody"))

    output = BytesIO()
    SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=2.1 * cm,
        leftMargin=2.1 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
        title="Informe de cumplimiento docente",
    ).build(story)
    output.seek(0)
    return output.getvalue()


def _docx_clear_body(document: Any) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _docx_paragraph(document: Any, text: str = "", bold: bool = False, justify: bool = False, space_after: int = 0) -> Any:
    paragraph = document.add_paragraph()
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return paragraph


def _docx_add_picture(document: Any, image_bytes: bytes, width_cm: float = 16.6) -> None:
    try:
        image = PILImage.open(BytesIO(image_bytes))
        image.verify()
    except Exception:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(BytesIO(image_bytes), width=Cm(width_cm))


def _teacher_compliance_report_docx(
    teacher: dict[str, Any],
    meta: dict[str, Any],
    students: list[dict[str, Any]],
    report_format: dict[str, Any],
    params: dict[str, Any],
    evidence_images: list[dict[str, Any]] | None = None,
) -> bytes:
    template_path = _TEACHER_COMPLIANCE_WORD_TEMPLATE_PATH
    document = Document(str(template_path)) if template_path.exists() else Document()
    _docx_clear_body(document)

    def evidence_group(*terms: str) -> list[dict[str, Any]]:
        lowered_terms = [term.lower() for term in terms]
        return [
            item
            for item in (evidence_images or [])
            if any(term in _clean(item.get("label")).lower() for term in lowered_terms)
        ]

    def add_evidence(*terms: str, width_cm: float = 16.6, always_example: bool = False) -> None:
        images = evidence_group(*terms)
        if images or always_example:
            _docx_paragraph(document, "Ejemplo:")
        for item in images:
            content = item.get("content")
            if content:
                _docx_add_picture(document, content, width_cm)

    teacher_name = _clean(teacher.get("docente"))
    name_parts = teacher_name.split()
    first_names = " ".join(name_parts[2:]) if len(name_parts) > 2 else teacher_name
    last_names = " ".join(name_parts[:2]) if len(name_parts) > 2 else "-"
    course_name = _clean(meta.get("nombre_materia")) or _clean(meta.get("cod_materia"))

    grade_values = [_number(item.get("promedio_final")) for item in students]
    grade_values = [value for value in grade_values if value is not None]
    failed = sum(1 for value in grade_values if value < 7)
    max_grade = max(grade_values) if grade_values else None
    min_grade = min(grade_values) if grade_values else None

    _docx_paragraph(document, "DATOS DEL DOCENTE:", bold=True)
    _docx_paragraph(document, f"Nombres del Docente: {first_names}")
    _docx_paragraph(document, f"Apellidos del Docente: {last_names}")
    _docx_paragraph(document, f"Cédula: {_clean(teacher.get('cedula'))}")
    _docx_paragraph(document, f"Correo institucional: {_clean(teacher.get('correo')) or _clean(teacher.get('correo_personal'))}")
    _docx_paragraph(document, f"Teléfono de contacto: {_clean(params.get('telefono')) or '-'}")
    _docx_paragraph(document)
    _docx_paragraph(document, f"DATOS DE LA ASIGNATRURA:  Asignatura: {course_name}", bold=True)
    _docx_paragraph(document, f"Fecha de inicio: {_clean(params.get('fecha_inicio')) or '-'}")
    _docx_paragraph(document, f"Fecha fin: {_clean(params.get('fecha_fin')) or '-'}")
    _docx_paragraph(document, f"Número de estudiantes matriculados: {len(students)}")
    for item in evidence_group("datos", "matriculados", "inicial"):
        if item.get("content"):
            _docx_add_picture(document, item["content"], 13.4)

    _docx_paragraph(document)
    _docx_paragraph(document, "REPORTE ACADÉMICO", bold=True)
    _docx_paragraph(document)
    _docx_paragraph(document, "Cumplimiento del PEA Y silabo (debidamente firmado)", bold=True)
    _docx_paragraph(
        document,
        "Evidenciar (captura de pantalla) silabo y PEA cargado en el sistema de Aula virtuales, debidamente firmado electrónicamente.",
    )
    add_evidence("pea", "sílabo", "silabo", width_cm=16.6, always_example=True)

    _docx_paragraph(document)
    _docx_paragraph(document, "Reporte de actualización del silabo (Describir actualizaciones realizadas al sílabo y su justificativo)", bold=True)
    _docx_paragraph(document, _clean(params.get("actualizaciones")) or "Sin cambios realizados.", bold=True)

    _docx_paragraph(document)
    _docx_paragraph(document, "Reporte del aula virtual.", bold=True)
    _docx_paragraph(
        document,
        "En el reporte consolidado evidencia en el sistema de aulas virtuales que se cargaron los siguientes recursos en material académico a través de capturas de pantalla:",
        justify=True,
    )
    for item in report_format.get("resources") or []:
        _docx_paragraph(document, item)
    for item in evidence_group("aula", "virtual", "recursos"):
        if item.get("content"):
            _docx_add_picture(document, item["content"], 16.6)

    _docx_paragraph(document, "Evidencia de clases grabadas en TEAMS:")
    if _clean(params.get("observaciones")):
        _docx_paragraph(document, _clean(params.get("observaciones")))
    for item in evidence_group("teams", "clases"):
        if item.get("content"):
            _docx_add_picture(document, item["content"], 16.6)

    _docx_paragraph(document)
    _docx_paragraph(document, "Asistencias", bold=True)
    for item in evidence_group("asistencia"):
        if item.get("content"):
            _docx_add_picture(document, item["content"], 16.6)

    _docx_paragraph(document)
    _docx_paragraph(document, "Reporte de Notas", bold=True)
    _docx_paragraph(
        document,
        (
            "Indicar la nota máxima obtenida y la nota mínima obtenida y si existieron casos de estudiantes reprobados, "
            "junto con captura de pantalla del reporte de notas debidamente firmado electrónicamente y de las notas subidas en el sistema académico:"
        ),
        justify=True,
    )
    add_evidence("notas", "reporte", width_cm=16.6, always_example=True)
    _docx_paragraph(
        document,
        f"Resumen generado: Nota máxima: {_grade_text(max_grade)}. Nota mínima: {_grade_text(min_grade)}. Estudiantes reprobados: {failed}.",
    )

    if students:
        _docx_paragraph(document, "Cuadro de notas", bold=True)
        table = document.add_table(rows=1, cols=7)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        headers = ["No.", "Carrera", "Cédula", "Apellidos y nombres", "P1", "P2", "Final"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row_index, item in enumerate(students, start=1):
            cells = table.add_row().cells
            values = [
                str(row_index),
                _clean(item.get("nombre_carrera")),
                _clean(item.get("cedula")),
                _clean(item.get("nombre_estudiante")),
                _legacy_grade_text(item.get("prom_p1")),
                _legacy_grade_text(item.get("prom_p2")),
                _legacy_grade_text(item.get("promedio_final")),
            ]
            for index, value in enumerate(values):
                cells[index].text = value

    _docx_paragraph(document)
    _docx_paragraph(document, "Anexos:", bold=True)
    _docx_paragraph(document, "El presente informe debe ir acompañado de la siguiente documentación de respaldo:")
    for item in report_format.get("annexes") or []:
        _docx_paragraph(document, item)
    _docx_paragraph(document)
    _docx_paragraph(document, "Saludos cordiales,")
    _docx_paragraph(document)
    _docx_paragraph(document)
    _docx_paragraph(document, "Firma electrónica")
    _docx_paragraph(document, teacher_name)
    _docx_paragraph(document, "DOCENTE")
    _docx_paragraph(document, f"Cédula: {_clean(teacher.get('cedula'))}")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


def _docx_bytes_to_pdf_bytes(docx_bytes: bytes, filename_stem: str) -> bytes:
    try:
        import pythoncom  # type: ignore[import-not-found]
        import win32com.client  # type: ignore[import-not-found]
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="No se puede convertir el informe a PDF porque falta pywin32 en el entorno del backend.",
        ) from exc

    with TemporaryDirectory(prefix="intec_compliance_") as temp_dir:
        temp_path = Path(temp_dir)
        safe_stem = (_safe_filename(filename_stem) or "informe-cumplimiento")[:90]
        docx_path = temp_path / f"{safe_stem}.docx"
        pdf_path = temp_path / f"{safe_stem}.pdf"
        docx_path.write_bytes(docx_bytes)

        pythoncom.CoInitialize()
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(str(docx_path), ReadOnly=True)
            try:
                document.ExportAsFixedFormat(str(pdf_path), 17)
            finally:
                document.Close(False)
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail="No se pudo convertir el informe Word a PDF. Verifica que Microsoft Word esté instalado y disponible en Windows.",
            ) from exc
        finally:
            if word is not None:
                word.Quit()
            pythoncom.CoUninitialize()

        if not pdf_path.exists() or pdf_path.stat().st_size == 0:
            raise HTTPException(status_code=500, detail="La conversión del informe a PDF no generó un archivo válido.")
        return pdf_path.read_bytes()


@router.get("/admin/teacher-compliance-format")
def get_teacher_compliance_format(
    current_user: Annotated[SessionUser, Depends(_PORTAL_ADMIN_ACCESS)],
) -> dict[str, Any]:
    return _read_teacher_compliance_format()


@router.put("/admin/teacher-compliance-format")
def update_teacher_compliance_format(
    payload: TeacherComplianceReportFormat,
    current_user: Annotated[SessionUser, Depends(_PORTAL_ADMIN_ACCESS)],
) -> dict[str, Any]:
    return _write_teacher_compliance_format(payload)


def _build_teacher_compliance_pdf(
    current_user: SessionUser,
    codigo_periodo: list[int],
    codigo_materia: str,
    paralelo: str,
    codigo_estud: list[int] | None = None,
    cod_anio_basica: int | None = None,
    cod_jornada: int | None = None,
    fecha_inicio: str = "",
    fecha_fin: str = "",
    telefono: str = "",
    actualizaciones: str = "",
    observaciones: str = "",
    teams_recordings: list[dict[str, Any]] | None = None,
    evidence_images: list[dict[str, Any]] | None = None,
) -> tuple[bytes, str]:
    codigo_doc = _teacher_code(current_user)
    parallel = paralelo.strip().upper()
    subject_filter = _clean(codigo_materia).upper()
    period_codes = list(dict.fromkeys(codigo_periodo))
    if not period_codes:
        raise HTTPException(status_code=400, detail="Debe seleccionar al menos un periodo")
    if len(period_codes) > 4:
        raise HTTPException(status_code=400, detail="Solo se pueden seleccionar hasta 4 periodos para el informe")
    if not subject_filter:
        raise HTTPException(status_code=400, detail="Debe seleccionar una materia")

    teacher = teacher_profile(current_user)["teacher"]
    students = _teacher_course_students_for_report(
        current_user=current_user,
        period_codes=period_codes,
        subject_filter=subject_filter,
        parallel=parallel,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
    )
    selected_student_codes = {str(code) for code in (codigo_estud or [])}
    if selected_student_codes:
        students = [item for item in students if _clean(item.get("codigo_estud")) in selected_student_codes]
    meta = _teacher_course_report_meta(codigo_doc, period_codes, subject_filter, parallel, cod_anio_basica)
    if students:
        first = students[0]
        period_names: list[str] = []
        career_names: list[str] = []
        for item in students:
            period = _clean(item.get("detalle_periodo")) or _clean(item.get("codigo_periodo"))
            if period and period not in period_names:
                period_names.append(period)
            career = _clean(item.get("nombre_carrera"))
            if career and career not in career_names:
                career_names.append(career)
        meta = {
            **meta,
            "nombre_carrera": meta.get("nombre_carrera") or (" / ".join(career_names) if len(career_names) <= 2 else f"{len(career_names)} carreras"),
            "detalle_periodo": meta.get("detalle_periodo") or " / ".join(period_names),
            "codigo_materia": meta.get("codigo_materia") or _clean(first.get("codigo_materia")),
            "cod_materia": meta.get("cod_materia") or _clean(first.get("cod_materia")),
            "nombre_materia": meta.get("nombre_materia") or _clean(first.get("nombre_materia")),
            "paralelo": meta.get("paralelo") or _clean(first.get("paralelo")),
            "jornada": meta.get("jornada") or _clean(first.get("jornada")),
            "semestre": meta.get("semestre") or _int(first.get("semestre")),
            "horas": meta.get("horas") or _number(first.get("horas")),
        }
    report_format = _read_teacher_compliance_format()
    filename_stem = (
        f"cumplimiento-docente-{_safe_filename(meta.get('nombre_materia') or subject_filter)}-"
        f"{_safe_filename(meta.get('detalle_periodo') or '-'.join(str(code) for code in period_codes))}"
    )[:110]
    report_params = {
        "fecha_inicio": fecha_inicio,
        "fecha_fin": fecha_fin,
        "telefono": telefono,
        "actualizaciones": actualizaciones,
        "observaciones": observaciones,
        "teams_recordings": teams_recordings or [],
    }
    pdf_bytes = _teacher_compliance_report_pdf(
        teacher,
        meta,
        students,
        report_format,
        report_params,
        evidence_images=evidence_images,
    )
    return pdf_bytes, filename_stem


def _teacher_compliance_response(
    current_user: SessionUser,
    codigo_periodo: list[int],
    codigo_materia: str,
    paralelo: str,
    codigo_estud: list[int] | None = None,
    cod_anio_basica: int | None = None,
    cod_jornada: int | None = None,
    fecha_inicio: str = "",
    fecha_fin: str = "",
    telefono: str = "",
    actualizaciones: str = "",
    observaciones: str = "",
    teams_recordings: list[dict[str, Any]] | None = None,
    evidence_images: list[dict[str, Any]] | None = None,
) -> StreamingResponse:
    pdf_bytes, filename_stem = _build_teacher_compliance_pdf(
        current_user=current_user,
        codigo_periodo=codigo_periodo,
        codigo_materia=codigo_materia,
        paralelo=paralelo,
        codigo_estud=codigo_estud,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
        fecha_inicio=fecha_inicio,
        fecha_fin=fecha_fin,
        telefono=telefono,
        actualizaciones=actualizaciones,
        observaciones=observaciones,
        teams_recordings=teams_recordings,
        evidence_images=evidence_images,
    )
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename_stem}.pdf"'},
    )


async def _read_compliance_evidence(
    uploads: list[UploadFile] | None,
    labels: list[str] | None,
) -> list[dict[str, Any]]:
    evidence_images: list[dict[str, Any]] = []
    evidence_labels = labels or []
    for index, upload in enumerate(uploads or []):
        if not upload.filename:
            continue
        content_type = (upload.content_type or "").lower()
        if content_type and not content_type.startswith("image/"):
            raise HTTPException(status_code=400, detail="Las evidencias deben ser imágenes")
        content = await upload.read()
        if len(content) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Cada captura debe pesar máximo 5 MB")
        evidence_images.append(
            {
                "label": evidence_labels[index] if index < len(evidence_labels) else upload.filename,
                "content": content,
            }
        )
    return evidence_images


def _pdf_pages_as_compliance_evidence(
    pdf_bytes: bytes,
    *,
    label: str = "Captura reporte de notas firmado",
) -> list[dict[str, Any]]:
    if not pdf_bytes.startswith(b"%PDF"):
        raise HTTPException(status_code=400, detail="El reporte de notas adjunto no es un PDF válido")

    try:
        import pypdfium2 as pdfium

        document = pdfium.PdfDocument(pdf_bytes)
        try:
            page_count = len(document)
            if page_count < 1:
                raise HTTPException(status_code=400, detail="El reporte de notas no contiene páginas")
            if page_count > 40:
                raise HTTPException(
                    status_code=400,
                    detail="El reporte de notas supera el límite de 40 páginas",
                )

            # La apariencia visible de una firma PDF es un campo AcroForm. Sin
            # inicializar formularios, PDFium conserva el contenido del reporte
            # pero omite el QR y los datos visibles del firmante al rasterizarlo.
            document.init_forms()
            evidence_images: list[dict[str, Any]] = []
            for page_index in range(page_count):
                page = document[page_index]
                try:
                    bitmap = page.render(
                        scale=1.5,
                        may_draw_forms=True,
                        draw_annots=True,
                    )
                    try:
                        image = bitmap.to_pil().convert("RGB")
                        image_output = BytesIO()
                        image.save(image_output, format="PNG", optimize=True)
                    finally:
                        bitmap.close()
                finally:
                    page.close()
                evidence_images.append(
                    {
                        "label": f"{label} - página {page_index + 1} de {page_count}",
                        "content": image_output.getvalue(),
                    }
                )
            return evidence_images
        finally:
            document.close()
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("No se pudo convertir el reporte de notas firmado a imágenes")
        raise HTTPException(
            status_code=400,
            detail="No se pudo convertir el reporte de notas firmado a imágenes",
        ) from exc


async def _read_signed_grade_report_evidence(upload: UploadFile | None) -> list[dict[str, Any]]:
    if upload is None or not upload.filename:
        return []
    filename = _clean(upload.filename).lower()
    content_type = _clean(upload.content_type).lower()
    if not filename.endswith(".pdf") and content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="El reporte de notas firmado debe ser un archivo PDF")
    content = await upload.read()
    if not content:
        raise HTTPException(status_code=400, detail="El reporte de notas firmado está vacío")
    if len(content) > 25 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="El reporte de notas firmado debe pesar máximo 25 MB")
    return _pdf_pages_as_compliance_evidence(content)


def _certificate_subject_text(certificate: Any) -> str:
    subject = certificate.subject.native
    values: list[str] = []
    for value in subject.values():
        if isinstance(value, (list, tuple)):
            values.extend(_clean(item) for item in value if _clean(item))
        elif _clean(value):
            values.append(_clean(value))
    return " | ".join(values)


def _certificate_signer_name(certificate: Any, current_user: SessionUser) -> str:
    subject = certificate.subject.native
    candidate = subject.get("common_name") or subject.get("name") or current_user.nombres or current_user.login
    if isinstance(candidate, (list, tuple)):
        candidate = " ".join(_clean(value) for value in candidate if _clean(value))
    normalized = unicodedata.normalize("NFKD", _clean(candidate)).encode("ascii", "ignore").decode("ascii")
    return (normalized or "DOCENTE INTEC").upper()[:100]


def _validate_signing_certificate(certificate: Any) -> str:
    validity = certificate["tbs_certificate"]["validity"]
    valid_from = validity["not_before"].native
    valid_until = validity["not_after"].native
    now = datetime.now(timezone.utc)
    if valid_from.tzinfo is None:
        valid_from = valid_from.replace(tzinfo=timezone.utc)
    if valid_until.tzinfo is None:
        valid_until = valid_until.replace(tzinfo=timezone.utc)
    if now < valid_from:
        raise HTTPException(status_code=400, detail="El certificado todavía no se encuentra vigente")
    if now > valid_until:
        raise HTTPException(status_code=400, detail="El certificado de firma electrónica está caducado")

    subject_text = _certificate_subject_text(certificate)
    return subject_text


def _pkcs12_private_key_entries(pkcs12_bytes: bytes, password: bytes) -> list[tuple[Any, str]]:
    """Read every private key bag; cryptography's PKCS#12 loader returns only one."""
    try:
        pfx = asn1_pkcs12.Pfx.load(pkcs12_bytes)
        authenticated_safe = asn1_pkcs12.AuthenticatedSafe.load(pfx["auth_safe"]["content"].native)
    except (TypeError, ValueError) as exc:
        raise HTTPException(status_code=400, detail="El archivo .p12 no tiene una estructura PKCS#12 válida") from exc

    entries: list[tuple[Any, str]] = []

    def read_bags(safe_contents: Any) -> None:
        for bag in safe_contents:
            bag_type = bag["bag_id"].native
            friendly_name = ""
            for attribute in bag["bag_attributes"] or []:
                if attribute["type"].native == "friendly_name" and len(attribute["values"]):
                    friendly_name = _clean(attribute["values"][0].native)
                    break

            try:
                if bag_type == "pkcs8_shrouded_key_bag":
                    key = serialization.load_der_private_key(bag["bag_value"].untag().dump(), password=password)
                    entries.append((key, friendly_name))
                elif bag_type == "key_bag":
                    key = serialization.load_der_private_key(bag["bag_value"].untag().dump(), password=None)
                    entries.append((key, friendly_name))
                elif bag_type == "safe_contents":
                    read_bags(bag["bag_value"])
            except (TypeError, ValueError) as exc:
                raise HTTPException(
                    status_code=400,
                    detail="No se pudo abrir el archivo .p12. Verifique el archivo y la contraseña",
                ) from exc

    for content_info in authenticated_safe:
        if content_info["content_type"].native != "data":
            continue
        read_bags(asn1_pkcs12.SafeContents.load(content_info["content"].native))
    return entries


def _public_key_der(value: Any) -> bytes:
    public_key = value.public_key()
    return public_key.public_bytes(
        encoding=serialization.Encoding.DER,
        format=serialization.PublicFormat.SubjectPublicKeyInfo,
    )


def _certificate_signature_usage(certificate: crypto_x509.Certificate) -> tuple[bool, bool]:
    try:
        usage = certificate.extensions.get_extension_for_class(crypto_x509.KeyUsage).value
    except crypto_x509.ExtensionNotFound:
        return True, False
    may_sign = usage.digital_signature or usage.content_commitment
    return may_sign, usage.digital_signature


def _certificate_chain(
    leaf: crypto_x509.Certificate,
    certificates: list[crypto_x509.Certificate],
) -> list[crypto_x509.Certificate]:
    chain: list[crypto_x509.Certificate] = []
    current = leaf
    remaining = [certificate for certificate in certificates if certificate != leaf]
    while current.issuer != current.subject:
        issuer = next((certificate for certificate in remaining if certificate.subject == current.issuer), None)
        if issuer is None:
            break
        chain.append(issuer)
        remaining.remove(issuer)
        current = issuer
    return chain


def _load_digital_signature_pkcs12(pkcs12_bytes: bytes, password: str) -> signers.SimpleSigner:
    password_bytes = password.encode("utf-8")
    try:
        primary_key, primary_certificate, additional_certificates = crypto_pkcs12.load_key_and_certificates(
            pkcs12_bytes,
            password_bytes,
        )
    except (TypeError, ValueError) as exc:
        raise HTTPException(
            status_code=400,
            detail="No se pudo abrir el archivo .p12. Verifique el archivo y la contraseña",
        ) from exc

    certificates = [
        certificate
        for certificate in [primary_certificate, *(additional_certificates or [])]
        if certificate is not None
    ]
    key_entries = _pkcs12_private_key_entries(pkcs12_bytes, password_bytes)
    if primary_key is not None and all(_public_key_der(key) != _public_key_der(primary_key) for key, _ in key_entries):
        key_entries.append((primary_key, ""))

    candidates: list[tuple[int, Any, crypto_x509.Certificate]] = []
    for key, friendly_name in key_entries:
        key_public = _public_key_der(key)
        certificate = next(
            (candidate for candidate in certificates if _public_key_der(candidate) == key_public),
            None,
        )
        if certificate is None:
            continue
        may_sign, has_digital_signature = _certificate_signature_usage(certificate)
        if not may_sign:
            continue
        normalized_name = friendly_name.casefold()
        score = (100 if has_digital_signature else 60) + (40 if "signing" in normalized_name else 0)
        candidates.append((score, key, certificate))

    if not candidates:
        raise HTTPException(
            status_code=400,
            detail=(
                "El archivo .p12 no contiene una clave habilitada para firma digital. "
                "No se puede firmar con una clave destinada únicamente a cifrado"
            ),
        )

    _, signing_key, signing_certificate = max(candidates, key=lambda candidate: candidate[0])
    chain = _certificate_chain(signing_certificate, certificates)
    normalized_pkcs12 = crypto_pkcs12.serialize_key_and_certificates(
        name=b"FirmaDocente",
        key=signing_key,
        cert=signing_certificate,
        cas=chain,
        encryption_algorithm=serialization.BestAvailableEncryption(password_bytes),
    )
    signer = signers.SimpleSigner.load_pkcs12_data(
        normalized_pkcs12,
        other_certs=[],
        passphrase=password_bytes,
    )
    if signer is None or signer.signing_cert is None:
        raise HTTPException(status_code=400, detail="No se pudo preparar el certificado para firma digital")
    return signer


def _pdf_signature_target_above_text(
    pdf_bytes: bytes,
    marker: str,
    *,
    box_width: float = 126,
    box_height: float = 48,
    vertical_gap: float = 3,
    page_margin: float = 18,
    fallback_box: tuple[float, float, float, float] = (243, 90, 369, 138),
) -> tuple[int, tuple[float, float, float, float]]:
    """Locate the last marker occurrence and reserve a visible stamp above it."""
    try:
        import pypdfium2 as pdfium

        document = pdfium.PdfDocument(pdf_bytes)
        try:
            for page_index in range(len(document) - 1, -1, -1):
                page = document[page_index]
                try:
                    page_width, page_height = page.get_size()
                    text_page = page.get_textpage()
                    try:
                        search = text_page.search(marker, match_case=False, match_whole_word=False)
                        try:
                            occurrence = search.get_next()
                        finally:
                            search.close()
                        if not occurrence:
                            continue
                        start, count = occurrence
                        boxes = [text_page.get_charbox(index) for index in range(start, start + count)]
                    finally:
                        text_page.close()
                finally:
                    page.close()

                marker_left = min(box[0] for box in boxes)
                marker_top = max(box[3] for box in boxes)
                marker_right = max(box[2] for box in boxes)
                marker_center = (marker_left + marker_right) / 2
                left = max(page_margin, min(marker_center - box_width / 2, page_width - page_margin - box_width))
                bottom = marker_top + vertical_gap
                top = bottom + box_height
                if top > page_height - page_margin:
                    top = page_height - page_margin
                    bottom = top - box_height
                return page_index, (left, bottom, left + box_width, top)
        finally:
            document.close()
    except Exception:
        logger.exception("No se pudo localizar el espacio visible de firma en el PDF")
    return -1, fallback_box


def _signature_stamp_layout(
    signature_box: tuple[float, float, float, float],
) -> tuple[int, int, int]:
    """Scale the visible QR stamp to the signature area of each document."""
    left, bottom, right, top = signature_box
    width = right - left
    height = top - bottom
    if width <= 0 or height <= 0:
        raise ValueError("El rectángulo visible de firma no es válido")

    qr_size = int(max(20, min(52, height - 10, width * 0.28)))
    font_size = int(max(5, min(8, height / 9, width / 22)))
    separation = max(1, min(3, int(height / 20)))
    return qr_size, font_size, separation


def _signature_stamp_signer_name(
    signer_name: str,
    signature_box: tuple[float, float, float, float],
    qr_size: int,
    font_size: int,
) -> str:
    width = signature_box[2] - signature_box[0]
    available_text_width = max(40.0, width - qr_size - 8)
    line_length = max(12, int(available_text_width / max(font_size * 0.55, 1)))
    lines = textwrap.wrap(
        signer_name,
        width=line_length,
        break_long_words=False,
        break_on_hyphens=False,
    )
    return "\n".join(lines or [signer_name])


def _assert_pdf_signature_field(pdf_bytes: bytes, field_name: str) -> None:
    try:
        reader = PdfFileReader(BytesIO(pdf_bytes), strict=False)
        field_names = {signature.field_name for signature in reader.embedded_signatures}
    except Exception as exc:
        raise HTTPException(status_code=500, detail="No se pudo verificar el PDF firmado") from exc
    if field_name not in field_names:
        raise HTTPException(
            status_code=500,
            detail=f"El documento no contiene la firma electrónica esperada ({field_name})",
        )


async def _sign_pdf_with_pkcs12(
    pdf_bytes: bytes,
    pkcs12_bytes: bytes,
    password: str,
    current_user: SessionUser,
    reason: str,
    location: str,
    contact: str,
    signature_box: tuple[float, float, float, float] = (72, 450, 392, 520),
    signature_page: int = -1,
    field_name: str = "FirmaDocente",
    readable_field_name: str = "Firma electrónica del docente",
) -> bytes:
    signer = _load_digital_signature_pkcs12(pkcs12_bytes, password)

    _validate_signing_certificate(signer.signing_cert)
    signer_name = _certificate_signer_name(signer.signing_cert, current_user)
    qr_size, stamp_font_size, stamp_separation = _signature_stamp_layout(signature_box)
    visible_signer_name = _signature_stamp_signer_name(
        signer_name,
        signature_box,
        qr_size,
        stamp_font_size,
    )
    stamp_reason = _clean(reason) or "Informe de cumplimiento docente"
    stamp_text = (
        "Firmado electronicamente por:\n"
        f"{visible_signer_name}\n"
        "Validar unicamente con FirmaEC"
    )
    writer = IncrementalPdfFileWriter(BytesIO(pdf_bytes))
    metadata = signers.PdfSignatureMetadata(
        field_name=field_name,
        md_algorithm="sha512",
        location=_clean(location)[:120] or None,
        reason=stamp_reason[:200],
        contact_info=_clean(contact)[:200] or None,
        subfilter=SigSeedSubFilter.ADOBE_PKCS7_DETACHED,
    )
    pdf_signer = signers.PdfSigner(
        metadata,
        signer=signer,
        stamp_style=QRStampStyle(
            border_width=0,
            stamp_text=stamp_text,
            qr_inner_size=qr_size,
            innsep=stamp_separation,
            text_box_style=TextBoxStyle(
                font_size=stamp_font_size,
                leading=max(stamp_font_size + 1, int(round(stamp_font_size * 1.12))),
            ),
        ),
        new_field_spec=SigFieldSpec(
            sig_field_name=field_name,
            on_page=signature_page,
            box=signature_box,
            readable_field_name=readable_field_name,
        ),
    )
    output = BytesIO()
    try:
        await pdf_signer.async_sign_pdf(
            writer,
            output=output,
            appearance_text_params={"url": "https://www.firmadigital.gob.ec/"},
        )
    except Exception as exc:
        logger.exception("No se pudo firmar el informe de cumplimiento con el certificado PKCS#12")
        technical_detail = _clean(str(exc))[:240]
        detail = f"No se pudo aplicar la firma electrónica al PDF ({type(exc).__name__})"
        if technical_detail:
            detail = f"{detail}: {technical_detail}"
        raise HTTPException(status_code=400, detail=detail) from exc
    signed_pdf = output.getvalue()
    _assert_pdf_signature_field(signed_pdf, field_name)
    return signed_pdf


def _planning_paragraph(value: Any, style: ParagraphStyle) -> Paragraph:
    text = escape(_clean(value)).replace("\n", "<br/>") or "-"
    return Paragraph(text, style)


def _teacher_academic_planning_pdf(
    payload: AcademicPlanningPayload,
    teacher: dict[str, Any],
    meta: dict[str, Any],
) -> bytes:
    dark = colors.black
    red = colors.HexColor("#CC0000")
    pale = colors.HexColor("#EEDDDD")
    header_fill = colors.HexColor("#D9DADA")
    border = colors.HexColor("#888888")
    page_size = A4 if payload.document_type == "pea" else landscape(A4)
    page_width, page_height = page_size
    background_path = (
        _ACADEMIC_PLANNING_PEA_BACKGROUND_PATH
        if payload.document_type == "pea"
        else _ACADEMIC_PLANNING_SYLLABUS_BACKGROUND_PATH
    )
    content_width = 17.8 * cm if payload.document_type == "pea" else 26.6 * cm
    narrative_width = 17.8 * cm
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="PlanningTitle", parent=styles["Title"], fontSize=20, leading=23, alignment=TA_CENTER, textColor=dark, spaceAfter=4))
    styles.add(ParagraphStyle(name="PlanningHeading", parent=styles["Heading2"], fontSize=10, leading=12, textColor=dark, alignment=TA_CENTER, spaceBefore=9, spaceAfter=5))
    styles.add(ParagraphStyle(name="PlanningBody", parent=styles["BodyText"], fontSize=7.2, leading=9.2, textColor=dark, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name="PlanningCell", parent=styles["BodyText"], fontSize=6.4, leading=7.8, textColor=dark))
    styles.add(ParagraphStyle(name="PlanningCellCenter", parent=styles["PlanningCell"], alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="PlanningCellBold", parent=styles["PlanningCell"], fontName="Helvetica-Bold"))

    def p(value: Any, style: str = "PlanningBody") -> Paragraph:
        return _planning_paragraph(value, styles[style])

    def section(title: str, value: Any = None) -> None:
        story.append(p(title, "PlanningHeading"))
        if _clean(value):
            story.append(p(value))
            story.append(Spacer(1, 4))

    def boxed_section(title: str, value: Any) -> None:
        story.append(table([
            [p(title, "PlanningCellBold")],
            [p(value)],
        ], [narrative_width]))
        story.append(Spacer(1, 5))

    def structured_section(title: str, value: Any) -> None:
        raw_lines = [line.strip(" -•\t") for line in _clean(value).splitlines() if line.strip(" -•\t")]
        pairs: list[tuple[str, str]] = []
        for line in raw_lines:
            if ":" in line:
                label, detail = line.split(":", 1)
                pairs.append((label.strip(), detail.strip()))
            else:
                pairs.append(("", line))
        if not pairs:
            pairs = [("", "-")]
        rows: list[list[Any]] = [[p(title, "PlanningCellBold"), ""]]
        rows.extend([
            [p(label, "PlanningCellBold") if label else "", p(detail)]
            for label, detail in pairs
        ])
        result = table(rows, [narrative_width * 0.30, narrative_width * 0.70])
        result.setStyle(TableStyle([
            ("SPAN", (0, 0), (1, 0)),
            ("ALIGN", (0, 0), (1, 0), "CENTER"),
        ]))
        story.append(result)
        story.append(Spacer(1, 6))

    def table(
        data: list[list[Any]],
        widths: list[float],
        header_rows: int = 0,
        header_background: Any = None,
        h_align: str = "LEFT",
    ) -> Table:
        result = Table(data, colWidths=widths, repeatRows=header_rows, hAlign=h_align)
        commands: list[tuple[Any, ...]] = [
            ("GRID", (0, 0), (-1, -1), 0.45, border),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]
        if header_rows:
            commands.extend([
                ("FONTNAME", (0, 0), (-1, header_rows - 1), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, header_rows - 1), "CENTER"),
            ])
            if header_background is not None:
                commands.append(("BACKGROUND", (0, 0), (-1, header_rows - 1), header_background))
        result.setStyle(TableStyle(commands))
        return result

    total_docencia = sum(topic.horas_docencia for unit in payload.unidades for topic in unit.temas)
    total_practica = sum(topic.horas_practica for unit in payload.unidades for topic in unit.temas)
    total_autonomo = sum(topic.horas_autonomo for unit in payload.unidades for topic in unit.temas)
    document_label = "PEA" if payload.document_type == "pea" else "Silabo"
    subject = _clean(meta.get("nombre_materia") or meta.get("cod_materia") or payload.codigo_materia)
    career = _clean(meta.get("nombre_carrera"))
    period = _clean(meta.get("detalle_periodo") or meta.get("detalle_periodos") or ", ".join(str(code) for code in payload.codigo_periodos))
    teacher_name = _clean(teacher.get("docente")) or "DOCENTE INTEC"
    coordinator_name = "Roberto Castro"
    spanish_months = (
        "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
        "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE",
    )
    change_date = f"{spanish_months[payload.fecha_elaboracion.month - 1]} {payload.fecha_elaboracion.year}"
    semester = _int(meta.get("semestre"))
    level = f"{semester}.º semestre" if semester else (_clean(payload.nivel) or "-")
    curricular_unit = _clean(meta.get("unidad_curricular")) or _clean(payload.unidad_curricular) or "-"

    story: list[Flowable] = [Spacer(1, 2.15 * cm if payload.document_type == "pea" else 0.68 * cm)]
    story.append(p(f"{document_label} DE LA ASIGNATURA", "PlanningTitle"))
    story.append(p(subject.upper(), "PlanningTitle"))
    story.append(Spacer(1, 1.05 * cm))
    control_total = 15.0 * cm
    control_widths = [control_total * 0.42, control_total * 0.14, control_total * 0.24, control_total * 0.20]
    control_table = table([
        [p("CONTROL DE CAMBIOS", "PlanningCellBold"), "", "", ""],
        [p("Descripción", "PlanningCellBold"), p("Versión", "PlanningCellBold"), p("Responsable", "PlanningCellBold"), p("Fecha", "PlanningCellBold")],
        [p(f"Desarrollo de {document_label}", "PlanningCell"), p(payload.version, "PlanningCellCenter"), p(teacher_name.upper(), "PlanningCell"), p(change_date, "PlanningCellCenter")],
        ["", "", "", ""],
        ["", "", "", ""],
    ], control_widths, 2, h_align="CENTER")
    control_table.setStyle(TableStyle([
        ("SPAN", (0, 0), (-1, 0)),
        ("ALIGN", (0, 0), (-1, 1), "CENTER"),
    ]))
    story.append(control_table)
    story.append(Spacer(1, 8))
    if payload.document_type == "silabo":
        story.append(Indenter(left=1.0 * cm))
    overview_widths = [
        narrative_width * 0.18, narrative_width * 0.15,
        narrative_width * 0.18, narrative_width * 0.15,
        narrative_width * 0.18, narrative_width * 0.16,
    ]
    overview_table = table([
        [p("PROGRAMA DE ESTUDIOS DE ASIGNATURA - PEA", "PlanningCellBold"), "", "", "", "", ""],
        [p("Carrera:", "PlanningCellBold"), p(career), "", "", "", ""],
        [p("Datos Generales:", "PlanningCellBold"), "", "", "", "", ""],
        [p("Código de la asignatura", "PlanningCellBold"), p(meta.get("cod_materia") or payload.codigo_materia), p("Nombre de la asignatura", "PlanningCellBold"), p(subject), "", ""],
        [p("Nivel de la asignatura", "PlanningCellBold"), p(level), p("Unidad de organización curricular", "PlanningCellBold"), p(curricular_unit), p("Campo de formación", "PlanningCellBold"), p(payload.campo_formacion)],
        [p("Distribución de horas en las actividades de aprendizaje", "PlanningCellBold"), "", "", "", "", ""],
        [p("Docencia:", "PlanningCellBold"), p(total_docencia, "PlanningCellCenter"), p("Trabajo Autónomo", "PlanningCellBold"), p(total_autonomo, "PlanningCellCenter"), p("Prácticas Aprendizaje", "PlanningCellBold"), p(total_practica, "PlanningCellCenter")],
        [p("Práctica profesional", "PlanningCellBold"), p(0, "PlanningCellCenter"), p("Vinculación", "PlanningCellBold"), p(0, "PlanningCellCenter"), p("Trabajo de titulación", "PlanningCellBold"), p(0, "PlanningCellCenter")],
        [p("Periodo Académico", "PlanningCellBold"), p(period, "PlanningCellCenter"), "", p("Modalidad", "PlanningCellBold"), p(payload.modalidad, "PlanningCellCenter"), ""],
        [p("Prerrequisitos de la asignatura", "PlanningCellBold"), "", "", p("Co Requisitos de la asignatura", "PlanningCellBold"), "", ""],
        [p(payload.prerrequisitos), "", "", p(payload.correquisitos), "", ""],
        [p("Horario de clases", "PlanningCellBold"), "", "", p("Horario atención de tutorías", "PlanningCellBold"), "", ""],
        [p(payload.horario_clases), "", "", p(payload.horario_tutorias), "", ""],
    ], overview_widths)
    overview_table.setStyle(TableStyle([
        ("TOPPADDING", (0, 0), (-1, -1), 1.4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1.4),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("SPAN", (0, 0), (-1, 0)),
        ("SPAN", (1, 1), (-1, 1)),
        ("SPAN", (0, 2), (-1, 2)),
        ("SPAN", (3, 3), (-1, 3)),
        ("SPAN", (0, 5), (-1, 5)),
        ("SPAN", (1, 8), (2, 8)),
        ("SPAN", (4, 8), (5, 8)),
        ("SPAN", (0, 9), (2, 9)),
        ("SPAN", (3, 9), (5, 9)),
        ("SPAN", (0, 10), (2, 10)),
        ("SPAN", (3, 10), (5, 10)),
        ("SPAN", (0, 11), (2, 11)),
        ("SPAN", (3, 11), (5, 11)),
        ("SPAN", (0, 12), (2, 12)),
        ("SPAN", (3, 12), (5, 12)),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("ALIGN", (0, 5), (-1, 5), "LEFT"),
        ("ALIGN", (0, 9), (-1, 9), "CENTER"),
        ("ALIGN", (0, 11), (-1, 11), "CENTER"),
    ]))
    story.append(overview_table)
    if payload.document_type == "silabo":
        story.append(PageBreak())
    boxed_section("Descripción de la asignatura:", payload.descripcion)
    boxed_section("Objetivo General:", payload.objetivo_general)
    boxed_section("Resultados de Aprendizaje de la asignatura y como aporta al perfil profesional:", payload.resultados_aprendizaje)
    if payload.document_type == "silabo":
        section("ALINEAMIENTO CURRICULAR")
        story.append(PageBreak())
    else:
        section("ALINEAMIENTO CURRICULAR")
    story.append(table([
        [p("Misión INTEC", "PlanningCellBold"), p("Misión Escuela", "PlanningCellBold"), p("Misión Carrera", "PlanningCellBold")],
        [p(payload.mision_intec), p(payload.mision_escuela), p(payload.mision_carrera)],
    ], [narrative_width / 3, narrative_width / 3, narrative_width / 3], 1))
    if payload.document_type == "silabo":
        story.append(Indenter(left=-1.0 * cm))
    section("CONTENIDOS DE LA ASIGNATURA")

    class VerticalPlanningLabel(Flowable):
        def __init__(self, text: str) -> None:
            super().__init__()
            self.text = text
            self.width = 10
            self.height = 45

        def wrap(self, available_width: float, available_height: float) -> tuple[float, float]:
            return min(self.width, available_width), min(self.height, available_height)

        def draw(self) -> None:
            self.canv.saveState()
            self.canv.setFont("Helvetica", 5.4)
            self.canv.rotate(90)
            self.canv.drawString(0, -6, self.text)
            self.canv.restoreState()

    silabo_header_added = False
    for unit_index, unit in enumerate(payload.unidades, start=1):
        if payload.document_type == "pea":
            rows = [
                [p(f"UNIDAD {unit_index}: {unit.nombre}", "PlanningCellBold"), "", "", "", "", ""],
                [p("Resultado de Aprendizaje:", "PlanningCellBold"), p(unit.resultado_aprendizaje), "", "", "", ""],
                [p("Contenidos", "PlanningCellBold"), p("Horas de la Unidad", "PlanningCellBold"), "", "", p("Observaciones", "PlanningCellBold"), ""],
                ["", p("Docencia", "PlanningCellBold"), p("Prácticas", "PlanningCellBold"), p("Autónomo", "PlanningCellBold"), p("Trabajo Autónomo del estudiante", "PlanningCellBold"), p("Mecanismo de Evaluación", "PlanningCellBold")],
            ]
            rows.extend([
                [p(topic.tema), p(topic.horas_docencia, "PlanningCellCenter"), p(topic.horas_practica, "PlanningCellCenter"),
                 p(topic.horas_autonomo, "PlanningCellCenter"), p(topic.actividad_autonoma), p(topic.evaluacion)]
                for topic in unit.temas
            ])
            rows.append([
                p("Total", "PlanningCellBold"),
                p(sum(topic.horas_docencia for topic in unit.temas), "PlanningCellCenter"),
                p(sum(topic.horas_practica for topic in unit.temas), "PlanningCellCenter"),
                p(sum(topic.horas_autonomo for topic in unit.temas), "PlanningCellCenter"), "", "",
            ])
            unit_table = table(rows, [content_width * 0.30, content_width * 0.09, content_width * 0.09, content_width * 0.10, content_width * 0.23, content_width * 0.19], 4)
            unit_table.setStyle(TableStyle([
                ("SPAN", (0, 0), (-1, 0)),
                ("SPAN", (1, 1), (-1, 1)),
                ("SPAN", (0, 2), (0, 3)),
                ("SPAN", (1, 2), (3, 2)),
                ("SPAN", (4, 2), (5, 2)),
                ("ALIGN", (0, 0), (-1, 3), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(unit_table)
        else:
            widths = [
                content_width * 0.117, content_width * 0.169, content_width * 0.081,
                content_width * 0.043, content_width * 0.043, content_width * 0.043,
                content_width * 0.157, content_width * 0.131, content_width * 0.102,
                content_width * 0.114,
            ]
            if not silabo_header_added:
                header_rows = [
                    [p("UNIDAD"), p("TEMA"), p("SEMANA"), p("No. Horas"), "", "", p("Componente de docencia"), p("Componente de práctica de aplicación y experimentación de los aprendizajes"), p("Componente de aprendizaje autónomo"), p("Actividad calificada")],
                    ["", "", "", VerticalPlanningLabel("DOCENCIA"), VerticalPlanningLabel("PRÁCTICA"), VerticalPlanningLabel("AUTÓNOMO"), "", "", "", ""],
                ]
                header_table = Table(header_rows, colWidths=widths, rowHeights=[3.2 * cm, 1.5 * cm], hAlign="LEFT")
                header_table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.6, border),
                    ("BACKGROUND", (0, 0), (-1, -1), header_fill),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("SPAN", (0, 0), (0, 1)),
                    ("SPAN", (1, 0), (1, 1)),
                    ("SPAN", (2, 0), (2, 1)),
                    ("SPAN", (3, 0), (5, 0)),
                    ("SPAN", (6, 0), (6, 1)),
                    ("SPAN", (7, 0), (7, 1)),
                    ("SPAN", (8, 0), (8, 1)),
                    ("SPAN", (9, 0), (9, 1)),
                ]))
                story.append(header_table)
                story.append(PageBreak())
                silabo_header_added = True

            rows = [
                [p(f"UNIDAD {unit_index}: {unit.nombre}" if topic_index == 0 else ""), p(topic.tema), p(topic.semana, "PlanningCellCenter"), p(topic.horas_docencia, "PlanningCellCenter"),
                 p(topic.horas_practica, "PlanningCellCenter"), p(topic.horas_autonomo, "PlanningCellCenter"),
                 p(topic.actividad_docencia), p(topic.actividad_practica), p(topic.actividad_autonoma), p(topic.evaluacion)]
                for topic_index, topic in enumerate(unit.temas)
            ]
            if not rows:
                rows = [[
                    p(f"UNIDAD {unit_index}: {unit.nombre}"),
                    p("Pendiente de registrar"),
                    p("-", "PlanningCellCenter"),
                    p("0", "PlanningCellCenter"),
                    p("0", "PlanningCellCenter"),
                    p("0", "PlanningCellCenter"),
                    p("-"),
                    p("-"),
                    p("-"),
                    p("-"),
                ]]
            unit_table = table(
                rows,
                widths,
            )
            evaluation_commands: list[tuple[Any, ...]] = []
            for row_index, topic in enumerate(unit.temas):
                if "EVALU" in f"{topic.tema} {topic.evaluacion}".upper() or "PARCIAL" in f"{topic.tema} {topic.evaluacion}".upper():
                    evaluation_commands.extend([
                        ("BACKGROUND", (0, row_index), (-1, row_index), red),
                        ("TEXTCOLOR", (0, row_index), (-1, row_index), colors.white),
                    ])
            if evaluation_commands:
                unit_table.setStyle(TableStyle(evaluation_commands))
            story.append(unit_table)
        story.append(Spacer(1, 5))

    if payload.document_type == "silabo":
        story.append(Indenter(left=1.0 * cm))
    structured_section("Estrategias Metodológicas", payload.estrategias_metodologicas)
    structured_section("1. Formación ciudadana / Desarrollo de habilidades blandas", payload.formacion_ciudadana)
    structured_section("1.2. Educación ambiental / Desarrollo sostenible", payload.sostenibilidad)
    structured_section("Recursos Didácticos", payload.recursos_didacticos)
    section("EVALUACIÓN")
    if payload.document_type == "silabo":
        story.append(Indenter(left=-1.0 * cm))
        story.append(PageBreak())
    evaluation_rows: list[list[Any]] = [["", p("Actividad", "PlanningCellBold"), p("Peso", "PlanningCellBold")]]
    for partial in range(1, 4):
        evaluation_rows.extend([
            [p(f"PARCIAL {partial}", "PlanningCellBold"), p("Tareas"), p(f"{payload.evaluacion_tareas}%", "PlanningCellCenter")],
            ["", p("Trabajo individual"), p(f"{payload.evaluacion_individual}%", "PlanningCellCenter")],
            ["", p("Trabajo colaborativo"), p(f"{payload.evaluacion_colaborativo}%", "PlanningCellCenter")],
            ["", p("Evaluación acumulativa"), p(f"{payload.evaluacion_acumulativa}%", "PlanningCellCenter")],
        ])
    evaluation_table = table(
        evaluation_rows,
        [4.0 * cm, 5.0 * cm, 2.0 * cm],
        1,
        h_align="CENTER",
    )
    evaluation_style: list[tuple[Any, ...]] = []
    for start in (1, 5, 9):
        evaluation_style.extend([
            ("SPAN", (0, start), (0, start + 3)),
            ("BACKGROUND", (0, start), (0, start + 3), red),
            ("TEXTCOLOR", (0, start), (0, start + 3), colors.white),
            ("ALIGN", (0, start), (0, start + 3), "CENTER"),
            ("VALIGN", (0, start), (0, start + 3), "MIDDLE"),
        ])
    evaluation_table.setStyle(TableStyle(evaluation_style))
    story.append(evaluation_table)
    if payload.document_type == "silabo":
        story.append(PageBreak())
    section("BIBLIOGRAFÍA BÁSICA", payload.bibliografia_basica)
    section("BIBLIOGRAFÍA COMPLEMENTARIA", payload.bibliografia_complementaria)
    section("PROYECTO DE APLICACIÓN PRÁCTICA")
    story.append(table([
        [p("Tema", "PlanningCellBold"), p(payload.proyecto_tema)],
        [p("Tiempo", "PlanningCellBold"), p(payload.proyecto_tiempo)],
        [p("Objetivo", "PlanningCellBold"), p(payload.proyecto_objetivo)],
        [p("Contexto", "PlanningCellBold"), p(payload.proyecto_contexto)],
    ], [narrative_width * 0.18, narrative_width * 0.82], h_align="CENTER"))
    story.append(PageBreak())
    story.append(Spacer(1, 4.1 * cm if payload.document_type == "pea" else 2.2 * cm))
    signature_table = Table([
        [p("Elaborado por", "PlanningCellBold"), p("Revisado por", "PlanningCellBold")],
        ["", ""],
        [p("Cargo: Docente\nNombre: " + teacher_name + "\nFecha: " + payload.fecha_elaboracion.isoformat()),
         p("Cargo: Coordinador Académico\nNombre: " + coordinator_name + "\nFecha: " + payload.fecha_elaboracion.isoformat())],
    ], [5.3 * cm, 5.3 * cm], rowHeights=[0.8 * cm, 3.1 * cm, 1.8 * cm], hAlign="CENTER")
    signature_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, border),
        ("ALIGN", (0, 0), (-1, 0), "LEFT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(signature_table)

    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=page_size,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=4.3 * cm,
        bottomMargin=1.5 * cm,
        title=f"{document_label} - {subject}",
        author=teacher_name,
    )

    def page_header(canvas: Canvas, doc: Any) -> None:
        canvas.saveState()
        width, height = page_size

        if background_path.exists():
            canvas.drawImage(
                ImageReader(str(background_path)),
                0,
                0,
                width=width,
                height=height,
                preserveAspectRatio=False,
                mask="auto",
            )

        # El fondo procede de los documentos oficiales. Solo se reemplazan las
        # celdas variables de materia y paginacion, conservando sus bordes.
        header_x = 262.0
        subject_split = 447.5
        header_right = 553.6
        header_y = height - 117.6
        bottom_row_height = 33.5
        canvas.setFillColor(colors.white)
        canvas.rect(
            header_x + 0.8,
            header_y + 0.8,
            subject_split - header_x - 1.6,
            bottom_row_height - 1.6,
            fill=1,
            stroke=0,
        )
        canvas.rect(
            subject_split + 0.8,
            header_y + 0.8,
            header_right - subject_split - 1.6,
            bottom_row_height - 1.6,
            fill=1,
            stroke=0,
        )
        canvas.setFillColor(dark)
        canvas.setFont("Helvetica", 9.5)
        canvas.drawCentredString((header_x + subject_split) / 2, header_y + 13.0, subject[:42])
        canvas.drawCentredString((subject_split + header_right) / 2, header_y + 18.0, f"Página {canvas.getPageNumber()} de")
        if payload.document_type == "silabo":
            canvas.setStrokeColor(border)
            canvas.setLineWidth(0.6)
            canvas.line(header_x, header_y, header_right, header_y)
        canvas.restoreState()

    total_x = (447.5 + 553.6) / 2
    total_y = page_height - 114.0

    class PlanningNumberedCanvas(Canvas):
        def __init__(self, *args: Any, **kwargs: Any) -> None:
            super().__init__(*args, **kwargs)
            self._saved_page_states: list[dict[str, Any]] = []

        def showPage(self) -> None:
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self) -> None:
            page_count = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.setFillColor(dark)
                self.setFont("Helvetica-Bold", 9.5)
                self.drawCentredString(total_x, total_y, str(page_count))
                super().showPage()
            super().save()

    document.build(
        story,
        onFirstPage=page_header,
        onLaterPages=page_header,
        canvasmaker=PlanningNumberedCanvas,
    )
    return output.getvalue()


@router.post("/teacher/academic-planning-pdf")
def teacher_academic_planning_pdf(
    payload: AcademicPlanningPayload,
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    preview: Annotated[bool, Query()] = False,
) -> StreamingResponse:
    if not preview and payload.evaluacion_tareas + payload.evaluacion_individual + payload.evaluacion_colaborativo + payload.evaluacion_acumulativa != 100:
        raise HTTPException(status_code=400, detail="Los porcentajes de evaluación deben sumar 100%")
    if not preview and not any(unit.temas for unit in payload.unidades):
        raise HTTPException(status_code=400, detail="Debe registrar al menos un tema en la planificación")
    codigo_doc = _teacher_code(current_user)
    meta = _teacher_course_report_meta(
        codigo_doc,
        payload.codigo_periodos,
        payload.codigo_materia.strip().upper(),
        payload.paralelo.strip().upper(),
        payload.cod_anio_basica,
    )
    if not meta:
        raise HTTPException(status_code=404, detail="La asignatura seleccionada no está vinculada al docente autenticado")
    teacher = teacher_profile(current_user)["teacher"]
    pdf_bytes = _teacher_academic_planning_pdf(payload, teacher, meta)
    document_label = "pea" if payload.document_type == "pea" else "silabo"
    filename = f"{document_label}-{_safe_filename(meta.get('nombre_materia') or payload.codigo_materia)}.pdf"
    disposition = "inline" if preview else "attachment"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'{disposition}; filename="{filename}"', "Cache-Control": "no-store"},
    )


@router.post("/teacher/academic-planning-sign")
async def teacher_sign_academic_planning_pdf(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    payload_json: Annotated[str, Form()],
    certificado: Annotated[UploadFile, File()],
    contrasena_certificado: Annotated[str, Form()],
    firma_motivo: Annotated[str, Form()] = "Planificación académica docente",
    firma_ubicacion: Annotated[str, Form()] = "Quito",
    firma_contacto: Annotated[str, Form()] = "",
) -> StreamingResponse:
    try:
        payload = AcademicPlanningPayload.model_validate_json(payload_json)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="La información de planificación no es válida") from exc
    if payload.evaluacion_tareas + payload.evaluacion_individual + payload.evaluacion_colaborativo + payload.evaluacion_acumulativa != 100:
        raise HTTPException(status_code=400, detail="Los porcentajes de evaluación deben sumar 100%")
    if not any(unit.temas for unit in payload.unidades):
        raise HTTPException(status_code=400, detail="Debe registrar al menos un tema en la planificación")
    if not certificado.filename or not certificado.filename.lower().endswith((".p12", ".pfx")):
        raise HTTPException(status_code=400, detail="Seleccione un certificado PKCS#12 con extensión .p12 o .pfx")
    certificate_bytes = await certificado.read()
    if not certificate_bytes or len(certificate_bytes) > 2 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="El certificado está vacío o supera el máximo de 2 MB")
    codigo_doc = _teacher_code(current_user)
    meta = _teacher_course_report_meta(
        codigo_doc,
        payload.codigo_periodos,
        payload.codigo_materia.strip().upper(),
        payload.paralelo.strip().upper(),
        payload.cod_anio_basica,
    )
    if not meta:
        raise HTTPException(status_code=404, detail="La asignatura seleccionada no está vinculada al docente autenticado")
    teacher = teacher_profile(current_user)["teacher"]
    pdf_bytes = _teacher_academic_planning_pdf(payload, teacher, meta)
    signed_pdf = await _sign_pdf_with_pkcs12(
        pdf_bytes=pdf_bytes,
        pkcs12_bytes=certificate_bytes,
        password=contrasena_certificado,
        current_user=current_user,
        reason=firma_motivo,
        location=firma_ubicacion,
        contact=firma_contacto,
        signature_box=(150, 510, 290, 590) if payload.document_type == "pea" else (275, 315, 415, 395),
        field_name=f"FirmaDocente{payload.document_type.upper()}",
        readable_field_name=f"Firma electrónica docente del {payload.document_type.upper()}",
    )
    document_label = "pea" if payload.document_type == "pea" else "silabo"
    filename = f"{document_label}-{_safe_filename(meta.get('nombre_materia') or payload.codigo_materia)}-firmado.pdf"
    return StreamingResponse(
        BytesIO(signed_pdf),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"', "Cache-Control": "no-store"},
    )


@router.get("/teacher/course-report-pdf")
def teacher_course_report_pdf(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    codigo_periodo: Annotated[list[int], Query()],
    codigo_materia: Annotated[str, Query()],
    paralelo: Annotated[str, Query(min_length=1)],
    cod_anio_basica: Annotated[int | None, Query()] = None,
    cod_jornada: Annotated[int | None, Query()] = None,
) -> StreamingResponse:
    codigo_doc = _teacher_code(current_user)
    parallel = paralelo.strip().upper()
    subject_filter = _clean(codigo_materia).upper()
    period_codes = list(dict.fromkeys(codigo_periodo))
    if not period_codes:
        raise HTTPException(status_code=400, detail="Debe seleccionar al menos un periodo")
    if not subject_filter:
        raise HTTPException(status_code=400, detail="Debe seleccionar una materia")

    teacher = teacher_profile(current_user)["teacher"]
    students = _teacher_course_students_for_report(
        current_user=current_user,
        period_codes=period_codes,
        subject_filter=subject_filter,
        parallel=parallel,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
    )
    meta = _teacher_course_report_meta(codigo_doc, period_codes, subject_filter, parallel, cod_anio_basica)
    if students:
        first = students[0]
        period_names: list[str] = []
        career_names: list[str] = []
        for item in students:
            period = _clean(item.get("detalle_periodo")) or _clean(item.get("codigo_periodo"))
            if period and period not in period_names:
                period_names.append(period)
            career = _clean(item.get("nombre_carrera"))
            if career and career not in career_names:
                career_names.append(career)
        meta = {
            **meta,
            "nombre_carrera": meta.get("nombre_carrera") or (" / ".join(career_names) if len(career_names) <= 2 else f"{len(career_names)} carreras"),
            "detalle_periodo": meta.get("detalle_periodo") or " / ".join(period_names),
            "codigo_materia": meta.get("codigo_materia") or _clean(first.get("codigo_materia")),
            "cod_materia": meta.get("cod_materia") or _clean(first.get("cod_materia")),
            "nombre_materia": meta.get("nombre_materia") or _clean(first.get("nombre_materia")),
            "paralelo": meta.get("paralelo") or _clean(first.get("paralelo")),
            "jornada": meta.get("jornada") or _clean(first.get("jornada")),
            "semestre": meta.get("semestre") or _int(first.get("semestre")),
            "horas": meta.get("horas") or _number(first.get("horas")),
            "es_homologacion": meta.get("es_homologacion") or any(item.get("es_homologacion") for item in students),
        }
    pdf_bytes = _teacher_notes_report_pdf(teacher, meta, students)
    filename = (
        f"notas-docente-{_safe_filename(meta.get('nombre_materia') or subject_filter)}-"
        f"{_safe_filename(meta.get('detalle_periodo') or '-'.join(str(code) for code in period_codes))}.pdf"
    )
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _build_teacher_student_grade_report_pdf(
    current_user: SessionUser,
    codigo_periodo: list[int],
    codigo_materia: str,
    paralelo: str,
    codigo_estud: list[int] | None = None,
    cod_anio_basica: int | None = None,
    cod_jornada: int | None = None,
) -> tuple[bytes, str]:
    codigo_doc = _teacher_code(current_user)
    parallel = paralelo.strip().upper()
    subject_filter = _clean(codigo_materia).upper()
    period_codes = list(dict.fromkeys(codigo_periodo))
    if not period_codes:
        raise HTTPException(status_code=400, detail="Debe seleccionar al menos un periodo")
    if not subject_filter:
        raise HTTPException(status_code=400, detail="Debe seleccionar una materia")

    teacher = teacher_profile(current_user)["teacher"]
    students = _teacher_course_students_for_report(
        current_user=current_user,
        period_codes=period_codes,
        subject_filter=subject_filter,
        parallel=parallel,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
    )
    selected_student_codes = {str(code) for code in (codigo_estud or [])}
    if selected_student_codes:
        students = [item for item in students if _clean(item.get("codigo_estud")) in selected_student_codes]
    meta = _teacher_course_report_meta(codigo_doc, period_codes, subject_filter, parallel, cod_anio_basica)
    if students:
        first = students[0]
        period_names: list[str] = []
        career_names: list[str] = []
        for item in students:
            period = _clean(item.get("detalle_periodo")) or _clean(item.get("codigo_periodo"))
            if period and period not in period_names:
                period_names.append(period)
            career = _clean(item.get("nombre_carrera"))
            if career and career not in career_names:
                career_names.append(career)
        meta = {
            **meta,
            "nombre_carrera": " / ".join(career_names) if career_names else meta.get("nombre_carrera"),
            "detalle_periodo": " / ".join(period_names) if period_names else meta.get("detalle_periodo"),
            "codigo_materia": meta.get("codigo_materia") or _clean(first.get("codigo_materia")),
            "cod_materia": meta.get("cod_materia") or _clean(first.get("cod_materia")),
            "nombre_materia": meta.get("nombre_materia") or _clean(first.get("nombre_materia")),
            "paralelo": meta.get("paralelo") or _clean(first.get("paralelo")),
            "jornada": meta.get("jornada") or _clean(first.get("jornada")),
            "semestre": meta.get("semestre") or _int(first.get("semestre")),
            "horas": meta.get("horas") or _number(first.get("horas")),
            "es_homologacion": meta.get("es_homologacion") or any(item.get("es_homologacion") for item in students),
        }
    pdf_bytes = _student_grade_report_pdf(teacher, meta, students)
    filename_stem = (
        f"reporte-notas-secretaria-{_safe_filename(meta.get('nombre_materia') or subject_filter)}-"
        f"{_safe_filename(meta.get('detalle_periodo') or '-'.join(str(code) for code in period_codes))}"
    )
    return pdf_bytes, filename_stem


@router.get("/teacher/student-grade-report-pdf")
def teacher_student_grade_report_pdf(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    codigo_periodo: Annotated[list[int], Query()],
    codigo_materia: Annotated[str, Query()],
    paralelo: Annotated[str, Query(min_length=1)],
    codigo_estud: Annotated[list[int] | None, Query()] = None,
    cod_anio_basica: Annotated[int | None, Query()] = None,
    cod_jornada: Annotated[int | None, Query()] = None,
) -> StreamingResponse:
    pdf_bytes, filename_stem = _build_teacher_student_grade_report_pdf(
        current_user=current_user,
        codigo_periodo=codigo_periodo,
        codigo_materia=codigo_materia,
        paralelo=paralelo,
        codigo_estud=codigo_estud,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
    )
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename_stem}.pdf"'},
    )


@router.post("/teacher/student-grade-report-sign")
async def teacher_sign_student_grade_report(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    codigo_periodo: Annotated[list[int], Form()],
    codigo_materia: Annotated[str, Form()],
    paralelo: Annotated[str, Form(min_length=1)],
    certificado: Annotated[UploadFile, File()],
    contrasena_certificado: Annotated[str, Form(min_length=1, max_length=256)],
    firma_motivo: Annotated[str, Form(max_length=200)] = "Reporte de notas formato Secretaría",
    firma_ubicacion: Annotated[str, Form(max_length=120)] = "Quito, Ecuador",
    firma_contacto: Annotated[str, Form(max_length=200)] = "",
    codigo_estud: Annotated[list[int] | None, Form()] = None,
    cod_anio_basica: Annotated[int | None, Form()] = None,
    cod_jornada: Annotated[int | None, Form()] = None,
) -> StreamingResponse:
    certificate_name = _clean(certificado.filename).lower()
    if not certificate_name.endswith((".p12", ".pfx")):
        raise HTTPException(status_code=400, detail="Seleccione un certificado con extensión .p12 o .pfx")
    certificate_bytes = await certificado.read()
    if not certificate_bytes:
        raise HTTPException(status_code=400, detail="El archivo de certificado está vacío")
    if len(certificate_bytes) > 2 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="El archivo .p12 debe pesar máximo 2 MB")

    pdf_bytes, filename_stem = _build_teacher_student_grade_report_pdf(
        current_user=current_user,
        codigo_periodo=codigo_periodo,
        codigo_materia=codigo_materia,
        paralelo=paralelo,
        codigo_estud=codigo_estud,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
    )
    signature_page, signature_box = _pdf_signature_target_above_text(
        pdf_bytes,
        "Firma del docente",
    )
    signed_pdf = await _sign_pdf_with_pkcs12(
        pdf_bytes=pdf_bytes,
        pkcs12_bytes=certificate_bytes,
        password=contrasena_certificado,
        current_user=current_user,
        reason=firma_motivo,
        location=firma_ubicacion,
        contact=firma_contacto,
        signature_box=signature_box,
        signature_page=signature_page,
        field_name="FirmaDocenteReporteNotas",
        readable_field_name="Firma electrónica docente del reporte de notas",
    )
    return StreamingResponse(
        BytesIO(signed_pdf),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}-firmado.pdf"',
            "Cache-Control": "no-store",
        },
    )


@router.get("/teacher/compliance-report-docx")
@router.get("/teacher/compliance-report-pdf")
def teacher_compliance_report_pdf(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    codigo_periodo: Annotated[list[int], Query()],
    codigo_materia: Annotated[str, Query()],
    paralelo: Annotated[str, Query(min_length=1)],
    codigo_estud: Annotated[list[int] | None, Query()] = None,
    cod_anio_basica: Annotated[int | None, Query()] = None,
    cod_jornada: Annotated[int | None, Query()] = None,
    fecha_inicio: Annotated[str, Query(max_length=40)] = "",
    fecha_fin: Annotated[str, Query(max_length=40)] = "",
    telefono: Annotated[str, Query(max_length=40)] = "",
    actualizaciones: Annotated[str, Query(max_length=1000)] = "",
    observaciones: Annotated[str, Query(max_length=1000)] = "",
) -> StreamingResponse:
    return _teacher_compliance_response(
        current_user=current_user,
        codigo_periodo=codigo_periodo,
        codigo_materia=codigo_materia,
        paralelo=paralelo,
        codigo_estud=codigo_estud,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
        fecha_inicio=fecha_inicio,
        fecha_fin=fecha_fin,
        telefono=telefono,
        actualizaciones=actualizaciones,
        observaciones=observaciones,
    )


@router.post("/teacher/compliance-report-docx")
@router.post("/teacher/compliance-report-pdf")
async def teacher_compliance_report_pdf_with_evidence(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    codigo_periodo: Annotated[list[int], Form()],
    codigo_materia: Annotated[str, Form()],
    paralelo: Annotated[str, Form(min_length=1)],
    codigo_estud: Annotated[list[int] | None, Form()] = None,
    cod_anio_basica: Annotated[int | None, Form()] = None,
    cod_jornada: Annotated[int | None, Form()] = None,
    fecha_inicio: Annotated[str, Form(max_length=40)] = "",
    fecha_fin: Annotated[str, Form(max_length=40)] = "",
    telefono: Annotated[str, Form(max_length=40)] = "",
    actualizaciones: Annotated[str, Form(max_length=1000)] = "",
    observaciones: Annotated[str, Form(max_length=1000)] = "",
    teams_recordings_json: Annotated[str, Form(max_length=100000)] = "",
    evidencia_label: Annotated[list[str] | None, Form()] = None,
    evidencia: Annotated[list[UploadFile] | None, File()] = None,
    reporte_notas_firmado: Annotated[UploadFile | None, File()] = None,
) -> StreamingResponse:
    evidence_images = await _read_compliance_evidence(evidencia, evidencia_label)
    evidence_images.extend(await _read_signed_grade_report_evidence(reporte_notas_firmado))
    teams_recordings = _parse_teacher_teams_recordings(teams_recordings_json)
    return _teacher_compliance_response(
        current_user=current_user,
        codigo_periodo=codigo_periodo,
        codigo_materia=codigo_materia,
        paralelo=paralelo,
        codigo_estud=codigo_estud,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
        fecha_inicio=fecha_inicio,
        fecha_fin=fecha_fin,
        telefono=telefono,
        actualizaciones=actualizaciones,
        observaciones=observaciones,
        teams_recordings=teams_recordings,
        evidence_images=evidence_images,
    )


@router.post("/teacher/compliance-report-sign")
async def teacher_sign_compliance_report(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    codigo_periodo: Annotated[list[int], Form()],
    codigo_materia: Annotated[str, Form()],
    paralelo: Annotated[str, Form(min_length=1)],
    certificado: Annotated[UploadFile, File()],
    contrasena_certificado: Annotated[str, Form(min_length=1, max_length=256)],
    firma_motivo: Annotated[str, Form(max_length=200)] = "Informe de cumplimiento docente",
    firma_ubicacion: Annotated[str, Form(max_length=120)] = "Quito, Ecuador",
    firma_contacto: Annotated[str, Form(max_length=200)] = "",
    codigo_estud: Annotated[list[int] | None, Form()] = None,
    cod_anio_basica: Annotated[int | None, Form()] = None,
    cod_jornada: Annotated[int | None, Form()] = None,
    fecha_inicio: Annotated[str, Form(max_length=40)] = "",
    fecha_fin: Annotated[str, Form(max_length=40)] = "",
    telefono: Annotated[str, Form(max_length=40)] = "",
    actualizaciones: Annotated[str, Form(max_length=1000)] = "",
    observaciones: Annotated[str, Form(max_length=1000)] = "",
    teams_recordings_json: Annotated[str, Form(max_length=100000)] = "",
    evidencia_label: Annotated[list[str] | None, Form()] = None,
    evidencia: Annotated[list[UploadFile] | None, File()] = None,
    reporte_notas_firmado: Annotated[UploadFile | None, File()] = None,
) -> StreamingResponse:
    certificate_name = _clean(certificado.filename).lower()
    if not certificate_name.endswith((".p12", ".pfx")):
        raise HTTPException(status_code=400, detail="Seleccione un certificado con extensión .p12 o .pfx")
    certificate_bytes = await certificado.read()
    if not certificate_bytes:
        raise HTTPException(status_code=400, detail="El archivo de certificado está vacío")
    if len(certificate_bytes) > 2 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="El archivo .p12 debe pesar máximo 2 MB")

    evidence_images = await _read_compliance_evidence(evidencia, evidencia_label)
    evidence_images.extend(await _read_signed_grade_report_evidence(reporte_notas_firmado))
    teams_recordings = _parse_teacher_teams_recordings(teams_recordings_json)
    pdf_bytes, filename_stem = _build_teacher_compliance_pdf(
        current_user=current_user,
        codigo_periodo=codigo_periodo,
        codigo_materia=codigo_materia,
        paralelo=paralelo,
        codigo_estud=codigo_estud,
        cod_anio_basica=cod_anio_basica,
        cod_jornada=cod_jornada,
        fecha_inicio=fecha_inicio,
        fecha_fin=fecha_fin,
        telefono=telefono,
        actualizaciones=actualizaciones,
        observaciones=observaciones,
        teams_recordings=teams_recordings,
        evidence_images=evidence_images,
    )
    signature_page, signature_box = _pdf_signature_target_above_text(
        pdf_bytes,
        "Firma electrónica",
        box_width=190,
        box_height=58,
        fallback_box=(72, 458, 262, 516),
    )
    signed_pdf = await _sign_pdf_with_pkcs12(
        pdf_bytes=pdf_bytes,
        pkcs12_bytes=certificate_bytes,
        password=contrasena_certificado,
        current_user=current_user,
        reason=firma_motivo,
        location=firma_ubicacion,
        contact=firma_contacto,
        signature_box=signature_box,
        signature_page=signature_page,
    )
    return StreamingResponse(
        BytesIO(signed_pdf),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}-firmado.pdf"',
            "Cache-Control": "no-store",
        },
    )


@router.post("/teacher/signed-documents-archive")
async def teacher_signed_documents_archive(
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
    informe: Annotated[UploadFile, File()],
    notas: Annotated[UploadFile, File()],
    contrato: Annotated[UploadFile, File()],
    codigo_materia: Annotated[str, Form()] = "",
    nombre_materia: Annotated[str, Form()] = "",
    codigo_periodo: Annotated[list[str] | None, Form()] = None,
) -> StreamingResponse:
    compliance_pdf = await informe.read()
    grades_pdf = await notas.read()
    contract_pdf = await contrato.read()
    archive_bytes = _signed_teacher_documents_archive(
        compliance_pdf,
        grades_pdf,
        contract_pdf,
    )
    identity = _teacher_contract_identity(current_user)
    try:
        stored = await run_in_threadpool(
            _store_signed_teacher_documents_onedrive,
            identity=identity,
            compliance_pdf=compliance_pdf,
            grades_pdf=grades_pdf,
            contract_pdf=contract_pdf,
            archive_bytes=archive_bytes,
            subject_code=codigo_materia,
            subject_name=nombre_materia,
            period_codes=codigo_periodo or [],
        )
    except Exception as exc:
        logger.exception(
            "No se pudieron guardar los documentos firmados del docente %s en OneDrive.",
            _clean(identity.get("codigo_doc")),
        )
        raise HTTPException(
            status_code=502,
            detail=(
                "Los documentos fueron firmados, pero Microsoft OneDrive no confirmó su almacenamiento "
                "en la carpeta DOCENTES. Verifique la conexión de Microsoft Graph e intente nuevamente."
            ),
        ) from exc

    logger.info(
        "Documentos firmados del docente %s guardados en %s (%s archivos).",
        _clean(identity.get("codigo_doc")),
        _clean(stored.get("folder_path")),
        len(stored.get("items") or []),
    )
    return StreamingResponse(
        BytesIO(archive_bytes),
        media_type="application/zip",
        headers={
            "Content-Disposition": 'attachment; filename="documentos-docente-firmados.zip"',
            "Cache-Control": "no-store, private",
            "X-Content-Type-Options": "nosniff",
            "X-OneDrive-Saved": "true",
            "X-OneDrive-Root": _TEACHER_DOCUMENT_ONEDRIVE_ROOT,
            "X-OneDrive-Item-Count": str(len(stored.get("items") or [])),
        },
    )


@router.put("/teacher/grades")
def teacher_save_grades(
    payload: TeacherGradePayload,
    current_user: Annotated[SessionUser, Depends(_TEACHER_ACCESS)],
) -> dict[str, Any]:
    codigo_doc = _teacher_code(current_user)
    parallel = payload.paralelo.strip().upper()
    values = payload.model_dump()
    final_grade: float | None = None
    if values.get("teoria_homo") is not None or values.get("practica_homo") is not None:
        homologation_grade = calculate_homologation_grade_with_recovery(
            values.get("teoria_homo"),
            values.get("practica_homo"),
            values.get("recuperacion"),
        )
        final_grade = homologation_grade.final
    else:
        regular_grade = calculate_regular_grade_with_recovery(
            (
                (values.get("p1_tareas"), values.get("p1_proyectos"), values.get("p1_examen")),
                (values.get("p2_tareas"), values.get("p2_proyectos"), values.get("p2_examen")),
                (values.get("p3_tareas"), values.get("p3_proyectos"), values.get("p3_examen")),
            ),
            values.get("recuperacion"),
        )
        prom_p1, prom_p2, prom_p3 = regular_grade.partials
        if prom_p1 is not None:
            values["prom_p1"] = prom_p1
        if prom_p2 is not None:
            values["prom_p2"] = prom_p2
        if prom_p3 is not None:
            values["prom_p3"] = prom_p3
        final_grade = regular_grade.final

    if final_grade is not None:
        values["promedio"] = final_grade
        values["promedio_final"] = final_grade
        values["caprueba"] = "A" if final_grade >= 7 else "R"

    assignments: list[str] = []
    params: list[Any] = []
    for payload_key, column in _GRADE_COLUMN_MAP.items():
        value = values.get(payload_key)
        if value is not None:
            assignments.append(f"{column} = ?")
            params.append(value)
    if not assignments:
        raise HTTPException(status_code=400, detail="No hay notas para actualizar")

    assignments.append("Usuario = ?")
    params.append(str(codigo_doc)[:10])

    where_parts = [
        "TRY_CONVERT(int, cxe.codigo_estud) = ?",
        "TRY_CONVERT(int, cxe.cod_anio_Basica) = ?",
        "TRY_CONVERT(int, cxe.codigo_materia) = ?",
        "TRY_CONVERT(int, cxe.codigo_periodo) = ?",
        "UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxe.paralelo)))) = ?",
    ]
    where_params: list[Any] = [
        payload.codigo_estud,
        payload.cod_anio_basica,
        payload.codigo_materia,
        payload.codigo_periodo,
        parallel,
    ]
    if payload.num_matricula is not None:
        where_parts.append("TRY_CONVERT(int, cxe.Num_Matricula) = ?")
        where_params.append(payload.num_matricula)
    if payload.num_grupo is not None:
        where_parts.append("TRY_CONVERT(int, cxe.NumGrupo) = ?")
        where_params.append(payload.num_grupo)

    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                """
                SELECT COUNT(*)
                FROM dbo.CARRERAXDOCENTE cxd
                LEFT JOIN dbo.PENSUM assigned_pensum
                  ON TRY_CONVERT(int, assigned_pensum.Cod_AnioBasica) = TRY_CONVERT(int, cxd.cod_Anio_Basica)
                 AND TRY_CONVERT(int, assigned_pensum.codigo_materia) = TRY_CONVERT(int, cxd.codigo_materia)
                LEFT JOIN dbo.PENSUM target_pensum
                  ON TRY_CONVERT(int, target_pensum.Cod_AnioBasica) = ?
                 AND TRY_CONVERT(int, target_pensum.codigo_materia) = ?
                WHERE TRY_CONVERT(int, cxd.codigo_doc) = ?
                  AND TRY_CONVERT(int, cxd.cod_Anio_Basica) = ?
                  AND TRY_CONVERT(int, cxd.codigo_periodo) = ?
                  AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), cxd.Paralelo)))) = ?
                  AND UPPER(LTRIM(RTRIM(COALESCE(
                        NULLIF(TRY_CONVERT(nvarchar(100), assigned_pensum.cod_materia), N''),
                        TRY_CONVERT(nvarchar(100), assigned_pensum.codigo_materia),
                        TRY_CONVERT(nvarchar(100), cxd.codigo_materia),
                        N''
                  )))) = UPPER(LTRIM(RTRIM(COALESCE(
                        NULLIF(TRY_CONVERT(nvarchar(100), target_pensum.cod_materia), N''),
                        TRY_CONVERT(nvarchar(100), target_pensum.codigo_materia),
                        TRY_CONVERT(nvarchar(100), ?),
                        N''
                  ))))
                """,
                payload.cod_anio_basica,
                payload.codigo_materia,
                codigo_doc,
                payload.cod_anio_basica,
                payload.codigo_periodo,
                parallel,
                payload.codigo_materia,
            )
            if int(cursor.fetchone()[0] or 0) == 0:
                raise HTTPException(status_code=403, detail="El curso no esta asignado al docente actual")

            cursor.execute(
                f"""
                UPDATE cxe
                SET {', '.join(assignments)}
                FROM dbo.CARRERAXESTUD AS cxe
                INNER JOIN dbo.DATOS_ESTUD AS de_active
                  ON TRY_CONVERT(int, de_active.codigo_estud) = TRY_CONVERT(int, cxe.codigo_estud)
                WHERE {' AND '.join(where_parts)}
                  AND UPPER(LTRIM(RTRIM(TRY_CONVERT(nvarchar(50), de_active.Estado))))
                      IN (N'A', N'ACTIVO', N'ACTIVA')
                """,
                *params,
                *where_params,
            )
            affected = int(cursor.rowcount or 0)
            if affected != 1:
                raise HTTPException(
                    status_code=409,
                    detail="Solo se pueden actualizar notas de una matricula unica con estudiante activo",
                )
            conn.commit()
        return {"ok": True, "message": "Notas actualizadas", "affected_rows": affected}
    except HTTPException:
        raise
    except pyodbc.Error as exc:
        raise HTTPException(status_code=500, detail=f"Error guardando notas: {exc}") from exc
